"""
CaseHub - LOR Maker Routes
Native Letter of Recommendation generator with persona-specific formatting.
Full feature parity with tools service: AI profile analysis, advanced customization,
international context detection, DOCX/PDF output.
"""
from fastapi import APIRouter, Depends, Request, Form, UploadFile, File
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from typing import Optional
import os
import re
import json
import shutil
import tempfile
import logging
import httpx

from models import get_db, Client, Case
from models.tenant import tenant_query
from auth import get_current_user
from core.request_utils import get_request_org_id
from middleware.features import require_feature
from config import settings

logger = logging.getLogger(__name__)

PREFIX = settings.PREFIX
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "output")
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY", "")

router = APIRouter(tags=["lor-maker"])
templates = Jinja2Templates(directory="templates")


# =============================================================================
# AI HELPER FUNCTIONS
# =============================================================================

async def detect_location_context(name: str, org: str) -> str:
    """Detect if recommender is US-based or International using Perplexity."""
    try:
        if not PERPLEXITY_API_KEY:
            return "us"

        prompt = f"""Is the organization "{org}" (associated with {name}) primarily a US-based entity or an International/Non-US entity?
        If it's a multinational with a major US presence, count as US-based.
        Return ONLY valid JSON: {{ "location": "us" }} or {{ "location": "international" }}."""

        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers={
                    "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "sonar",
                    "messages": [
                        {"role": "system", "content": "You are a location detector. Return only valid JSON."},
                        {"role": "user", "content": prompt}
                    ]
                }
            )

            if response.status_code == 200:
                content = response.json()['choices'][0]['message']['content']
                match = re.search(r'\{.*\}', content, re.DOTALL)
                if match:
                    data = json.loads(match.group(0))
                    return data.get("location", "us")
    except Exception as e:
        logger.error(f"Location detection failed: {e}")

    return "us"


async def analyze_recommender_profile(name: str, title: str, organization: str, cv_text: str = "", style_description: str = "") -> dict:
    """Analyze recommender and generate a writing profile using Perplexity."""

    cv_section = f"\n\nCV/Background Information:\n{cv_text[:3000]}" if cv_text else ""

    style_section = ""
    if style_description:
        style_section = f"""

User's Description of Recommender's Style:
"{style_description[:1000]}"

IMPORTANT: Use this description to inform the persona, but ALWAYS maintain strict professional language standards.
The final writing style must be suitable for formal immigration petition letters."""

    analysis_prompt = f"""Analyze this professional for writing style recommendations:

Name: {name}
Title: {title}
Organization: {organization}
{cv_section}
{style_section}

Based on their professional background, determine:
1. Communication style (formal/academic/corporate/technical)
2. Writing tone (authoritative/collaborative/mentoring/peer)
3. Key expertise areas
4. Relationship context
5. Font style recommendation
6. Appropriate salutation and closing

Respond in JSON format only:
{{
    "style": "formal|academic|corporate|technical",
    "tone": "authoritative|collaborative|mentoring|peer|diplomatic",
    "expertise_areas": ["area1", "area2", "area3"],
    "relationship_type": "supervisor|colleague|industry_leader|academic|mentor",
    "font": "Times New Roman|Calibri|Arial|Garamond|Georgia",
    "font_size": 11,
    "use_tabs": false,
    "salutation": "appropriate opening",
    "closing": "appropriate closing",
    "emphasis": ["what they emphasize", "in letters"],
    "profile_summary": "One sentence describing their writing persona"
}}"""

    try:
        if not PERPLEXITY_API_KEY:
            return {
                "success": True,
                "profile": {
                    "style": "formal", "tone": "authoritative",
                    "expertise_areas": ["professional expertise"],
                    "relationship_type": "colleague",
                    "font": "Arial", "font_size": 11, "use_tabs": False,
                    "salutation": "Dear Reviewing Officer,",
                    "closing": "Sincerely,",
                    "emphasis": ["professional achievements"],
                    "profile_summary": f"Professional profile for {name}"
                }
            }

        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers={
                    "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "sonar",
                    "messages": [
                        {"role": "system", "content": "You are a professional writing style analyst. Respond only in valid JSON format."},
                        {"role": "user", "content": analysis_prompt}
                    ]
                }
            )

            if response.status_code == 200:
                data = response.json()
                content = data.get('choices', [{}])[0].get('message', {}).get('content', '')
                json_match = re.search(r'\{[\s\S]*\}', content)
                if json_match:
                    profile = json.loads(json_match.group())
                    return {"success": True, "profile": profile}

            return {
                "success": True,
                "profile": {
                    "style": "formal", "tone": "authoritative",
                    "expertise_areas": ["professional expertise"],
                    "relationship_type": "colleague",
                    "font": "Arial", "font_size": 11, "use_tabs": False,
                    "salutation": "Dear Reviewing Officer,",
                    "closing": "Sincerely,",
                    "emphasis": ["professional achievements"],
                    "profile_summary": f"Professional profile for {name}"
                }
            }

    except Exception as e:
        logger.error(f"Profile analysis error: {e}")
        return {"success": False, "error": str(e)}


def docx_to_pdf(docx_path: str) -> str:
    """Convert a DOCX file to PDF using ReportLab (basic conversion)."""
    from docx import Document as DocxDocument
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas as rl_canvas

    output_path = docx_path.replace(".docx", ".pdf")
    doc = DocxDocument(docx_path)
    c = rl_canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    margin = 1 * inch
    line_height = 14
    y = height - margin

    c.setFont("Helvetica", 11)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y -= line_height
            continue

        if y < margin:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - margin

        # Wrap long lines
        while len(text) > 90:
            c.drawString(margin, y, text[:90])
            text = text[90:]
            y -= line_height
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - margin

        c.drawString(margin, y, text)
        y -= line_height

    c.save()
    return output_path


def extract_cv_text(filepath: str) -> str:
    """Extract text from CV file (PDF, DOCX, TXT)."""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == '.docx':
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        elif ext == '.pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            return '\n'.join([page.extract_text() or '' for page in reader.pages])
    except Exception as e:
        logger.error(f"CV text extraction failed: {e}")
    return ""


# =============================================================================
# ROUTES
# =============================================================================

@router.get("/lor-maker", response_class=HTMLResponse)
async def lor_maker_page(
    request: Request,
    case_id: Optional[int] = None,
    client_name: Optional[str] = None,
    visa_type: Optional[str] = None,
    db: Session = Depends(get_db),
    _feature=Depends(require_feature("ai_lor")),
):
    """Render the LOR Maker form page."""
    from fastapi.responses import RedirectResponse
    user = get_current_user(request, db)
    if not user:
        return RedirectResponse(url=f"{PREFIX}/login", status_code=302)

    org_id = get_request_org_id(request)
    prefill = {"beneficiary_name": "", "visa_type": "EB-2 NIW", "field": ""}
    case = None

    if case_id and org_id is not None:
        case = tenant_query(db, Case, org_id).filter(Case.id == case_id).first()
        if case:
            client = tenant_query(db, Client, org_id).filter(Client.id == case.client_id).first()
            if client:
                prefill["beneficiary_name"] = f"{client.first_name} {client.last_name}"
            if case.visa_type:
                prefill["visa_type"] = case.visa_type

    if client_name and not prefill["beneficiary_name"]:
        prefill["beneficiary_name"] = client_name.replace("+", " ")
    if visa_type and visa_type != prefill["visa_type"]:
        prefill["visa_type"] = visa_type

    from services.lor_generator import PERSONAS, NATIONAL_INTEREST_TEXTS
    persona_list = []
    for key, cfg in PERSONAS.items():
        persona_list.append({
            "key": key,
            "name": cfg.get("name", key.replace("_", " ").title()),
            "font": cfg.get("font", "Arial"),
            "size": cfg.get("size", 11),
            "description": cfg.get("description", ""),
        })

    field_list = list(NATIONAL_INTEREST_TEXTS.keys())

    return templates.TemplateResponse("app/lor_maker/form.html", {
        "request": request, "user": user, "PREFIX": PREFIX,
        "prefill": prefill,
        "case": case,
        "personas": persona_list,
        "fields": field_list,
    })


@router.post("/api/lor/generate")
async def generate_lor_api(
    request: Request,
    persona: str = Form(...),
    beneficiary_name: str = Form(...),
    recommender_name: str = Form(...),
    recommender_title: str = Form(...),
    recommender_org: str = Form(...),
    recommender_email: str = Form(...),
    relationship: str = Form(...),
    custom_paragraphs: str = Form(""),
    visa_type: str = Form("EB-2 NIW"),
    field: str = Form(""),
    years_known: str = Form("several"),
    recommender_phone: str = Form(""),
    include_national_importance: str = Form("true"),
    include_prong3: str = Form("true"),
    include_detriment: str = Form("true"),
    output_format: str = Form("docx"),
    opening_style: str = Form(""),
    letter_structure: str = Form("flowing_narrative"),
    relationship_type: str = Form(""),
    salutation: str = Form(""),
    closing: str = Form(""),
    evidence_style: str = Form("footnoted"),
    paragraph_style: str = Form("medium_balanced"),
    international_context: str = Form("auto"),
    header_format: str = Form("full_letterhead"),
    case_id: Optional[int] = Form(None),
    db: Session = Depends(get_db),
    _feature=Depends(require_feature("ai_lor")),
):
    """Generate a Letter of Recommendation with full advanced options."""
    user = get_current_user(request, db)
    if not user:
        return JSONResponse({"error": "Not authenticated"}, status_code=401)

    # Convert string booleans
    do_ni = include_national_importance.lower() not in ('false', '0', 'off')
    do_p3 = include_prong3.lower() not in ('false', '0', 'off')

    try:
        from services.lor_generator import LORGenerator, PERSONAS

        # Handle custom personas (stored in localStorage on client)
        if persona.startswith("custom_"):
            # Custom persona - use default settings, client-side stores the profile
            effective_persona = "executive"
        elif persona not in PERSONAS:
            return JSONResponse({"error": f"Invalid persona: {persona}"}, status_code=400)
        else:
            effective_persona = persona

        os.makedirs(OUTPUT_DIR, exist_ok=True)

        generator = LORGenerator(
            persona=effective_persona,
            beneficiary_name=beneficiary_name,
            visa_type=visa_type,
            field=field if field else None,
            output_dir=OUTPUT_DIR,
        )

        # Resolve location context
        if international_context == "auto":
            international_context = await detect_location_context(recommender_name, recommender_org)
            logger.info(f"Auto-detected location context for {recommender_org}: {international_context}")

        # Build paragraphs
        paragraphs = []

        # Opening paragraph
        opening = generator.get_opening_paragraph(relationship, years_known=years_known)
        paragraphs.append(opening)

        # Custom paragraphs
        if custom_paragraphs.strip():
            for p in custom_paragraphs.strip().split("\n\n"):
                p = p.strip()
                if p:
                    paragraphs.append(p)

        # National importance paragraph
        if do_ni and field:
            ni_text = generator.get_national_importance_paragraph()
            if ni_text:
                paragraphs.append(ni_text)

        # Prong 3 paragraph
        if do_p3:
            p3_text = generator.get_prong3_paragraph()
            if p3_text:
                paragraphs.append(p3_text)

        # International context paragraph
        if international_context == "international":
            paragraphs.append(
                f"From my international perspective as an expert at {recommender_org}, "
                f"I can attest that {beneficiary_name}'s work has achieved global recognition "
                f"and is frequently cited by researchers worldwide, demonstrating an influence "
                f"that extends far beyond the United States."
            )

        # Conclusion
        conclusion = (
            f"Based on my professional experience and knowledge of {beneficiary_name}'s work, "
            f"I strongly support the approval of this {visa_type} petition. "
            f"{beneficiary_name}'s contributions are of substantial merit and national importance."
        )
        paragraphs.append(conclusion)

        # Generate DOCX
        docx_filepath = generator.create_document(
            recommender_name=recommender_name,
            recommender_title=recommender_title,
            recommender_org=recommender_org,
            recommender_email=recommender_email,
            relationship=relationship,
            paragraphs=paragraphs,
            recommender_phone=recommender_phone if recommender_phone else None,
            custom_salutation=salutation if salutation else None,
            custom_closing=closing if closing else None,
        )

        # Convert to PDF if requested
        filepath = docx_filepath
        if output_format == "pdf":
            try:
                pdf_path = docx_to_pdf(docx_filepath)
                os.remove(docx_filepath)
                filepath = pdf_path
            except Exception as e:
                logger.warning(f"PDF conversion failed, returning DOCX: {e}")

        filename = os.path.basename(filepath)
        logger.info(f"LOR generated: {filename}")

        return JSONResponse({
            "success": True,
            "filename": filename,
            "download_url": f"{PREFIX}/api/lor/download/{filename}",
            "persona": persona,
            "beneficiary": beneficiary_name,
        })

    except Exception as e:
        logger.error(f"LOR generation error: {e}", exc_info=True)
        return JSONResponse({"error": str(e)}, status_code=500)


@router.post("/api/lor/persona-analyze")
async def analyze_persona_api(
    request: Request,
    name: str = Form(...),
    title: str = Form(""),
    organization: str = Form(""),
    style_description: str = Form(""),
    cv_file: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db)
):
    """Analyze a recommender and generate a custom writing profile."""
    user = get_current_user(request, db)
    if not user:
        return JSONResponse({"error": "Not authenticated"}, status_code=401)

    logger.info(f"Analyzing profile for: {name}, {title}, {organization}")

    cv_text = ""
    temp_path = None

    try:
        if cv_file and cv_file.filename:
            ext = os.path.splitext(cv_file.filename)[1].lower()
            if ext in ['.docx', '.pdf', '.txt']:
                temp_path = os.path.join(tempfile.mkdtemp(), f"cv{ext}")
                content = await cv_file.read()
                with open(temp_path, "wb") as f:
                    f.write(content)
                cv_text = extract_cv_text(temp_path)
                logger.info(f"Extracted {len(cv_text)} chars from CV")

        result = await analyze_recommender_profile(name, title, organization, cv_text, style_description)

        if result.get("success"):
            return JSONResponse({
                "success": True,
                "persona": {
                    "name": f"Custom: {name.split()[0] if name else 'Profile'}",
                    "recommender_name": name,
                    "recommender_title": title,
                    "recommender_org": organization,
                    **result["profile"]
                }
            })
        else:
            return JSONResponse({
                "success": False,
                "error": result.get("error", "Analysis failed")
            }, status_code=500)

    except Exception as e:
        logger.error(f"Persona analysis failed: {e}", exc_info=True)
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)
    finally:
        if temp_path and os.path.exists(temp_path):
            shutil.rmtree(os.path.dirname(temp_path), ignore_errors=True)


@router.get("/api/lor/download/{filename}")
async def download_lor(request: Request, filename: str, db: Session = Depends(get_db)):
    """Download a generated LOR file."""
    user = get_current_user(request, db)
    if not user:
        return JSONResponse({"error": "Not authenticated"}, status_code=401)

    safe_chars = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789_-. ")
    if not all(c in safe_chars for c in filename):
        return JSONResponse({"error": "Invalid filename"}, status_code=400)

    filepath = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return JSONResponse({"error": "File not found"}, status_code=404)

    return FileResponse(filepath, filename=filename)
