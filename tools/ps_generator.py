#!/usr/bin/env python3
"""
ILC Personal Statement Generator
=================================

Generates professionally formatted Personal Statements for EB-2 NIW petitions.

MANDATORY STRUCTURE (5 sections):
I.   Overview of the Proposed Endeavor
II.  National Importance of the Endeavor (Prong 1)
III. Practical Impact and Innovation (Prong 1/2)
IV.  Why I Am Well-Positioned to Execute This Endeavor (Prong 2)
V.   Conclusion (with perjury declaration)

CRITICAL RULES:
- Times New Roman 12pt (standard for legal documents)
- 1 inch margins
- 1.5 or double spacing
- NO em-dash (-) ever
- MUST include perjury declaration at the end
- MUST cite government documents (EO, strategies, reports)
- NOT a resume - it's a NARRATIVE
"""

import os
from datetime import datetime
from typing import Optional, List, Dict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# GOVERNMENT DOCUMENTS BY FIELD
# =============================================================================

GOVERNMENT_DOCS = {
    "cybersecurity": [
        {
            "name": "Executive Order 14028",
            "title": "Improving the Nation's Cybersecurity",
            "date": "May 12, 2021",
            "citation": "86 Fed. Reg. 26633",
            "key_points": [
                "Modernize Federal Government cybersecurity",
                "Enhance software supply chain security",
                "Establish Cyber Safety Review Board",
            ],
        },
        {
            "name": "Executive Order 14144",
            "title": "Strengthening and Promoting Innovation in the Nation's Cybersecurity",
            "date": "January 16, 2025",
            "citation": "90 Fed. Reg. 4551",
            "key_points": [
                "Address software vulnerabilities",
                "Promote cybersecurity innovation",
                "Strengthen critical infrastructure protection",
            ],
        },
        {
            "name": "National Cybersecurity Strategy",
            "title": "National Cybersecurity Strategy",
            "date": "March 2023",
            "citation": "The White House",
            "key_points": [
                "Defend critical infrastructure",
                "Disrupt threat actors",
                "Shape market forces",
                "Invest in a resilient future",
                "Forge international partnerships",
            ],
        },
    ],
    "ai_ml": [
        {
            "name": "Executive Order 14110",
            "title": "Safe, Secure, and Trustworthy Development and Use of Artificial Intelligence",
            "date": "October 30, 2023",
            "citation": "88 Fed. Reg. 75191",
            "key_points": [
                "Establish new standards for AI safety and security",
                "Protect Americans' privacy",
                "Advance equity and civil rights",
                "Support workers",
                "Promote innovation and competition",
            ],
        },
        {
            "name": "National AI R&D Strategic Plan",
            "title": "National Artificial Intelligence Research and Development Strategic Plan: 2023 Update",
            "date": "May 2023",
            "citation": "The White House OSTP",
            "key_points": [
                "Make long-term investments in AI research",
                "Develop effective methods for human-AI collaboration",
                "Understand and address the ethical, legal, and societal implications of AI",
            ],
        },
    ],
    "clean_energy": [
        {
            "name": "Inflation Reduction Act",
            "title": "Inflation Reduction Act of 2022",
            "date": "August 16, 2022",
            "citation": "Pub. L. No. 117-169",
            "key_points": [
                "Largest investment in clean energy in American history",
                "Tax credits for renewable energy",
                "Support for domestic manufacturing",
            ],
        },
        {
            "name": "Critical Technologies List",
            "title": "Critical and Emerging Technologies List Update",
            "date": "February 2024",
            "citation": "Executive Office of the President",
            "key_points": [
                "Advanced clean energy technologies",
                "Energy storage",
                "Grid modernization",
            ],
        },
    ],
    "stem_education": [
        {
            "name": "Executive Order 14081",
            "title": "Advancing Biotechnology and Biomanufacturing Innovation",
            "date": "September 12, 2022",
            "citation": "87 Fed. Reg. 56849",
            "key_points": [
                "Strengthen STEM education workforce",
                "Support innovation in life sciences",
                "Build domestic manufacturing capacity",
            ],
        },
    ],
    "biotech": [
        {
            "name": "Executive Order on Biotechnology",
            "title": "Advancing Biotechnology and Biomanufacturing Innovation for a Sustainable, Safe, and Secure American Bioeconomy",
            "date": "September 12, 2022",
            "citation": "87 Fed. Reg. 56849",
            "key_points": [
                "Accelerate biotechnology innovation",
                "Grow the U.S. bioeconomy",
                "Drive the bioeconomy toward solutions",
            ],
        },
    ],
}


# =============================================================================
# SECTION TEMPLATES FOR CONTENT EXPANSION (Added 2026-01-04)
# =============================================================================
# These templates are used to generate rich content for each PS section
# Target: 2,000-4,000 words per PS (currently generating ~550)

PS_SECTION_TEMPLATES = {
    "section_1_overview": {
        "opening": [
            "My academic journey and professional experience to date have consistently focused on advancing {field}. During my {early_career}, I began developing expertise in {initial_focus}, which laid the foundation for my current work. Growing interest in {advanced_area} has since driven broader collaborations and innovations in {current_focus}.",
            "Throughout my career spanning {career_duration}, I have dedicated myself to advancing {field} through both research and practical application. My journey began at {early_institution} where I {early_achievement}, and has evolved into a comprehensive endeavor focused on {current_endeavor}.",
        ],
        "endeavor_components": [
            "My proposed endeavor in the United States has three main components. First, I will continue to {component_1}, building on the foundation I have established through {component_1_evidence}. Second, I will expand my work in {component_2}, leveraging my unique combination of {skills_combination}. Finally, I will contribute to {component_3}, ensuring knowledge transfer and workforce development in this critical field.",
            "The endeavor I propose consists of multiple interconnected objectives. My primary focus is {primary_focus}, which addresses {primary_challenge}. Additionally, I am committed to {secondary_focus}, which supports broader national objectives in {national_area}. These efforts are unified by my overarching goal of {overarching_goal}.",
        ],
        "current_role": [
            "I currently serve as {current_position} at {current_org}, where I lead {current_responsibilities}. My primary contributions include the development of {technical_contributions}, which are directly aligned with national priorities in {priority_areas}. This role allows me to apply my expertise in {expertise_areas} to address real-world challenges affecting {stakeholders}.",
            "In my present capacity as {current_position} at {current_org}, I oversee {scope_of_work}. My work involves {work_description}, which has resulted in {achievements}. Through this position, I have the platform to {future_impact}.",
        ],
    },
    "section_2_national_importance": {
        "government_alignment": [
            "The national importance of my endeavor is substantial. {government_doc_citation} establishes {field} as a national priority, citing the need for {national_need}. My work directly supports these objectives by {how_supports}. As stated in {additional_citation}, the United States must {must_do} to maintain its competitive position.",
            "My proposed endeavor aligns with federal priorities documented across multiple authoritative sources. {primary_citation} specifically identifies {priority_area} as critical to national interests. Furthermore, {secondary_citation} emphasizes the importance of {emphasis_area}. My contributions to {contribution_area} directly advance these established national goals.",
        ],
        "statistics": [
            "The scale of the challenge my work addresses is significant. {national_statistic_1}. Furthermore, {national_statistic_2}. These statistics underscore the urgent need for professionals who can {professional_need}. My endeavor directly contributes to addressing this challenge by {how_contributes}.",
            "Quantitative evidence demonstrates the national significance of this endeavor. According to {source_1}, {statistic_1}. {source_2} reports that {statistic_2}. The economic implications are substantial, with {economic_impact}. My work addresses these challenges through {approach}.",
        ],
        "twofold_benefit": [
            "The benefit of my endeavor to the United States is twofold. First, {benefit_1_detail}. This directly addresses {challenge_1}. Second, {benefit_2_detail}. By {how_second_benefit}, my work contributes to {broader_impact}. Together, these benefits support both the immediate needs of {immediate_beneficiaries} and the long-term strategic interests of the United States.",
        ],
    },
    "section_3_practical_impact": {
        "innovations": [
            "My work has produced tangible innovations with measurable impact. I developed {innovation_1}, which {innovation_1_impact}. This innovation has been {innovation_1_adoption}. Additionally, my work on {innovation_2} has resulted in {innovation_2_outcome}, benefiting {innovation_2_beneficiaries}.",
            "The practical applications of my expertise are demonstrated through multiple concrete achievements. {achievement_1_description}. This resulted in {achievement_1_outcome}. Furthermore, {achievement_2_description}. These innovations have been {adoption_status}.",
        ],
        "publications": [
            "I have disseminated my research through {publication_count} peer-reviewed publications in journals including {journal_names}. My work has been cited {citation_count} times, indicating its influence on the field. Key publications include {key_publication_1}, which {publication_1_impact}, and {key_publication_2}, which {publication_2_impact}.",
            "My scholarly contributions include publications in {journals_or_venues}. These works have contributed to the body of knowledge in {knowledge_area} by {contribution_type}. The recognition of this work is evidenced by {recognition_evidence}.",
        ],
        "recognition": [
            "My contributions have been recognized through {recognition_type}. I received {award_or_recognition_1} for {award_reason_1}. Additionally, {recognition_2_description}. This recognition confirms the value and impact of my work in {field}.",
        ],
    },
    "section_4_well_positioned": {
        "academic_credentials": [
            "I am uniquely qualified to advance this endeavor due to my combination of academic training and practical experience. I hold a {highest_degree} in {degree_field} from {institution}, where I focused on {research_focus}. During my doctoral research, I {doctoral_achievement}, which {doctoral_impact}. This academic foundation provides me with the theoretical knowledge necessary to {theoretical_application}.",
            "My academic credentials include {degree_1} from {institution_1} and {degree_2} from {institution_2}. My academic work focused on {academic_focus}, resulting in {academic_outcomes}. This rigorous training equipped me with {skills_acquired}.",
        ],
        "dual_role": [
            "My unique position at the intersection of {intersection_1} and {intersection_2} enables me to bridge the gap between {bridge_from} and {bridge_to}. Unlike professionals who focus exclusively on one aspect, I can {unique_capability}. This dual perspective is essential for {why_essential}.",
            "I possess the rare combination of {combination_1} and {combination_2}. While many professionals excel in {common_area}, few have demonstrated ability in {rare_area}. My background allows me to {allows_to_do}, which is precisely what is needed to {needed_for}.",
        ],
        "track_record": [
            "My track record demonstrates my ability to execute this endeavor successfully. At {organization_1}, I {accomplishment_1}. At {organization_2}, I {accomplishment_2}. These experiences confirm my capacity to {capacity_confirmed}. I have consistently delivered results that {results_description}.",
        ],
    },
    "section_5_conclusion": {
        "summary": [
            "In light of my track record of innovation, my strategic position at {current_org}, and the national importance of advancing {field}, I respectfully request approval of my EB-2 National Interest Waiver petition. My endeavor provides twofold benefits to the United States: it advances {benefit_1} and it strengthens {benefit_2}. Granting this waiver would allow me to continue this nationally significant work without interruption.",
            "For the reasons stated above, I submit that my proposed endeavor satisfies the three prongs established in Matter of Dhanasar. My work has substantial merit and national importance, I am well-positioned to advance the endeavor, and on balance, it would benefit the United States to waive the job offer and labor certification requirements. I am committed to contributing to {commitment_area} for the benefit of the American people.",
        ],
    },
}

# Minimum words per section to reach 2,000+ word target
PS_SECTION_MIN_WORDS = {
    "section_1_overview": 600,
    "section_2_national_importance": 500,
    "section_3_practical_impact": 400,
    "section_4_well_positioned": 400,
    "section_5_conclusion": 200,
}


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def set_font(run, font_name: str = "Times New Roman", font_size: int = 12):
    """Apply font styling to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)


def validate_content(text: str) -> List[str]:
    """Validate content. Returns list of issues."""
    issues = []

    # Check for em-dash
    if "\u2014" in text:
        issues.append("Content contains em-dash. Use regular dash (-) or rewrite.")

    return issues


# =============================================================================
# PERSONAL STATEMENT GENERATOR CLASS
# =============================================================================

class PSGenerator:
    """
    Generate Personal Statements for EB-2 NIW petitions.

    Example usage:
        generator = PSGenerator(
            beneficiary_name="John Doe",
            field="cybersecurity",
        )

        filepath = generator.create_document(
            sections={
                "overview": "My proposed endeavor...",
                "national_importance": "The United States...",
                "practical_impact": "My work has resulted in...",
                "well_positioned": "I am uniquely qualified...",
                "conclusion": "In summary...",
            }
        )
    """

    def __init__(
        self,
        beneficiary_name: str,
        field: str,
        output_dir: str = "output",
    ):
        self.beneficiary_name = beneficiary_name
        self.field = field
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_document(
        self,
        sections: Dict[str, str],
        date: Optional[str] = None,
        salutation: Optional[str] = None,
        closing: Optional[str] = None,
    ) -> str:
        """
        Create a complete Personal Statement document.

        Args:
            sections: Dictionary with keys:
                - overview: Section I content
                - national_importance: Section II content
                - practical_impact: Section III content
                - well_positioned: Section IV content
                - conclusion: Section V content (without perjury - added automatically)
            date: Optional date string, defaults to today

        Returns:
            Path to generated document
        """
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # =====================================================================
        # HEADER
        # =====================================================================
        date_para = doc.add_paragraph()
        run = date_para.add_run(date or datetime.now().strftime("%B %d, %Y"))
        set_font(run)
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()

        # Salutation
        sal = doc.add_paragraph()
        salutation_text = salutation if salutation else "Dear Immigration Official,"
        run = sal.add_run(salutation_text)
        set_font(run)

        doc.add_paragraph()

        # =====================================================================
        # INTRODUCTION
        # =====================================================================
        intro = doc.add_paragraph()
        intro_text = f"""I, {self.beneficiary_name}, respectfully submit this personal statement in support of my Petition for Alien of Exceptional Ability and National Interest Waiver (Form I-140), seeking classification under the employment-based second preference category (EB-2 NIW)."""
        run = intro.add_run(intro_text)
        set_font(run)
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        intro.paragraph_format.line_spacing = 1.5

        doc.add_paragraph()

        # =====================================================================
        # SECTION I: Overview of the Proposed Endeavor
        # =====================================================================
        self._add_section(
            doc,
            "I. Overview of the Proposed Endeavor",
            sections.get("overview", ""),
        )

        # =====================================================================
        # SECTION II: National Importance of the Endeavor
        # =====================================================================
        self._add_section(
            doc,
            "II. National Importance of the Endeavor",
            sections.get("national_importance", ""),
        )

        # =====================================================================
        # SECTION III: Practical Impact and Innovation
        # =====================================================================
        self._add_section(
            doc,
            "III. Practical Impact and Innovation",
            sections.get("practical_impact", ""),
        )

        # =====================================================================
        # SECTION IV: Why I Am Well-Positioned
        # =====================================================================
        self._add_section(
            doc,
            "IV. Why I Am Well-Positioned to Execute This Endeavor",
            sections.get("well_positioned", ""),
        )

        # =====================================================================
        # SECTION V: Conclusion
        # =====================================================================
        self._add_section(
            doc,
            "V. Conclusion",
            sections.get("conclusion", ""),
        )

        # =====================================================================
        # PERJURY DECLARATION (MANDATORY)
        # =====================================================================
        doc.add_paragraph()
        perjury = doc.add_paragraph()
        run = perjury.add_run("I declare this under penalty of perjury under the laws of the United States of America.")
        run.bold = True
        set_font(run)
        perjury.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # =====================================================================
        # SIGNATURE
        # =====================================================================
        doc.add_paragraph()
        doc.add_paragraph()

        closing_para = doc.add_paragraph()
        closing_text = closing if closing else "Sincerely,"
        run = closing_para.add_run(closing_text)
        set_font(run)

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        sig = doc.add_paragraph()
        run = sig.add_run(self.beneficiary_name)
        run.bold = True
        set_font(run)

        # =====================================================================
        # SAVE
        # =====================================================================
        safe_name = "".join(c for c in self.beneficiary_name if c.isalnum() or c in " _-").replace(" ", "_")
        filename = f"Personal_Statement_{safe_name}.docx"
        filepath = os.path.join(self.output_dir, filename)

        doc.save(filepath)
        print(f"Generated: {filepath}")
        return filepath

    def _add_section(self, doc: Document, title: str, content: str):
        """Add a numbered section to the document."""
        # Validate content
        issues = validate_content(content)
        if issues:
            print(f"WARNING: {issues}")

        # Section title
        doc.add_paragraph()
        title_para = doc.add_paragraph()
        run = title_para.add_run(title)
        run.bold = True
        set_font(run)

        doc.add_paragraph()

        # Section content - split by double newlines for paragraphs
        paragraphs = content.split("\n\n") if content else ["[Content to be added]"]

        for para_text in paragraphs:
            if not para_text.strip():
                continue

            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(para_text.strip())
            set_font(run)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5

    def get_government_docs_paragraph(self) -> str:
        """Generate paragraph citing relevant government documents."""
        if self.field not in GOVERNMENT_DOCS:
            return ""

        docs = GOVERNMENT_DOCS[self.field]

        paragraphs = []
        for doc in docs[:2]:  # Cite first 2 documents
            key_points_text = ", ".join(doc["key_points"][:3])
            para = f"The {doc['name']}, \"{doc['title']}\" ({doc['date']}), identifies key priorities that are directly supported by my work. These include: {key_points_text}."
            paragraphs.append(para)

        return "\n\n".join(paragraphs)

    def get_section_template(self, section: str) -> str:
        """Get a template for a specific section."""
        templates = {
            "overview": f"""My proposed endeavor in the United States is to continue advancing [SPECIFIC FIELD/TECHNOLOGY].

Building on my role as [CURRENT POSITION] at [ORGANIZATION], I will further develop and implement [SPECIFIC TECHNOLOGIES/METHODS]. These innovations directly address critical challenges in [AREA], aiming to [SPECIFIC GOALS].

In parallel, I plan to expand the applicability of my work to related domains such as [RELATED AREAS].""",

            "national_importance": f"""The proposed endeavor possesses substantial merit and national importance.

{self.get_government_docs_paragraph()}

My work directly aligns with these national priorities by [SPECIFIC ALIGNMENT].""",

            "practical_impact": f"""The real-world impact of my work is already evident.

In [DATE], [SPECIFIC PROJECT/SYSTEM] successfully [SPECIFIC ACHIEVEMENT] using my [CONTRIBUTION]. These results demonstrated [MEASURABLE OUTCOME].

My methodology has been adopted by [OTHER ORGANIZATIONS/TEAMS] for [SPECIFIC USES], validating its practical value and scalability.""",

            "well_positioned": f"""What uniquely qualifies me to lead this endeavor is the combination of [EXPERTISE 1] and [EXPERTISE 2].

My academic background in [FIELD], combined with [NUMBER] years of industry experience developing [TYPE OF WORK], has equipped me to solve challenges that have limited progress in the field.

I carry critical tacit design knowledge from [SPECIFIC EXPERIENCES] that ensures [SPECIFIC BENEFIT].""",

            "conclusion": f"""In light of my academic background, track record of professional accomplishments, and sustained commitment to innovation and public-interest-driven work, I respectfully submit this personal statement in support of my EB-2 National Interest Waiver petition.

My continued work in the United States will [SPECIFIC BENEFIT TO US]. I remain committed to advancing [FIELD] and contributing to American leadership in [AREA].""",
        }

        return templates.get(section, "[Section content to be added]")

    def generate_full_ps(
        self,
        user_inputs: Dict[str, str],
        background_info: Optional[Dict[str, str]] = None,
    ) -> Dict[str, str]:
        """
        Generate a complete PS with expanded content.
        Target: 2,000-4,000 words.

        Args:
            user_inputs: Dictionary with user-provided content for each section
            background_info: Optional dict with beneficiary background details

        Returns:
            Dictionary with expanded section content ready for create_document()
        """
        background_info = background_info or {}
        sections = {}

        # Section I: Overview
        sections["overview"] = self._expand_overview(
            user_inputs.get("overview", ""),
            background_info
        )

        # Section II: National Importance
        sections["national_importance"] = self._expand_national_importance(
            user_inputs.get("national_importance", ""),
            background_info
        )

        # Section III: Practical Impact
        sections["practical_impact"] = self._expand_practical_impact(
            user_inputs.get("practical_impact", ""),
            background_info
        )

        # Section IV: Well-Positioned
        sections["well_positioned"] = self._expand_well_positioned(
            user_inputs.get("well_positioned", ""),
            background_info
        )

        # Section V: Conclusion
        sections["conclusion"] = self._expand_conclusion(
            user_inputs.get("conclusion", ""),
            background_info
        )

        # Validate total word count
        total_words = sum(len(s.split()) for s in sections.values())
        if total_words < 2000:
            sections = self._add_elaboration_to_ps(sections, 2100)

        return sections

    def _expand_overview(self, user_content: str, info: Dict[str, str]) -> str:
        """Expand Section I: Overview of the Proposed Endeavor."""
        paragraphs = []

        # Opening paragraph
        templates = PS_SECTION_TEMPLATES.get("section_1_overview", {}).get("opening", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "field": self.field or "this field",
                "early_career": info.get("early_career", "early career"),
                "initial_focus": info.get("initial_focus", "foundational work in this area"),
                "advanced_area": info.get("advanced_area", "advanced applications"),
                "current_focus": info.get("current_focus", "cutting-edge developments"),
                "career_duration": info.get("career_duration", "over a decade"),
                "early_institution": info.get("early_institution", "leading institutions"),
                "early_achievement": info.get("early_achievement", "developed foundational skills"),
                "current_endeavor": info.get("current_endeavor", "advancing national priorities"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        # User content (priority)
        if user_content:
            paragraphs.append(user_content)

        # Endeavor components
        templates = PS_SECTION_TEMPLATES.get("section_1_overview", {}).get("endeavor_components", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "component_1": info.get("component_1", "continue my current research and development"),
                "component_1_evidence": info.get("component_1_evidence", "my track record of innovation"),
                "component_2": info.get("component_2", "expand applications to related domains"),
                "skills_combination": info.get("skills_combination", "technical expertise and practical experience"),
                "component_3": info.get("component_3", "knowledge transfer and workforce development"),
                "primary_focus": info.get("primary_focus", "advancing this critical field"),
                "primary_challenge": info.get("primary_challenge", "key challenges facing the nation"),
                "secondary_focus": info.get("secondary_focus", "disseminating best practices"),
                "national_area": info.get("national_area", "national competitiveness"),
                "overarching_goal": info.get("overarching_goal", "contributing to American leadership"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        # Current role
        templates = PS_SECTION_TEMPLATES.get("section_1_overview", {}).get("current_role", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "current_position": info.get("current_position", "my current role"),
                "current_org": info.get("current_org", "my organization"),
                "current_responsibilities": info.get("current_responsibilities", "key initiatives"),
                "technical_contributions": info.get("technical_contributions", "innovative solutions"),
                "priority_areas": info.get("priority_areas", self.field or "this critical area"),
                "expertise_areas": info.get("expertise_areas", "my areas of specialization"),
                "stakeholders": info.get("stakeholders", "key stakeholders"),
                "scope_of_work": info.get("scope_of_work", "strategic initiatives"),
                "work_description": info.get("work_description", "developing and implementing solutions"),
                "achievements": info.get("achievements", "measurable improvements"),
                "future_impact": info.get("future_impact", "continue driving innovation"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        return "\n\n".join(paragraphs) if paragraphs else user_content

    def _expand_national_importance(self, user_content: str, info: Dict[str, str]) -> str:
        """Expand Section II: National Importance."""
        paragraphs = []

        # Government document citations
        gov_para = self.get_government_docs_paragraph()
        if gov_para:
            paragraphs.append(gov_para)

        # User content
        if user_content:
            paragraphs.append(user_content)

        # Statistics template
        templates = PS_SECTION_TEMPLATES.get("section_2_national_importance", {}).get("statistics", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "national_statistic_1": info.get("statistic_1", "The challenge in this field is significant"),
                "national_statistic_2": info.get("statistic_2", "Demand for qualified professionals continues to grow"),
                "professional_need": info.get("professional_need", "address these critical challenges"),
                "how_contributes": info.get("how_contributes", "developing innovative solutions"),
                "source_1": info.get("source_1", "industry reports"),
                "statistic_1": info.get("stat_detail_1", "the need is substantial"),
                "source_2": info.get("source_2", "government data"),
                "statistic_2": info.get("stat_detail_2", "growth continues"),
                "economic_impact": info.get("economic_impact", "billions of dollars at stake"),
                "approach": info.get("approach", "innovative methodologies"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        # Twofold benefit
        templates = PS_SECTION_TEMPLATES.get("section_2_national_importance", {}).get("twofold_benefit", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "benefit_1_detail": info.get("benefit_1", "my work advances the field technically"),
                "challenge_1": info.get("challenge_1", "critical gaps in current capabilities"),
                "benefit_2_detail": info.get("benefit_2", "it strengthens American competitiveness"),
                "how_second_benefit": info.get("how_second_benefit", "developing domestic expertise"),
                "broader_impact": info.get("broader_impact", "national strategic interests"),
                "immediate_beneficiaries": info.get("immediate_beneficiaries", "organizations and individuals"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        return "\n\n".join(paragraphs) if paragraphs else user_content

    def _expand_practical_impact(self, user_content: str, info: Dict[str, str]) -> str:
        """Expand Section III: Practical Impact and Innovation."""
        paragraphs = []

        # User content first (priority)
        if user_content:
            paragraphs.append(user_content)

        # Innovations template
        templates = PS_SECTION_TEMPLATES.get("section_3_practical_impact", {}).get("innovations", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "innovation_1": info.get("innovation_1", "innovative solutions"),
                "innovation_1_impact": info.get("innovation_1_impact", "significantly improved outcomes"),
                "innovation_1_adoption": info.get("innovation_1_adoption", "adopted by leading organizations"),
                "innovation_2": info.get("innovation_2", "advanced methodologies"),
                "innovation_2_outcome": info.get("innovation_2_outcome", "measurable improvements"),
                "innovation_2_beneficiaries": info.get("innovation_2_beneficiaries", "key stakeholders"),
                "achievement_1_description": info.get("achievement_1_description", "I developed key innovations"),
                "achievement_1_outcome": info.get("achievement_1_outcome", "significant improvements"),
                "achievement_2_description": info.get("achievement_2_description", "my work was recognized"),
                "adoption_status": info.get("adoption_status", "widely adopted"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        # Publications template
        templates = PS_SECTION_TEMPLATES.get("section_3_practical_impact", {}).get("publications", [])
        if templates and info.get("has_publications", True):
            import random
            template = random.choice(templates)
            defaults = {
                "publication_count": info.get("publication_count", "multiple"),
                "journal_names": info.get("journal_names", "peer-reviewed journals"),
                "citation_count": info.get("citation_count", "numerous"),
                "key_publication_1": info.get("key_publication_1", "my primary research"),
                "publication_1_impact": info.get("publication_1_impact", "advanced the field"),
                "key_publication_2": info.get("key_publication_2", "my follow-up work"),
                "publication_2_impact": info.get("publication_2_impact", "expanded applications"),
                "journals_or_venues": info.get("journals_or_venues", "leading venues"),
                "knowledge_area": info.get("knowledge_area", self.field or "this field"),
                "contribution_type": info.get("contribution_type", "providing new insights"),
                "recognition_evidence": info.get("recognition_evidence", "citations and adoption"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        return "\n\n".join(paragraphs) if paragraphs else user_content

    def _expand_well_positioned(self, user_content: str, info: Dict[str, str]) -> str:
        """Expand Section IV: Why I Am Well-Positioned."""
        paragraphs = []

        # Academic credentials
        templates = PS_SECTION_TEMPLATES.get("section_4_well_positioned", {}).get("academic_credentials", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "highest_degree": info.get("highest_degree", "advanced degree"),
                "degree_field": info.get("degree_field", self.field or "this field"),
                "institution": info.get("institution", "a leading institution"),
                "research_focus": info.get("research_focus", "key research areas"),
                "doctoral_achievement": info.get("doctoral_achievement", "developed novel approaches"),
                "doctoral_impact": info.get("doctoral_impact", "advanced understanding in the field"),
                "theoretical_application": info.get("theoretical_application", "address complex challenges"),
                "degree_1": info.get("degree_1", "my first degree"),
                "institution_1": info.get("institution_1", "a respected institution"),
                "degree_2": info.get("degree_2", "my advanced degree"),
                "institution_2": info.get("institution_2", "a leading university"),
                "academic_focus": info.get("academic_focus", "critical areas"),
                "academic_outcomes": info.get("academic_outcomes", "significant contributions"),
                "skills_acquired": info.get("skills_acquired", "essential capabilities"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        # User content
        if user_content:
            paragraphs.append(user_content)

        # Dual role template
        templates = PS_SECTION_TEMPLATES.get("section_4_well_positioned", {}).get("dual_role", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "intersection_1": info.get("intersection_1", "academic research"),
                "intersection_2": info.get("intersection_2", "industry practice"),
                "bridge_from": info.get("bridge_from", "theoretical knowledge"),
                "bridge_to": info.get("bridge_to", "practical implementation"),
                "unique_capability": info.get("unique_capability", "translate research into real-world solutions"),
                "why_essential": info.get("why_essential", "advancing this field effectively"),
                "combination_1": info.get("combination_1", "deep technical expertise"),
                "combination_2": info.get("combination_2", "practical experience"),
                "common_area": info.get("common_area", "technical skills"),
                "rare_area": info.get("rare_area", "implementation at scale"),
                "allows_to_do": info.get("allows_to_do", "address challenges comprehensively"),
                "needed_for": info.get("needed_for", "national priorities in this area"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        return "\n\n".join(paragraphs) if paragraphs else user_content

    def _expand_conclusion(self, user_content: str, info: Dict[str, str]) -> str:
        """Expand Section V: Conclusion."""
        paragraphs = []

        # User content first
        if user_content:
            paragraphs.append(user_content)

        # Summary template
        templates = PS_SECTION_TEMPLATES.get("section_5_conclusion", {}).get("summary", [])
        if templates:
            import random
            template = random.choice(templates)
            defaults = {
                "current_org": info.get("current_org", "my current organization"),
                "field": self.field or "this field",
                "benefit_1": info.get("benefit_1", "technical capabilities in this critical area"),
                "benefit_2": info.get("benefit_2", "American leadership and competitiveness"),
                "commitment_area": info.get("commitment_area", "advancing national priorities"),
            }
            try:
                paragraphs.append(template.format(**defaults))
            except KeyError:
                pass

        return "\n\n".join(paragraphs) if paragraphs else user_content

    def _add_elaboration_to_ps(self, sections: Dict[str, str], target_words: int) -> Dict[str, str]:
        """Add elaboration to reach target word count."""
        current_words = sum(len(s.split()) for s in sections.values())
        if current_words >= target_words:
            return sections

        # Add more government citations to national importance
        if self.field in GOVERNMENT_DOCS:
            docs = GOVERNMENT_DOCS[self.field]
            if len(docs) > 2:
                additional = f"Further supporting the national priority of this work, {docs[2]['name']} emphasizes the importance of continued investment in this area."
                sections["national_importance"] += "\n\n" + additional

        return sections


# =============================================================================
# CONVENIENCE FUNCTION
# =============================================================================

def generate_ps(
    beneficiary_name: str,
    field: str,
    sections: Dict[str, str],
    output_dir: str = "output",
    **kwargs,
) -> str:
    """
    High-level function to generate a Personal Statement.

    Args:
        beneficiary_name: Name of petitioner
        field: Field for government document citations
        sections: Dictionary with section content
        output_dir: Directory for output files
        **kwargs: Additional arguments

    Returns:
        Path to generated document
    """
    generator = PSGenerator(
        beneficiary_name=beneficiary_name,
        field=field,
        output_dir=output_dir,
    )

    return generator.create_document(sections=sections, **kwargs)


# =============================================================================
# CLI INTERFACE
# =============================================================================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="ILC Personal Statement Generator")
    parser.add_argument("--name", required=True, help="Beneficiary name")
    parser.add_argument("--field", choices=GOVERNMENT_DOCS.keys(), required=True, help="Field")
    parser.add_argument("--output", default="output", help="Output directory")

    args = parser.parse_args()

    generator = PSGenerator(
        beneficiary_name=args.name,
        field=args.field,
        output_dir=args.output,
    )

    # Create with templates
    sections = {
        "overview": generator.get_section_template("overview"),
        "national_importance": generator.get_section_template("national_importance"),
        "practical_impact": generator.get_section_template("practical_impact"),
        "well_positioned": generator.get_section_template("well_positioned"),
        "conclusion": generator.get_section_template("conclusion"),
    }

    filepath = generator.create_document(sections=sections)

    print(f"\nGenerated Personal Statement: {filepath}")
    print("\nNote: Replace bracketed sections with actual content.")
