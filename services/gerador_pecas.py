"""
CaseHub Lite - Gerador de Pecas Juridicas (Legal Document Generator)

Generates legal documents (peticoes, contestacoes, recursos, etc.) using:
1. Template-based generation (fill in fields from form)
2. LLM enhancement via Ollama/Maestro (optional)
3. Export as TXT or DOCX

Supported document types:
- Peticao Inicial Civel
- Contestacao
- Recurso de Apelacao
- Agravo de Instrumento
- Embargos de Declaracao
- Recurso Ordinario Trabalhista
- Replica
- Mandado de Seguranca
"""
import logging
import os
import io
from datetime import date, datetime
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Ollama / Maestro integration
# ---------------------------------------------------------------------------
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://host.docker.internal:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.2:3b")


# ---------------------------------------------------------------------------
# Templates registry
# ---------------------------------------------------------------------------
PECAS_TEMPLATES = {
    "peticao_inicial_civel": {
        "nome": "Peticao Inicial Civel",
        "descricao": "Peticao para iniciar acao civel",
        "icone": "fas fa-gavel",
        "cor": "primary",
        "campos": [
            {"id": "autor", "label": "Autor (nome completo)", "tipo": "text", "required": True},
            {"id": "qualificacao_autor", "label": "Qualificacao do Autor", "tipo": "textarea",
             "placeholder": "Nacionalidade, estado civil, profissao, CPF, RG, endereco"},
            {"id": "reu", "label": "Reu (nome completo)", "tipo": "text", "required": True},
            {"id": "qualificacao_reu", "label": "Qualificacao do Reu", "tipo": "textarea",
             "placeholder": "Nacionalidade, estado civil, profissao, CNPJ/CPF, endereco"},
            {"id": "tipo_acao", "label": "Tipo de Acao", "tipo": "text",
             "placeholder": "Ex: Indenizatoria, Cobranca, Obrigacao de Fazer"},
            {"id": "comarca", "label": "Comarca", "tipo": "text", "required": True},
            {"id": "vara", "label": "Vara", "tipo": "text",
             "placeholder": "Ex: 1a Vara Civel"},
            {"id": "valor_causa", "label": "Valor da Causa (R$)", "tipo": "number"},
            {"id": "fatos", "label": "I - DOS FATOS", "tipo": "textarea", "required": True,
             "placeholder": "Descreva os fatos que fundamentam a acao..."},
            {"id": "fundamentos", "label": "II - DO DIREITO", "tipo": "textarea", "required": True,
             "placeholder": "Fundamentos juridicos (artigos de lei, jurisprudencia)..."},
            {"id": "pedidos", "label": "III - DOS PEDIDOS", "tipo": "textarea", "required": True,
             "placeholder": "Liste os pedidos (ex: condenacao ao pagamento de...)"},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text", "placeholder": "Ex: MG"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {vara} DA COMARCA DE {comarca}",
            "",
            "",
            "{autor}, {qualificacao_autor}, vem, respeitosamente, a presenca de Vossa Excelencia, por seu advogado que esta subscreve (procuracao anexa), com escritorio profissional no endereco constante do rodape, onde recebe intimacoes, propor a presente",
            "",
            "ACAO {tipo_acao}",
            "",
            "em face de {reu}, {qualificacao_reu}, pelos fatos e fundamentos a seguir expostos:",
            "",
            "",
            "I - DOS FATOS",
            "",
            "{fatos}",
            "",
            "",
            "II - DO DIREITO",
            "",
            "{fundamentos}",
            "",
            "",
            "III - DOS PEDIDOS",
            "",
            "Diante do exposto, requer:",
            "",
            "{pedidos}",
            "",
            "",
            "IV - DO VALOR DA CAUSA",
            "",
            "Da-se a causa o valor de R$ {valor_causa} ({valor_causa_extenso}).",
            "",
            "",
            "V - DAS PROVAS",
            "",
            "Protesta provar o alegado por todos os meios de prova em direito admitidos, especialmente prova documental, testemunhal e pericial, se necessario.",
            "",
            "",
            "Termos em que,",
            "Pede deferimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
    "contestacao": {
        "nome": "Contestacao",
        "descricao": "Resposta do reu a peticao inicial",
        "icone": "fas fa-shield-alt",
        "cor": "danger",
        "campos": [
            {"id": "processo", "label": "Numero do Processo", "tipo": "text", "required": True},
            {"id": "reu", "label": "Reu / Contestante", "tipo": "text", "required": True},
            {"id": "qualificacao_reu", "label": "Qualificacao do Reu", "tipo": "textarea"},
            {"id": "autor", "label": "Autor", "tipo": "text", "required": True},
            {"id": "vara", "label": "Vara", "tipo": "text"},
            {"id": "comarca", "label": "Comarca", "tipo": "text"},
            {"id": "preliminares", "label": "Preliminares", "tipo": "textarea",
             "placeholder": "Inepcia da inicial, ilegitimidade passiva, falta de interesse processual..."},
            {"id": "fatos_defesa", "label": "Dos Fatos / Merito", "tipo": "textarea", "required": True,
             "placeholder": "Contraponha os fatos narrados pelo autor e apresente a versao do reu..."},
            {"id": "merito", "label": "Do Direito (Fundamentacao Juridica)", "tipo": "textarea",
             "placeholder": "Fundamentos juridicos para a defesa..."},
            {"id": "pedidos", "label": "Dos Pedidos", "tipo": "textarea",
             "placeholder": "Ex: improcedencia total dos pedidos autorais..."},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {vara} DA COMARCA DE {comarca}",
            "",
            "Processo n. {processo}",
            "",
            "",
            "{reu}, {qualificacao_reu}, nos autos da acao movida por {autor}, processo em epigrafe, vem, respeitosamente, perante Vossa Excelencia, por seu advogado que esta subscreve, apresentar",
            "",
            "CONTESTACAO",
            "",
            "pelos fatos e fundamentos a seguir expostos:",
            "",
            "",
            "I - DAS PRELIMINARES",
            "",
            "{preliminares}",
            "",
            "",
            "II - DOS FATOS",
            "",
            "{fatos_defesa}",
            "",
            "",
            "III - DO DIREITO",
            "",
            "{merito}",
            "",
            "",
            "IV - DOS PEDIDOS",
            "",
            "Diante do exposto, requer:",
            "",
            "{pedidos}",
            "",
            "",
            "Protesta provar o alegado por todos os meios de prova em direito admitidos.",
            "",
            "Termos em que,",
            "Pede deferimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
    "recurso_apelacao": {
        "nome": "Recurso de Apelacao",
        "descricao": "Recurso contra sentenca de primeiro grau",
        "icone": "fas fa-level-up-alt",
        "cor": "warning",
        "campos": [
            {"id": "processo", "label": "Numero do Processo", "tipo": "text", "required": True},
            {"id": "apelante", "label": "Apelante", "tipo": "text", "required": True},
            {"id": "qualificacao_apelante", "label": "Qualificacao do Apelante", "tipo": "textarea"},
            {"id": "apelado", "label": "Apelado", "tipo": "text", "required": True},
            {"id": "vara", "label": "Vara de Origem", "tipo": "text"},
            {"id": "comarca", "label": "Comarca", "tipo": "text"},
            {"id": "sentenca_recorrida", "label": "Sentenca Recorrida (resumo)", "tipo": "textarea",
             "placeholder": "Resuma a sentenca que esta sendo impugnada..."},
            {"id": "razoes", "label": "Razoes do Recurso", "tipo": "textarea", "required": True,
             "placeholder": "Fundamentos de fato e de direito para a reforma da sentenca..."},
            {"id": "pedidos", "label": "Dos Pedidos", "tipo": "textarea",
             "placeholder": "Ex: reforma total da sentenca, provimento do recurso..."},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {vara} DA COMARCA DE {comarca}",
            "",
            "Processo n. {processo}",
            "",
            "",
            "{apelante}, {qualificacao_apelante}, inconformado com a r. sentenca proferida nos autos da acao em epigrafe, vem, respeitosamente, interpor o presente",
            "",
            "RECURSO DE APELACAO",
            "",
            "com fundamento no art. 1.009 e seguintes do CPC, requerendo seja recebido e processado, remetendo-se os autos ao Egr. Tribunal de Justica, conforme as razoes anexas.",
            "",
            "",
            "Termos em que,",
            "Pede deferimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
            "",
            "",
            "=" * 60,
            "",
            "RAZOES DE APELACAO",
            "",
            "Processo n. {processo}",
            "Apelante: {apelante}",
            "Apelado: {apelado}",
            "",
            "",
            "EGREGIO TRIBUNAL DE JUSTICA",
            "COLENDA CAMARA",
            "",
            "",
            "I - DA SENTENCA RECORRIDA",
            "",
            "{sentenca_recorrida}",
            "",
            "",
            "II - DAS RAZOES DO RECURSO",
            "",
            "{razoes}",
            "",
            "",
            "III - DOS PEDIDOS",
            "",
            "{pedidos}",
            "",
            "",
            "Termos em que,",
            "Pede provimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
    "agravo_instrumento": {
        "nome": "Agravo de Instrumento",
        "descricao": "Recurso contra decisao interlocutoria",
        "icone": "fas fa-bolt",
        "cor": "info",
        "campos": [
            {"id": "processo", "label": "Numero do Processo (origem)", "tipo": "text", "required": True},
            {"id": "agravante", "label": "Agravante", "tipo": "text", "required": True},
            {"id": "qualificacao_agravante", "label": "Qualificacao do Agravante", "tipo": "textarea"},
            {"id": "agravado", "label": "Agravado", "tipo": "text", "required": True},
            {"id": "vara", "label": "Vara / Juizo de Origem", "tipo": "text"},
            {"id": "comarca", "label": "Comarca", "tipo": "text"},
            {"id": "decisao_agravada", "label": "Decisao Agravada", "tipo": "textarea", "required": True,
             "placeholder": "Transcreva ou resuma a decisao interlocutoria impugnada..."},
            {"id": "razoes", "label": "Razoes do Agravo", "tipo": "textarea", "required": True,
             "placeholder": "Fundamentos para a reforma da decisao..."},
            {"id": "tutela_urgencia", "label": "Pedido de Efeito Suspensivo / Tutela de Urgencia", "tipo": "textarea",
             "placeholder": "Se aplicavel, fundamente o periculum in mora e fumus boni iuris..."},
            {"id": "pedidos", "label": "Dos Pedidos", "tipo": "textarea",
             "placeholder": "Ex: concessao de efeito suspensivo e reforma da decisao..."},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DESEMBARGADOR(A) PRESIDENTE DO TRIBUNAL DE JUSTICA DO ESTADO",
            "",
            "Processo de Origem n. {processo}",
            "Juizo de Origem: {vara} da Comarca de {comarca}",
            "",
            "",
            "{agravante}, {qualificacao_agravante}, inconformado(a) com a r. decisao interlocutoria proferida pelo MM. Juizo da {vara} da Comarca de {comarca}, vem, respeitosamente, interpor o presente",
            "",
            "AGRAVO DE INSTRUMENTO",
            "",
            "com fundamento no art. 1.015 do CPC, contra a decisao que segue transcrita, pelas razoes de fato e de direito a seguir aduzidas.",
            "",
            "",
            "I - DA DECISAO AGRAVADA",
            "",
            "{decisao_agravada}",
            "",
            "",
            "II - DO CABIMENTO",
            "",
            "O presente recurso e cabivel nos termos do art. 1.015 do CPC, tratando-se de decisao interlocutoria que causa gravame a parte agravante.",
            "",
            "",
            "III - DO PEDIDO DE EFEITO SUSPENSIVO / TUTELA DE URGENCIA",
            "",
            "{tutela_urgencia}",
            "",
            "",
            "IV - DAS RAZOES DO RECURSO",
            "",
            "{razoes}",
            "",
            "",
            "V - DOS PEDIDOS",
            "",
            "{pedidos}",
            "",
            "",
            "Termos em que,",
            "Pede provimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
    "embargos_declaracao": {
        "nome": "Embargos de Declaracao",
        "descricao": "Pedido de esclarecimento de decisao",
        "icone": "fas fa-search-plus",
        "cor": "secondary",
        "campos": [
            {"id": "processo", "label": "Numero do Processo", "tipo": "text", "required": True},
            {"id": "embargante", "label": "Embargante", "tipo": "text", "required": True},
            {"id": "qualificacao_embargante", "label": "Qualificacao do Embargante", "tipo": "textarea"},
            {"id": "vara", "label": "Vara", "tipo": "text"},
            {"id": "comarca", "label": "Comarca", "tipo": "text"},
            {"id": "tipo_vicio", "label": "Tipo de Vicio", "tipo": "select",
             "opcoes": ["Omissao", "Contradicao", "Obscuridade", "Erro Material"],
             "required": True},
            {"id": "decisao_embargada", "label": "Decisao Embargada (resumo)", "tipo": "textarea",
             "placeholder": "Resuma a decisao que contem o vicio..."},
            {"id": "fundamentacao", "label": "Fundamentacao", "tipo": "textarea", "required": True,
             "placeholder": "Aponte o vicio e fundamente o pedido de esclarecimento..."},
            {"id": "pedidos", "label": "Dos Pedidos", "tipo": "textarea",
             "placeholder": "Ex: acolhimento dos embargos para sanar a omissao..."},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {vara} DA COMARCA DE {comarca}",
            "",
            "Processo n. {processo}",
            "",
            "",
            "{embargante}, {qualificacao_embargante}, nos autos do processo em epigrafe, vem, respeitosamente, perante Vossa Excelencia, opor os presentes",
            "",
            "EMBARGOS DE DECLARACAO",
            "",
            "com fundamento no art. 1.022 do CPC, em razao de {tipo_vicio} na r. decisao de fls., conforme razoes a seguir expostas.",
            "",
            "",
            "I - DA DECISAO EMBARGADA",
            "",
            "{decisao_embargada}",
            "",
            "",
            "II - DO VICIO ({tipo_vicio})",
            "",
            "{fundamentacao}",
            "",
            "",
            "III - DOS PEDIDOS",
            "",
            "{pedidos}",
            "",
            "",
            "Termos em que,",
            "Pede deferimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
    "recurso_ordinario_trabalhista": {
        "nome": "Recurso Ordinario Trabalhista",
        "descricao": "Recurso trabalhista para TRT",
        "icone": "fas fa-hard-hat",
        "cor": "success",
        "campos": [
            {"id": "processo", "label": "Numero do Processo", "tipo": "text", "required": True},
            {"id": "recorrente", "label": "Recorrente", "tipo": "text", "required": True},
            {"id": "qualificacao_recorrente", "label": "Qualificacao do Recorrente", "tipo": "textarea"},
            {"id": "recorrido", "label": "Recorrido", "tipo": "text", "required": True},
            {"id": "vara", "label": "Vara do Trabalho de Origem", "tipo": "text"},
            {"id": "comarca", "label": "Municipio", "tipo": "text"},
            {"id": "sentenca_recorrida", "label": "Sentenca Recorrida (resumo)", "tipo": "textarea"},
            {"id": "razoes", "label": "Razoes do Recurso", "tipo": "textarea", "required": True,
             "placeholder": "Fundamentos para a reforma da sentenca trabalhista..."},
            {"id": "pedidos", "label": "Dos Pedidos", "tipo": "textarea",
             "placeholder": "Ex: reforma total da sentenca, provimento do recurso..."},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DO TRABALHO DA {vara} DE {comarca}",
            "",
            "Processo n. {processo}",
            "",
            "",
            "{recorrente}, {qualificacao_recorrente}, nos autos da Reclamacao Trabalhista movida por/contra {recorrido}, inconformado com a r. sentenca proferida, vem, respeitosamente, interpor",
            "",
            "RECURSO ORDINARIO",
            "",
            "com fundamento no art. 895, I, da CLT, requerendo seja recebido e processado, remetendo-se os autos ao Egr. Tribunal Regional do Trabalho, conforme as razoes anexas.",
            "",
            "",
            "Termos em que,",
            "Pede deferimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
            "",
            "",
            "=" * 60,
            "",
            "RAZOES DE RECURSO ORDINARIO",
            "",
            "Processo n. {processo}",
            "Recorrente: {recorrente}",
            "Recorrido: {recorrido}",
            "",
            "",
            "EGREGIO TRIBUNAL REGIONAL DO TRABALHO",
            "COLENDA TURMA",
            "",
            "",
            "I - DA SENTENCA RECORRIDA",
            "",
            "{sentenca_recorrida}",
            "",
            "",
            "II - DAS RAZOES DO RECURSO",
            "",
            "{razoes}",
            "",
            "",
            "III - DOS PEDIDOS",
            "",
            "{pedidos}",
            "",
            "",
            "Termos em que,",
            "Pede provimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
    "replica": {
        "nome": "Replica",
        "descricao": "Manifestacao do autor sobre a contestacao",
        "icone": "fas fa-reply",
        "cor": "dark",
        "campos": [
            {"id": "processo", "label": "Numero do Processo", "tipo": "text", "required": True},
            {"id": "autor", "label": "Autor / Replicante", "tipo": "text", "required": True},
            {"id": "qualificacao_autor", "label": "Qualificacao do Autor", "tipo": "textarea"},
            {"id": "reu", "label": "Reu / Contestante", "tipo": "text", "required": True},
            {"id": "vara", "label": "Vara", "tipo": "text"},
            {"id": "comarca", "label": "Comarca", "tipo": "text"},
            {"id": "pontos_contestados", "label": "Refutacao dos Pontos da Contestacao", "tipo": "textarea",
             "required": True,
             "placeholder": "Rebata cada argumento da contestacao, ponto a ponto..."},
            {"id": "pedidos", "label": "Dos Pedidos", "tipo": "textarea",
             "placeholder": "Ex: ratifica todos os termos da peticao inicial..."},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA {vara} DA COMARCA DE {comarca}",
            "",
            "Processo n. {processo}",
            "",
            "",
            "{autor}, {qualificacao_autor}, nos autos da acao movida em face de {reu}, processo em epigrafe, vem, respeitosamente, apresentar",
            "",
            "REPLICA A CONTESTACAO",
            "",
            "nos seguintes termos:",
            "",
            "",
            "I - DA REFUTACAO DOS ARGUMENTOS DA CONTESTACAO",
            "",
            "{pontos_contestados}",
            "",
            "",
            "II - DOS PEDIDOS",
            "",
            "{pedidos}",
            "",
            "",
            "Protesta provar o alegado por todos os meios de prova em direito admitidos.",
            "",
            "Termos em que,",
            "Pede deferimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
    "mandado_seguranca": {
        "nome": "Mandado de Seguranca",
        "descricao": "Protecao contra ato ilegal de autoridade",
        "icone": "fas fa-balance-scale-right",
        "cor": "purple",
        "campos": [
            {"id": "impetrante", "label": "Impetrante", "tipo": "text", "required": True},
            {"id": "qualificacao_impetrante", "label": "Qualificacao do Impetrante", "tipo": "textarea"},
            {"id": "autoridade_coatora", "label": "Autoridade Coatora", "tipo": "text", "required": True,
             "placeholder": "Ex: Delegado da Receita Federal de Juiz de Fora/MG"},
            {"id": "orgao", "label": "Orgao a que pertence a Autoridade", "tipo": "text",
             "placeholder": "Ex: Uniao Federal, Estado de Minas Gerais, Municipio de..."},
            {"id": "vara", "label": "Vara", "tipo": "text"},
            {"id": "comarca", "label": "Comarca / Secao Judiciaria", "tipo": "text"},
            {"id": "ato_coator", "label": "Ato Coator", "tipo": "textarea", "required": True,
             "placeholder": "Descreva o ato ilegal ou abusivo da autoridade..."},
            {"id": "direito_liquido", "label": "Do Direito Liquido e Certo", "tipo": "textarea", "required": True,
             "placeholder": "Demonstre o direito liquido e certo violado..."},
            {"id": "liminar", "label": "Pedido Liminar", "tipo": "textarea",
             "placeholder": "Fundamente o pedido de medida liminar (periculum in mora + fumus boni iuris)..."},
            {"id": "pedidos", "label": "Dos Pedidos", "tipo": "textarea",
             "placeholder": "Ex: concessao da seguranca para cessar o ato coator..."},
            {"id": "valor_causa", "label": "Valor da Causa (R$)", "tipo": "number"},
            {"id": "advogado", "label": "Advogado(a)", "tipo": "text"},
            {"id": "oab_uf", "label": "UF da OAB", "tipo": "text"},
            {"id": "oab_numero", "label": "Numero OAB", "tipo": "text"},
            {"id": "cidade", "label": "Cidade", "tipo": "text"},
        ],
        "estrutura": [
            "EXCELENTISSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) FEDERAL DA {vara} DA SECAO JUDICIARIA DE {comarca}",
            "",
            "",
            "{impetrante}, {qualificacao_impetrante}, vem, respeitosamente, perante Vossa Excelencia, com fundamento no art. 5., LXIX, da Constituicao Federal, e na Lei n. 12.016/2009, impetrar",
            "",
            "MANDADO DE SEGURANCA COM PEDIDO LIMINAR",
            "",
            "contra ato ilegal e abusivo praticado por {autoridade_coatora} ({orgao}), pelos fatos e fundamentos a seguir expostos:",
            "",
            "",
            "I - DO ATO COATOR",
            "",
            "{ato_coator}",
            "",
            "",
            "II - DO DIREITO LIQUIDO E CERTO",
            "",
            "{direito_liquido}",
            "",
            "",
            "III - DA MEDIDA LIMINAR",
            "",
            "Nos termos do art. 7., III, da Lei n. 12.016/2009, estando presentes o fumus boni iuris e o periculum in mora:",
            "",
            "{liminar}",
            "",
            "",
            "IV - DOS PEDIDOS",
            "",
            "Diante do exposto, requer:",
            "",
            "a) a concessao de medida liminar, inaudita altera parte, para suspender os efeitos do ato coator;",
            "",
            "b) a notificacao da autoridade coatora para que preste as informacoes no prazo legal;",
            "",
            "c) a oitiva do Ministerio Publico;",
            "",
            "d) a concessao definitiva da seguranca, confirmando-se a liminar;",
            "",
            "{pedidos}",
            "",
            "",
            "Da-se a causa o valor de R$ {valor_causa}.",
            "",
            "Termos em que,",
            "Pede deferimento.",
            "",
            "{cidade}, {data}.",
            "",
            "",
            "___________________________________",
            "{advogado}",
            "OAB/{oab_uf} n. {oab_numero}",
        ],
    },
}


# ---------------------------------------------------------------------------
# Generator functions
# ---------------------------------------------------------------------------
def gerar_peca_template(tipo: str, dados: dict) -> str:
    """
    Generate a legal document from template + field data.
    Returns the complete document as plain text.
    """
    template = PECAS_TEMPLATES.get(tipo)
    if not template:
        raise ValueError(f"Tipo de peca desconhecido: {tipo}")

    # Set defaults for common fields
    dados.setdefault("data", date.today().strftime("%d de %B de %Y").replace(
        "January", "janeiro").replace("February", "fevereiro").replace(
        "March", "marco").replace("April", "abril").replace(
        "May", "maio").replace("June", "junho").replace(
        "July", "julho").replace("August", "agosto").replace(
        "September", "setembro").replace("October", "outubro").replace(
        "November", "novembro").replace("December", "dezembro"))

    # Format date in Portuguese
    hoje = date.today()
    meses = {
        1: "janeiro", 2: "fevereiro", 3: "marco", 4: "abril",
        5: "maio", 6: "junho", 7: "julho", 8: "agosto",
        9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro",
    }
    dados["data"] = f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"

    # valor_causa_extenso (basic)
    if "valor_causa" in dados and dados["valor_causa"]:
        try:
            val = float(dados["valor_causa"])
            dados["valor_causa"] = f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        except (ValueError, TypeError):
            pass
    dados.setdefault("valor_causa_extenso", "valor por extenso")

    # Fill missing fields with placeholder
    for campo in template.get("campos", []):
        campo_id = campo["id"] if isinstance(campo, dict) else campo
        dados.setdefault(campo_id, "_______________")

    # Also set common fields that may not be in campos list
    for key in ["qualificacao_autor", "qualificacao_reu", "qualificacao_apelante",
                "qualificacao_agravante", "qualificacao_embargante", "qualificacao_recorrente",
                "qualificacao_impetrante", "orgao", "sentenca_recorrida", "decisao_embargada",
                "tutela_urgencia", "liminar", "tipo_vicio"]:
        dados.setdefault(key, "_______________")

    # Build the document from structure
    lines = []
    for line in template.get("estrutura", []):
        if isinstance(line, str):
            try:
                lines.append(line.format(**dados))
            except (KeyError, IndexError) as e:
                logger.warning("Template field missing: %s", e)
                lines.append(line)
        else:
            lines.append(str(line))

    return "\n".join(lines)


async def gerar_peca_llm(tipo: str, dados: dict, texto_base: str) -> dict:
    """
    Enhance a template-generated document using Ollama LLM.
    Returns dict with 'texto' and 'status'.
    """
    import httpx

    template = PECAS_TEMPLATES.get(tipo)
    nome_peca = template["nome"] if template else tipo

    prompt = f"""Voce e um advogado brasileiro experiente. Recebeu o rascunho abaixo de uma {nome_peca}.

Sua tarefa:
1. Revise e melhore a redacao juridica, mantendo a estrutura
2. Adicione citacoes de artigos de lei relevantes onde apropriado
3. Melhore a argumentacao juridica onde houver espaco
4. Mantenha todos os dados faticos fornecidos (nomes, numeros, datas)
5. NAO invente fatos - apenas melhore a redacao e fundamentacao
6. Mantenha o formato de peticao (cabecalho, secoes, fecho)
7. Responda APENAS com o texto da peca melhorada, sem comentarios adicionais

RASCUNHO:

{texto_base}"""

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(
                f"{OLLAMA_URL}/api/chat",
                json={
                    "model": OLLAMA_MODEL,
                    "messages": [
                        {"role": "system", "content": "Voce e um advogado brasileiro com 20 anos de experiencia em direito processual civil. Redija pecas juridicas com linguagem tecnica precisa e fundamentacao solida."},
                        {"role": "user", "content": prompt},
                    ],
                    "stream": False,
                },
            )
            if resp.status_code == 200:
                data = resp.json()
                return {
                    "texto": data["message"]["content"],
                    "status": "ok",
                    "model": OLLAMA_MODEL,
                }
            else:
                logger.error("Ollama returned status %s", resp.status_code)
                return {"texto": texto_base, "status": "error", "error": f"Ollama retornou status {resp.status_code}"}
    except Exception as e:
        logger.error("Ollama error: %s", e)
        return {
            "texto": texto_base,
            "status": "offline",
            "error": "Ollama nao esta disponivel. Usando texto do template.",
        }


async def verificar_ollama() -> bool:
    """Check if Ollama is running and accessible."""
    import httpx
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{OLLAMA_URL}/api/tags")
            return resp.status_code == 200
    except Exception:
        return False


def gerar_docx(texto: str, titulo: str) -> Optional[io.BytesIO]:
    """
    Generate a DOCX file from plain text.
    Returns BytesIO buffer or None if python-docx is not available.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx not installed. DOCX export unavailable.")
        return None

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Process text line by line
    lines = texto.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Section headers (Roman numerals)
        if line.strip().startswith(("I -", "II -", "III -", "IV -", "V -", "VI -",
                                     "EXCELENTISSIMO", "ACAO ", "CONTESTACAO",
                                     "RECURSO DE APELACAO", "AGRAVO DE INSTRUMENTO",
                                     "EMBARGOS DE DECLARACAO", "RECURSO ORDINARIO",
                                     "REPLICA A CONTESTACAO", "MANDADO DE SEGURANCA",
                                     "RAZOES DE APELACAO", "RAZOES DE RECURSO",
                                     "EGREGIO TRIBUNAL", "COLENDA")):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            if line.strip().startswith("EXCELENTISSIMO"):
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run.font.size = Pt(12)
            elif line.strip() in ("CONTESTACAO", "RECURSO DE APELACAO", "AGRAVO DE INSTRUMENTO",
                                  "EMBARGOS DE DECLARACAO", "RECURSO ORDINARIO",
                                  "REPLICA A CONTESTACAO", "MANDADO DE SEGURANCA COM PEDIDO LIMINAR",
                                  "RAZOES DE APELACAO", "RAZOES DE RECURSO ORDINARIO"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = Pt(14)
            elif "ACAO " in line.strip() and len(line.strip()) < 50:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = Pt(14)
        elif line.strip().startswith("=" * 10):
            doc.add_paragraph()  # page break equivalent - separator
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("_" * 40)
            doc.add_paragraph()
        elif line.strip().startswith("___"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line.strip())
        elif line.strip() == "":
            doc.add_paragraph()
        elif line.strip().startswith(("Termos em que,", "Pede deferimento.", "Pede provimento.")):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(line.strip())
        elif line.strip().startswith(("OAB/", "Processo n.")):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if line.strip().startswith("OAB/") else WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(line.strip())
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(line)
            p.paragraph_format.first_line_indent = Cm(1.5)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def listar_pecas() -> list:
    """Return list of all available pecas with metadata."""
    pecas = []
    for tipo, template in PECAS_TEMPLATES.items():
        pecas.append({
            "tipo": tipo,
            "nome": template["nome"],
            "descricao": template["descricao"],
            "icone": template.get("icone", "fas fa-file-alt"),
            "cor": template.get("cor", "primary"),
            "num_campos": len(template.get("campos", [])),
        })
    return pecas


def get_template(tipo: str) -> Optional[dict]:
    """Get a specific template by type."""
    return PECAS_TEMPLATES.get(tipo)
