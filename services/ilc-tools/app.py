#!/usr/bin/env python3
"""
CaseHub - Web Interface
Dark theme, minimalist design
Version: 2.1.0 - Security Update with Auth, CORS, Rate Limiting
"""

import os
import uuid
import shutil
import logging
import secrets
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional, List

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException, Depends, Cookie
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware

# PDF/Document handling
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.units import inch
from docx import Document as DocxDocument
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
import io

# Import tools
from tools.lor_generator import LORGenerator, PERSONAS, NATIONAL_INTEREST_TEXTS
from tools.ps_generator import PSGenerator, GOVERNMENT_DOCS
from tools.package_builder import PackageBuilder, EXHIBITS, quick_merge, convert_images_to_pdf

# Import client mapping manager for email processor clients
import client_mapping_manager

# Import leads CRM manager
import leads_manager

# Create app
app = FastAPI(title="CaseHub", version="2.1.0-alpha")

# =============================================================================
# SECURITY MIDDLEWARE
# =============================================================================

# CORS - Restrict to specific origins
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "https://casehub.app").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["Authorization", "Content-Type"],
)

# Security Headers Middleware
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "geolocation=(), microphone=(), camera=()"
        # HSTS only in production
        if os.getenv("ENVIRONMENT") == "production":
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response

app.add_middleware(SecurityHeadersMiddleware)

# Global exception handler for API routes (prevents HTML errors in JSON responses)
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    """Catch unhandled exceptions and return JSON instead of HTML for API routes."""
    logger.error(f"Unhandled exception on {request.url.path}: {str(exc)}", exc_info=True)

    # Return JSON for API endpoints (prevents "Unexpected token '<'" errors)
    if request.url.path.startswith('/api/'):
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": f"Internal server error: {str(exc)}"
            }
        )

    # For non-API routes, re-raise to get default HTML error page
    raise exc

# Simple in-memory rate limiter (for production, use Redis-based solution)
from collections import defaultdict
import time

class RateLimiter:
    def __init__(self):
        self.requests = defaultdict(list)
        self.limits = {
            "default": (60, 60),      # 60 requests per 60 seconds
            "auth": (5, 900),         # 5 login attempts per 15 minutes
            "generate": (20, 3600),   # 20 document generations per hour
        }

    def is_allowed(self, ip: str, limit_type: str = "default") -> bool:
        max_requests, window = self.limits.get(limit_type, self.limits["default"])
        now = time.time()

        # Clean old requests
        self.requests[ip] = [t for t in self.requests[ip] if now - t < window]

        if len(self.requests[ip]) >= max_requests:
            return False

        self.requests[ip].append(now)
        return True

rate_limiter = RateLimiter()

async def check_rate_limit(request: Request, limit_type: str = "default"):
    """Check rate limit for the request."""
    client_ip = request.client.host if request.client else "unknown"
    if not rate_limiter.is_allowed(client_ip, limit_type):
        raise HTTPException(status_code=429, detail="Too many requests. Please try again later.")

# Directories
BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "output"
UPLOAD_DIR = BASE_DIR / "uploads"
TEMPLATES_DIR = BASE_DIR / "templates"
LOG_DIR = BASE_DIR / "logs"

OUTPUT_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)
TEMPLATES_DIR.mkdir(exist_ok=True)
LOG_DIR.mkdir(exist_ok=True)

# =============================================================================
# LOGGING SETUP
# =============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_DIR / f"ilc_{datetime.now().strftime('%Y%m%d')}.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
logger.info("CaseHub v2.1.0 (alpha) starting...")

# Import authentication module
import auth

# Templates
templates = Jinja2Templates(directory=str(TEMPLATES_DIR))

# Static files (for logo, etc.)
STATIC_DIR = BASE_DIR / "static"
STATIC_DIR.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


# =============================================================================
# AUTHENTICATION DEPENDENCY
# =============================================================================

from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr

security = HTTPBearer(auto_error=False)

class UserRegister(BaseModel):
    email: EmailStr
    username: str
    full_name: str

class UserLogin(BaseModel):
    email_or_username: str
    password: str

class TwoFactorVerify(BaseModel):
    email: str
    code: str
    method: str = "totp"  # "totp" or "email"

class PasswordChange(BaseModel):
    current_password: str
    new_password: str

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> Optional[dict]:
    """Get current authenticated user from JWT token."""
    if not credentials:
        return None

    payload = auth.verify_token(credentials.credentials)
    if not payload:
        return None

    email = payload.get("sub")
    if not email:
        return None

    user = auth.get_user_by_email(email)
    if not user or user.get("status") != "active":
        return None

    return user

async def require_auth(user: dict = Depends(get_current_user)):
    """Require authentication for endpoint."""
    if not user:
        raise HTTPException(status_code=401, detail="Authentication required")
    return user

async def require_admin(user: dict = Depends(require_auth)):
    """Require admin privileges."""
    if not user.get("is_admin") and not user.get("is_super_admin"):
        raise HTTPException(status_code=403, detail="Admin privileges required")
    return user

async def require_super_admin(user: dict = Depends(require_auth)):
    """Require super admin privileges."""
    if not user.get("is_super_admin"):
        raise HTTPException(status_code=403, detail="Super admin privileges required")
    return user


def verify_page_auth(request: Request) -> Optional[dict]:
    """
    Verify authentication for page requests using cookies.
    Returns user dict if valid, None otherwise.
    """
    # Check for token in cookie
    token = request.cookies.get("ilc_access_token")
    if not token:
        return None

    # Verify token
    payload = auth.verify_token(token)
    if not payload:
        return None

    email = payload.get("sub")
    if not email:
        return None

    user = auth.get_user_by_email(email)
    if not user or user.get("status") != "active":
        return None

    return user


# =============================================================================
# AUTHENTICATION ENDPOINTS
# =============================================================================

@app.post("/api/auth/register")
async def register_user(data: UserRegister, request: Request):
    """Register a new user (requires admin approval)."""
    await check_rate_limit(request, "auth")

    try:
        result = auth.create_user(data.email, data.username, data.full_name)

        # Notify admin about new registration
        logger.info(f"New registration from {request.client.host}: {data.email}")

        return {
            "success": True,
            "message": "Registration submitted. Waiting for admin approval.",
            "email": data.email
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/auth/login")
async def login_user(data: UserLogin, request: Request):
    """First step of login - verify credentials."""
    await check_rate_limit(request, "auth")

    ip = request.client.host if request.client else "unknown"

    # Check rate limiting
    if not auth.check_login_allowed(data.email_or_username, ip):
        raise HTTPException(status_code=429, detail="Too many login attempts. Try again in 15 minutes.")

    # Find user by email or username
    user = auth.get_user_by_email(data.email_or_username)
    if not user:
        user = auth.get_user_by_username(data.email_or_username)

    if not user:
        auth.record_login_attempt(data.email_or_username, ip, False)
        raise HTTPException(status_code=401, detail="Invalid credentials")

    # Check user status
    if user["status"] == "pending_approval":
        raise HTTPException(status_code=403, detail="Account pending approval")
    if user["status"] == "suspended":
        raise HTTPException(status_code=403, detail="Account suspended")

    # Verify password
    if not auth.verify_password(data.password, user["password_hash"]):
        auth.record_login_attempt(user["email"], ip, False)
        raise HTTPException(status_code=401, detail="Invalid credentials")

    # If 2FA is completely disabled (no secret and no preferred method), skip 2FA
    if not user.get("totp_secret") and not user.get("preferred_2fa"):
        # No 2FA configured - login directly
        ip = request.client.host if request.client else "unknown"
        user_agent = request.headers.get("user-agent", "unknown")
        auth.record_login_attempt(user["email"], ip, True)
        session = auth.create_session(user["email"], ip, user_agent)
        return {
            "success": True,
            "must_change_password": user.get("must_change_password", False),
            **session
        }

    # If 2FA not set up but required, require setup
    if not user.get("totp_secret") and user.get("preferred_2fa") == "totp":
        return {
            "success": True,
            "requires_2fa_setup": True,
            "email": user["email"],
            "message": "Please set up two-factor authentication"
        }

    # Return which 2FA method to use
    return {
        "success": True,
        "requires_2fa": True,
        "email": user["email"],
        "preferred_2fa": user.get("preferred_2fa", "totp"),
        "message": "Enter your 2FA code"
    }


@app.post("/api/auth/send-email-otp")
async def send_email_otp(request: Request, email: str = Form(...)):
    """Send OTP via email for 2FA."""
    await check_rate_limit(request, "auth")

    user = auth.get_user_by_email(email)
    if not user:
        # Don't reveal if user exists
        return {"success": True, "message": "If the email exists, a code was sent"}

    otp = auth.generate_email_otp(email)
    auth.send_otp_email(email, otp)

    return {"success": True, "message": "OTP sent to your email"}


@app.post("/api/auth/verify-2fa")
async def verify_2fa(data: TwoFactorVerify, request: Request):
    """Verify 2FA code and complete login."""
    await check_rate_limit(request, "auth")

    ip = request.client.host if request.client else "unknown"
    user_agent = request.headers.get("user-agent", "unknown")

    user = auth.get_user_by_email(data.email)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid request")

    # Verify code based on method
    if data.method == "totp":
        valid = auth.verify_totp(data.email, data.code)
    else:
        valid = auth.verify_email_otp(data.email, data.code)

    if not valid:
        auth.record_login_attempt(data.email, ip, False)
        raise HTTPException(status_code=401, detail="Invalid 2FA code")

    # Success! Create session
    auth.record_login_attempt(data.email, ip, True)
    session = auth.create_session(data.email, ip, user_agent)

    return {
        "success": True,
        "must_change_password": user.get("must_change_password", False),
        **session
    }


@app.post("/api/auth/setup-2fa")
async def setup_2fa(request: Request, user: dict = Depends(require_auth)):
    """Set up TOTP 2FA (returns QR code)."""
    result = auth.setup_totp(user["email"])

    return {
        "success": True,
        "qr_code": result["qr_code"],
        "secret": result["secret"],
        "message": "Scan QR code with Google Authenticator"
    }


@app.post("/api/auth/change-password")
async def change_password(data: PasswordChange, user: dict = Depends(require_auth)):
    """Change user password."""
    # Verify current password
    if not auth.verify_password(data.current_password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Current password is incorrect")

    # Validate new password strength
    if len(data.new_password) < 12:
        raise HTTPException(status_code=400, detail="Password must be at least 12 characters")

    auth.change_password(user["email"], data.new_password)

    return {"success": True, "message": "Password changed successfully"}


@app.post("/api/auth/refresh")
async def refresh_token(request: Request, credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Refresh access token using refresh token."""
    if not credentials:
        raise HTTPException(status_code=401, detail="Refresh token required")

    payload = auth.verify_token(credentials.credentials, token_type="refresh")
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid or expired refresh token")

    email = payload.get("sub")
    user = auth.get_user_by_email(email)
    if not user or user["status"] != "active":
        raise HTTPException(status_code=401, detail="User not found or inactive")

    # Create new access token
    token_data = {"sub": email, "session_id": payload.get("session_id")}
    access_token = auth.create_access_token(token_data)

    return {
        "access_token": access_token,
        "token_type": "bearer",
        "expires_in": 3600
    }


@app.post("/api/auth/logout")
async def logout(request: Request, credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Logout and invalidate session."""
    if credentials:
        payload = auth.verify_token(credentials.credentials)
        if payload and payload.get("session_id"):
            auth.invalidate_session(payload["session_id"])

    return {"success": True, "message": "Logged out successfully"}


@app.get("/api/auth/me")
async def get_current_user_info(user: dict = Depends(require_auth)):
    """Get current user information."""
    # Remove sensitive fields
    safe_user = {k: v for k, v in user.items() if k not in ["password_hash", "totp_secret"]}
    return safe_user


# =============================================================================
# ADMIN ENDPOINTS
# =============================================================================

@app.get("/api/admin/pending-users")
async def get_pending_users(admin: dict = Depends(require_admin)):
    """Get all users pending approval."""
    return {"users": auth.get_pending_users()}


@app.post("/api/admin/approve/{email}")
async def approve_user(email: str, admin: dict = Depends(require_admin)):
    """Approve a pending user."""
    try:
        result = auth.approve_user(email, admin["email"])
        # Send approval email with temp password
        auth.send_otp_email(email, f"Your account was approved! Temporary password: {result['temp_password']}")
        return {"success": True, "message": f"User {email} approved"}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/admin/reject/{email}")
async def reject_user(email: str, admin: dict = Depends(require_admin)):
    """Reject a pending user."""
    users = auth.load_users()
    if email.lower() in users:
        del users[email.lower()]
        auth.save_users(users)
        return {"success": True, "message": f"User {email} rejected"}
    raise HTTPException(status_code=404, detail="User not found")


@app.get("/api/admin/users")
async def list_all_users(admin: dict = Depends(require_admin)):
    """List all users."""
    return {"users": auth.get_all_users()}


@app.post("/api/admin/suspend/{email}")
async def suspend_user(email: str, admin: dict = Depends(require_admin)):
    """Suspend a user account."""
    try:
        auth.suspend_user(email, admin["email"])
        return {"success": True, "message": f"User {email} suspended"}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/admin/reset-2fa/{email}")
async def reset_user_2fa(email: str, admin: dict = Depends(require_admin)):
    """Reset user's 2FA."""
    try:
        auth.reset_user_2fa(email, admin["email"])
        return {"success": True, "message": f"2FA reset for {email}"}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/admin/activate/{email}")
async def activate_user(email: str, admin: dict = Depends(require_admin)):
    """Reactivate a suspended user."""
    try:
        users = auth.load_users()
        email = email.lower()
        if email not in users:
            raise HTTPException(status_code=404, detail="User not found")

        users[email]["status"] = "active"
        auth.save_users(users)
        logger.info(f"User reactivated: {email} by {admin['email']}")
        return {"success": True, "message": f"User {email} reactivated"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# =============================================================================
# COMMUNICATIONS ENDPOINTS
# =============================================================================

import communications
import attachment_metadata

# Pydantic models for communications
class ClientExclude(BaseModel):
    reason: str = "Manual exclusion"

class SendEmailRequest(BaseModel):
    test_mode: bool = False
    test_recipient: Optional[str] = None

class SendSingleEmailRequest(BaseModel):
    client_id: str
    template_type: str  # "weekly" or "monthly"
    test_mode: bool = False
    test_recipient: Optional[str] = None

class ScheduleUpdate(BaseModel):
    is_active: Optional[bool] = None
    hour: Optional[int] = None
    minute: Optional[int] = None

class TemplateUpdate(BaseModel):
    subject: Optional[str] = None
    body_html: Optional[str] = None
    body_text: Optional[str] = None
    is_active: Optional[bool] = None


# --- Client Management ---

@app.get("/api/communications/clients")
async def get_comm_clients(admin: dict = Depends(require_admin)):
    """Get all communication clients."""
    clients = communications.get_all_comm_clients()
    return {"clients": clients, "total": len(clients)}


@app.get("/api/communications/clients/{client_id}")
async def get_comm_client(client_id: str, admin: dict = Depends(require_admin)):
    """Get a single communication client."""
    client = communications.get_comm_client_by_id(client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    return client


@app.post("/api/communications/clients/sync")
async def sync_comm_clients(admin: dict = Depends(require_admin)):
    """Sync clients from active-clients.json source file."""
    try:
        # Path to active-clients.json
        source_file = str(BASE_DIR.parent / "immigrant-law-ops" / "whatsapp-bot" / "client-followup" / "active-clients.json")
        result = communications.sync_clients_from_json(source_file)
        logger.info(f"Clients synced by {admin['email']}: {result}")
        return {"success": True, **result}
    except Exception as e:
        logger.error(f"Client sync failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/api/communications/clients/{client_id}")
async def update_comm_client(client_id: str, updates: dict, admin: dict = Depends(require_admin)):
    """Update a communication client."""
    try:
        client = communications.update_comm_client(client_id, updates)
        logger.info(f"Client {client_id} updated by {admin['email']}")
        return client
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.post("/api/communications/clients/{client_id}/exclude")
async def exclude_comm_client(client_id: str, data: ClientExclude, admin: dict = Depends(require_admin)):
    """Add client to exclusion list."""
    try:
        client = communications.exclude_client(client_id, data.reason, admin["email"])
        logger.info(f"Client {client_id} excluded by {admin['email']}: {data.reason}")
        return {"success": True, "client": client}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.delete("/api/communications/clients/{client_id}/exclude")
async def include_comm_client(client_id: str, admin: dict = Depends(require_admin)):
    """Remove client from exclusion list."""
    try:
        client = communications.include_client(client_id)
        logger.info(f"Client {client_id} removed from exclusion by {admin['email']}")
        return {"success": True, "client": client}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


# --- Email Processor Client Management (CLIENT_MAPPING) ---

@app.get("/api/email-processor/clients")
async def get_email_processor_clients(admin: dict = Depends(require_admin)):
    """Get all clients from email_processor.py CLIENT_MAPPING."""
    try:
        clients = client_mapping_manager.get_all_clients()
        return {"clients": clients, "total": len(clients)}
    except Exception as e:
        logger.error(f"Error getting email processor clients: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/email-processor/clients/{email:path}")
async def update_email_processor_client(
    email: str,
    updates: dict,
    admin: dict = Depends(require_admin)
):
    """Update a client in email_processor.py CLIENT_MAPPING."""
    try:
        # Validate paralegal
        valid_paralegals = ["Ana Clara", "Juliana"]
        if "paralegal" in updates and updates["paralegal"] not in valid_paralegals:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid paralegal. Must be one of: {valid_paralegals}"
            )

        client = client_mapping_manager.update_client(email, updates)
        logger.info(f"Email processor client {email} updated by {admin['email']}: {updates}")
        return {"success": True, "client": client}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error updating email processor client: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/email-processor/clients")
async def add_email_processor_client(
    data: dict,
    admin: dict = Depends(require_admin)
):
    """Add a new client to email_processor.py CLIENT_MAPPING."""
    try:
        required = ["email", "name", "paralegal"]
        for field in required:
            if not data.get(field):
                raise HTTPException(status_code=400, detail=f"Missing required field: {field}")

        valid_paralegals = ["Ana Clara", "Juliana"]
        if data["paralegal"] not in valid_paralegals:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid paralegal. Must be one of: {valid_paralegals}"
            )

        client = client_mapping_manager.add_client(
            email=data["email"].lower().strip(),
            name=data["name"].strip(),
            paralegal=data["paralegal"],
            case=data.get("case", "")
        )
        logger.info(f"Email processor client added by {admin['email']}: {data['email']}")
        return {"success": True, "client": client}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error adding email processor client: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/email-processor/clients/{email:path}")
async def delete_email_processor_client(
    email: str,
    admin: dict = Depends(require_admin)
):
    """Remove a client from email_processor.py CLIENT_MAPPING."""
    try:
        client_mapping_manager.remove_client(email)
        logger.info(f"Email processor client {email} removed by {admin['email']}")
        return {"success": True}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Error removing email processor client: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Communications History ---

@app.get("/api/communications/history")
async def get_comm_history(
    client_id: Optional[str] = None,
    comm_type: Optional[str] = None,
    status: Optional[str] = None,
    limit: int = 100,
    offset: int = 0,
    admin: dict = Depends(require_admin)
):
    """Get communications history with filters."""
    history = communications.get_communications_history(
        client_id=client_id,
        comm_type=comm_type,
        status=status,
        limit=limit,
        offset=offset
    )
    return {"communications": history, "count": len(history)}


@app.get("/api/communications/status")
async def get_comm_status(admin: dict = Depends(require_admin)):
    """Get dashboard statistics for communications."""
    return communications.get_communications_status()


# --- Email Sending ---

@app.post("/api/communications/send-weekly")
async def send_weekly_comms(data: SendEmailRequest, admin: dict = Depends(require_admin)):
    """Send weekly follow-up emails to all eligible clients."""
    try:
        result = communications.send_batch_followups(
            template_type="weekly",
            test_mode=data.test_mode,
            test_recipient=data.test_recipient,
            sent_by=admin["email"]
        )
        logger.info(f"Weekly emails sent by {admin['email']}: {result}")
        return {"success": True, **result}
    except Exception as e:
        logger.error(f"Weekly email send failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/communications/send-monthly")
async def send_monthly_comms(data: SendEmailRequest, admin: dict = Depends(require_admin)):
    """Send monthly follow-up emails to all eligible clients."""
    try:
        result = communications.send_batch_followups(
            template_type="monthly",
            test_mode=data.test_mode,
            test_recipient=data.test_recipient,
            sent_by=admin["email"]
        )
        logger.info(f"Monthly emails sent by {admin['email']}: {result}")
        return {"success": True, **result}
    except Exception as e:
        logger.error(f"Monthly email send failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/communications/send-single")
async def send_single_comm(data: SendSingleEmailRequest, admin: dict = Depends(require_admin)):
    """Send a single email to a specific client."""
    try:
        client = communications.get_comm_client_by_id(data.client_id)
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")

        result = communications.send_followup_email(
            client=client,
            template_type=data.template_type,
            test_mode=data.test_mode,
            test_recipient=data.test_recipient,
            sent_by=admin["email"]
        )
        logger.info(f"Single email sent to {client.get('name')} by {admin['email']}: {result}")
        return result
    except Exception as e:
        logger.error(f"Single email send failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/api/communications/send-test")
async def send_test_comm(admin: dict = Depends(require_admin)):
    """Send a test email to verify configuration."""
    try:
        # Get first active client for test
        clients = communications.get_all_comm_clients()
        test_client = next((c for c in clients if c.get("email")), None)

        if not test_client:
            test_client = {
                "id": "test",
                "name": "Test Client",
                "email": admin["email"]
            }

        result = communications.send_followup_email(
            client=test_client,
            template_type="weekly",
            test_mode=True,
            test_recipient=admin["email"],
            sent_by=admin["email"]
        )
        logger.info(f"Test email sent to {admin['email']}")
        return result
    except Exception as e:
        logger.error(f"Test email send failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))


# --- Email Threads (Gmail IMAP) ---

@app.get("/api/communications/threads")
async def get_email_threads(client_id: Optional[str] = None, admin: dict = Depends(require_admin)):
    """Get email threads from center@casehub.app."""
    threads = communications.get_email_threads(client_id=client_id)
    return {"threads": threads, "count": len(threads)}


@app.get("/api/communications/threads/{thread_id}")
async def get_email_thread(thread_id: str, admin: dict = Depends(require_admin)):
    """Get a specific email thread with messages."""
    threads = communications.get_email_threads()
    thread = next((t for t in threads if t.get("id") == thread_id), None)
    if not thread:
        raise HTTPException(status_code=404, detail="Thread not found")

    messages = communications.get_thread_messages(thread_id)
    return {"thread": thread, "messages": messages}


@app.post("/api/communications/threads/sync")
async def sync_email_threads(client_email: Optional[str] = None, admin: dict = Depends(require_admin)):
    """Sync email threads from Gmail IMAP."""
    try:
        result = communications.sync_gmail_threads(client_email=client_email)
        logger.info(f"Email threads synced by {admin['email']}: {result}")
        return result
    except Exception as e:
        logger.error(f"Email thread sync failed: {e}")
        raise HTTPException(status_code=400, detail=str(e))


# --- Email Attachments ---

@app.get("/api/communications/threads/{thread_id}/messages")
async def get_thread_messages_with_attachments(thread_id: str, user: dict = Depends(require_auth)):
    """Get all messages in a thread with attachment info."""
    messages = communications.get_thread_messages(thread_id)

    # Enrich messages with attachment data
    for msg in messages:
        gmail_id = msg.get("gmail_message_id")
        if gmail_id:
            attachments = attachment_metadata.get_attachments_by_message_id(gmail_id)
            msg["has_attachments"] = len(attachments) > 0
            msg["attachment_count"] = len(attachments)
            msg["attachments"] = attachments

    return {"messages": messages, "count": len(messages)}


@app.get("/api/attachments/{attachment_id}")
async def get_attachment_info(attachment_id: str, user: dict = Depends(require_auth)):
    """Get attachment metadata."""
    att = attachment_metadata.get_attachment_by_id(attachment_id)
    if not att:
        raise HTTPException(status_code=404, detail="Attachment not found")
    return att


@app.get("/api/attachments/{attachment_id}/download")
async def download_attachment(attachment_id: str, user: dict = Depends(require_auth)):
    """Download an attachment file."""
    file_path = attachment_metadata.get_attachment_path(attachment_id)
    if not file_path:
        raise HTTPException(status_code=404, detail="Attachment not found or file missing")

    att = attachment_metadata.get_attachment_by_id(attachment_id)
    filename = att.get("original_filename", file_path.name) if att else file_path.name
    mime_type = att.get("mime_type", "application/octet-stream") if att else "application/octet-stream"

    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type=mime_type
    )


@app.get("/api/attachments/{attachment_id}/preview")
async def preview_attachment(attachment_id: str, user: dict = Depends(require_auth)):
    """Preview an attachment (for images/PDFs - displays inline)."""
    file_path = attachment_metadata.get_attachment_path(attachment_id)
    if not file_path:
        raise HTTPException(status_code=404, detail="Attachment not found or file missing")

    att = attachment_metadata.get_attachment_by_id(attachment_id)
    mime_type = att.get("mime_type", "application/octet-stream") if att else "application/octet-stream"

    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        headers={"Content-Disposition": "inline"}
    )


@app.get("/api/attachments")
async def list_attachments(
    client_email: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    user: dict = Depends(require_auth)
):
    """List all attachments with optional filtering."""
    if client_email:
        attachments = attachment_metadata.get_attachments_by_client(client_email)
        attachments = attachments[offset:offset + limit]
    else:
        attachments = attachment_metadata.get_all_attachments(limit=limit, offset=offset)

    stats = attachment_metadata.get_attachment_stats()
    return {
        "attachments": attachments,
        "count": len(attachments),
        "total": stats["total_count"]
    }


# --- Schedules ---

@app.get("/api/communications/schedules")
async def get_comm_schedules(admin: dict = Depends(require_admin)):
    """Get all communication schedules."""
    schedules = communications.get_schedules()
    return {"schedules": schedules}


@app.put("/api/communications/schedules/{schedule_id}")
async def update_comm_schedule(schedule_id: str, data: ScheduleUpdate, admin: dict = Depends(require_admin)):
    """Update a communication schedule."""
    try:
        updates = {k: v for k, v in data.dict().items() if v is not None}
        schedule = communications.update_schedule(schedule_id, updates)
        logger.info(f"Schedule {schedule_id} updated by {admin['email']}")
        return {"success": True, "schedule": schedule}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.post("/api/communications/schedules/{schedule_id}/pause")
async def pause_comm_schedule(schedule_id: str, admin: dict = Depends(require_admin)):
    """Pause a communication schedule."""
    try:
        schedule = communications.update_schedule(schedule_id, {"is_active": False})
        logger.info(f"Schedule {schedule_id} paused by {admin['email']}")
        return {"success": True, "schedule": schedule}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.post("/api/communications/schedules/{schedule_id}/resume")
async def resume_comm_schedule(schedule_id: str, admin: dict = Depends(require_admin)):
    """Resume a communication schedule."""
    try:
        schedule = communications.update_schedule(schedule_id, {"is_active": True})
        logger.info(f"Schedule {schedule_id} resumed by {admin['email']}")
        return {"success": True, "schedule": schedule}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


# --- Templates ---

@app.get("/api/communications/templates")
async def get_comm_templates(admin: dict = Depends(require_admin)):
    """Get all email templates."""
    templates_list = communications.get_templates()
    return {"templates": templates_list}


@app.put("/api/communications/templates/{template_id}")
async def update_comm_template(template_id: str, data: TemplateUpdate, admin: dict = Depends(require_admin)):
    """Update an email template."""
    try:
        updates = {k: v for k, v in data.dict().items() if v is not None}
        template = communications.update_template(template_id, updates)
        logger.info(f"Template {template_id} updated by {admin['email']}")
        return {"success": True, "template": template}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@app.post("/api/communications/templates/preview")
async def preview_comm_template(
    template_type: str = "weekly",
    client_name: str = "John Doe",
    admin: dict = Depends(require_admin)
):
    """Preview an email template with a sample client name."""
    template = communications.get_template_by_type(template_type)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    preview = {
        "subject": template["subject"],
        "body_html": template["body_html"].replace("{client_name}", client_name),
        "body_text": template["body_text"].replace("{client_name}", client_name)
    }
    return preview


# =============================================================================
# LOGIN PAGE
# =============================================================================

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """Serve login page."""
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    """Serve registration page."""
    return templates.TemplateResponse("login.html", {"request": request, "show_register": True})


@app.get("/admin", response_class=HTMLResponse)
async def admin_page(request: Request):
    """Serve admin panel page - PROTECTED (requires admin)."""
    # Server-side authentication check
    user = verify_page_auth(request)
    if not user:
        return RedirectResponse(url="/tools/login", status_code=302)

    # Check admin privileges
    if not user.get("is_admin") and not user.get("is_super_admin"):
        return RedirectResponse(url="/tools/dashboard", status_code=302)

    return templates.TemplateResponse("admin.html", {"request": request, "user": user})


# =============================================================================
# COVER PAGE GENERATOR
# =============================================================================

def create_cover_page(
    beneficiary_name: str,
    visa_type: str,
    case_number: str = None,
    date_str: str = None,
    output_path: str = None
) -> str:
    """Create a professional USCIS cover page PDF."""
    logger.info(f"Creating cover page for {beneficiary_name} - {visa_type}")

    c = rl_canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    # Border rectangle
    c.setStrokeColorRGB(0.3, 0.3, 0.3)
    c.setLineWidth(2)
    c.rect(0.75*inch, 0.75*inch, width - 1.5*inch, height - 1.5*inch)

    # Header - "PETITION FOR"
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width/2, height - 2.5*inch, "PETITION FOR")

    # Visa Type (large)
    c.setFont("Helvetica-Bold", 32)
    c.drawCentredString(width/2, height - 3.3*inch, visa_type)

    # Decorative line
    c.setLineWidth(1)
    c.line(width*0.25, height - 3.7*inch, width*0.75, height - 3.7*inch)

    # "ON BEHALF OF"
    c.setFont("Helvetica", 14)
    c.drawCentredString(width/2, height - 4.3*inch, "ON BEHALF OF")

    # Beneficiary Name (large)
    c.setFont("Helvetica-Bold", 26)
    c.drawCentredString(width/2, height - 5.1*inch, beneficiary_name.upper())

    # Case number if provided
    if case_number:
        c.setFont("Helvetica", 12)
        c.drawCentredString(width/2, height - 5.8*inch, f"Case Number: {case_number}")

    # Date
    c.setFont("Helvetica", 12)
    date_display = date_str or datetime.now().strftime("%B %d, %Y")
    c.drawCentredString(width/2, height - 7*inch, date_display)

    # Bottom decorative line
    c.line(width*0.25, height - 7.5*inch, width*0.75, height - 7.5*inch)

    c.save()
    logger.info(f"Cover page created: {output_path}")
    return output_path


# =============================================================================
# TABLE OF CONTENTS GENERATOR
# =============================================================================

def create_toc_page(
    exhibits_info: List[dict],
    output_path: str
) -> str:
    """Create a Table of Contents page.

    Args:
        exhibits_info: List of dicts with keys: letter, name, page_start
        output_path: Path to save the PDF
    """
    logger.info(f"Creating TOC with {len(exhibits_info)} exhibits")

    c = rl_canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    # Title
    c.setFont("Helvetica-Bold", 20)
    c.drawCentredString(width/2, height - 1.5*inch, "TABLE OF CONTENTS")

    # Separator line
    c.setStrokeColorRGB(0.3, 0.3, 0.3)
    c.setLineWidth(1)
    c.line(1*inch, height - 1.8*inch, width - 1*inch, height - 1.8*inch)

    # Column headers
    y = height - 2.3*inch
    c.setFont("Helvetica-Bold", 12)
    c.drawString(1*inch, y, "EXHIBIT")
    c.drawString(2*inch, y, "DESCRIPTION")
    c.drawRightString(width - 1*inch, y, "PAGE")

    # Separator line under headers
    y -= 0.15*inch
    c.line(1*inch, y, width - 1*inch, y)

    # TOC entries
    y -= 0.4*inch
    c.setFont("Helvetica", 11)

    for exhibit in exhibits_info:
        if y < 1*inch:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 1.5*inch

        letter_code = exhibit.get('letter', '')
        name = exhibit.get('name', '')
        page_start = exhibit.get('page_start', '')

        # Exhibit letter
        c.drawString(1*inch, y, f"Exhibit {letter_code}")

        # Description
        c.drawString(2*inch, y, name)

        # Page number (right-aligned)
        c.drawRightString(width - 1*inch, y, str(page_start))

        # Dotted line connecting name to page
        name_width = c.stringWidth(name, "Helvetica", 11)
        page_width = c.stringWidth(str(page_start), "Helvetica", 11)
        dot_start = 2*inch + name_width + 10
        dot_end = width - 1*inch - page_width - 10

        if dot_end > dot_start:
            c.setDash(1, 3)
            c.line(dot_start, y + 3, dot_end, y + 3)
            c.setDash()

        y -= 0.35*inch

    c.save()
    logger.info(f"TOC created: {output_path}")
    return output_path


# =============================================================================
# PAGE NUMBERING
# =============================================================================

def add_page_numbers(input_path: str, output_path: str, start_page: int = 1) -> str:
    """Add page numbers to all pages of a PDF.

    Args:
        input_path: Path to input PDF
        output_path: Path to save numbered PDF
        start_page: Starting page number (default 1)
    """
    logger.info(f"Adding page numbers to {input_path}")

    reader = PdfReader(input_path)
    writer = PdfWriter()
    total_pages = len(reader.pages)

    for i, page in enumerate(reader.pages):
        page_num = start_page + i

        # Create overlay with page number
        packet = io.BytesIO()
        c = rl_canvas.Canvas(packet, pagesize=letter)
        width, height = letter

        # Footer: "Page X of Y"
        c.setFont("Helvetica", 10)
        c.setFillColorRGB(0.4, 0.4, 0.4)
        c.drawCentredString(width/2, 0.4*inch, f"Page {page_num} of {total_pages}")
        c.save()

        packet.seek(0)
        overlay = PdfReader(packet)

        # Merge overlay with page
        page.merge_page(overlay.pages[0])
        writer.add_page(page)

    with open(output_path, 'wb') as f:
        writer.write(f)

    logger.info(f"Page numbers added: {output_path}")
    return output_path


# =============================================================================
# WATERMARK FUNCTION
# =============================================================================

def add_watermark(input_path: str, output_path: str, text: str = "DRAFT") -> str:
    """Add diagonal watermark to all pages."""
    logger.info(f"Adding watermark '{text}' to {input_path}")

    reader = PdfReader(input_path)
    writer = PdfWriter()

    for page in reader.pages:
        packet = io.BytesIO()
        c = rl_canvas.Canvas(packet, pagesize=letter)
        width, height = letter

        # Watermark - diagonal, semi-transparent
        c.saveState()
        c.setFillColorRGB(0.85, 0.85, 0.85)  # Light gray
        c.setFont("Helvetica-Bold", 72)
        c.translate(width/2, height/2)
        c.rotate(45)
        c.drawCentredString(0, 0, text)
        c.restoreState()
        c.save()

        packet.seek(0)
        watermark_pdf = PdfReader(packet)
        page.merge_page(watermark_pdf.pages[0])
        writer.add_page(page)

    with open(output_path, 'wb') as f:
        writer.write(f)

    logger.info(f"Watermark added: {output_path}")
    return output_path


# =============================================================================
# MAIN PAGE
# =============================================================================

# =============================================================================
# MAIN PAGE
# =============================================================================

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Landing page."""
    return templates.TemplateResponse("landing.html", {"request": request})

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):
    """Main tools interface - PROTECTED."""
    # Server-side authentication check
    user = verify_page_auth(request)
    if not user:
        # Redirect to login if not authenticated
        # Use full path since nginx proxies /tools/ to this app
        return RedirectResponse(url="/tools/login", status_code=302)

    return templates.TemplateResponse("index.html", {
        "request": request,
        "personas": PERSONAS,
        "fields": list(NATIONAL_INTEREST_TEXTS.keys()),
        "exhibits": EXHIBITS,
        "user": user,  # Pass user info to template
        # Google Drive Picker browser credentials — injected from env, never hardcoded.
        # The Picker API key MUST be HTTP-referrer restricted in the Cloud Console.
        "google_api_key": os.getenv("GOOGLE_DRIVE_API_KEY", ""),
        "google_client_id": os.getenv("GOOGLE_OAUTH_CLIENT_ID", ""),
        "google_app_id": os.getenv("GOOGLE_DRIVE_APP_ID", "152327852610"),
    })


@app.get("/package-builder", response_class=HTMLResponse)
async def package_builder_page(request: Request):
    """Package Builder UI - PROTECTED."""
    user = verify_page_auth(request)
    if not user:
        return RedirectResponse(url="/tools/login", status_code=302)

    return templates.TemplateResponse("package_builder.html", {
        "request": request,
        "exhibits": EXHIBITS,
        "user": user,
    })


# =============================================================================
# DEBUG ENDPOINTS (SECURED - Production should disable these)
# =============================================================================

# Load debug token from environment - NEVER use default in production
DEBUG_TOKEN = os.getenv("DEBUG_TOKEN")
ENVIRONMENT = os.getenv("ENVIRONMENT", "development")
SUPER_ADMIN_IP = os.getenv("SUPER_ADMIN_IP", "")

def verify_debug_access(request: Request, token: str):
    """Verify debug access with token and optional IP whitelist."""
    # In production, debug endpoints are disabled unless explicitly enabled
    if ENVIRONMENT == "production" and not DEBUG_TOKEN:
        raise HTTPException(status_code=404, detail="Not found")

    if not DEBUG_TOKEN or not secrets.compare_digest(token, DEBUG_TOKEN):
        logger.warning(f"Unauthorized debug access attempt from {request.client.host}")
        raise HTTPException(status_code=403, detail="Unauthorized")

    # Optional: IP whitelist for extra security
    if SUPER_ADMIN_IP and request.client.host != SUPER_ADMIN_IP:
        logger.warning(f"Debug access from non-whitelisted IP: {request.client.host}")
        raise HTTPException(status_code=403, detail="IP not authorized")


# =============================================================================
# HEALTH CHECK ENDPOINT (PUBLIC)
# =============================================================================

@app.get("/api/health")
async def health_check():
    """Public health check endpoint for monitoring and integrations."""
    return {
        "status": "healthy",
        "version": "2.1.0-alpha",
        "service": "ilc-tools"
    }


@app.get("/api/debug/logs")
async def get_logs(request: Request, lines: int = 100, token: str = ""):
    """Get recent log entries for debugging. PROTECTED."""
    verify_debug_access(request, token)

    # Limit lines to prevent DoS
    lines = min(lines, 500)

    log_file = LOG_DIR / f"ilc_{datetime.now().strftime('%Y%m%d')}.log"
    if not log_file.exists():
        return {"logs": [], "message": "No logs for today"}

    with open(log_file, 'r') as f:
        all_lines = f.readlines()
        recent = all_lines[-lines:]

    return {"logs": recent, "total_lines": len(all_lines)}


@app.get("/api/debug/status")
async def get_status(request: Request, token: str = ""):
    """Get system status. PROTECTED."""
    verify_debug_access(request, token)

    output_files = list(OUTPUT_DIR.glob("*.pdf"))
    upload_files = list(UPLOAD_DIR.glob("*"))
    log_files = list(LOG_DIR.glob("*.log"))

    # Don't expose full paths in response
    return {
        "status": "running",
        "version": "2.1.0-alpha",
        "counts": {
            "output_files": len(output_files),
            "upload_files": len(upload_files),
            "log_files": len(log_files),
        },
        "environment": ENVIRONMENT
    }


# =============================================================================
# LOR ENDPOINTS
# =============================================================================

def docx_to_pdf(docx_path: str, pdf_path: Optional[str] = None) -> str:
    """Convert a DOCX file to PDF using reportlab.

    Args:
        docx_path: Path to input DOCX file
        pdf_path: Optional output PDF path. If None, auto-generates as docx_path + ".pdf"

    Returns:
        Path to the generated PDF file
    """
    from docx import Document as DocxDoc

    # Auto-generate output path if not provided
    if pdf_path is None:
        pdf_path = str(docx_path) + ".pdf"

    doc = DocxDoc(docx_path)

    c = rl_canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    margin = 1 * inch
    y = height - margin
    line_height = 14

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y -= line_height
            continue

        # Check for bold (header)
        is_bold = any(run.bold for run in para.runs if run.bold)
        font_name = "Helvetica-Bold" if is_bold else "Helvetica"
        font_size = 11

        c.setFont(font_name, font_size)

        # Word wrap
        words = text.split()
        current_line = ""
        for word in words:
            test_line = current_line + " " + word if current_line else word
            if c.stringWidth(test_line, font_name, font_size) < (width - 2 * margin):
                current_line = test_line
            else:
                if y < margin:
                    c.showPage()
                    y = height - margin
                    c.setFont(font_name, font_size)
                c.drawString(margin, y, current_line)
                y -= line_height
                current_line = word

        if current_line:
            if y < margin:
                c.showPage()
                y = height - margin
                c.setFont(font_name, font_size)
            c.drawString(margin, y, current_line)
            y -= line_height

        y -= 6  # Paragraph spacing

    c.save()
    logger.info(f"Converted DOCX to PDF: {pdf_path}")
    return pdf_path


async def detect_location_context(name: str, org: str, api_key: str) -> str:
    """Detect if recommender is US-based or International using Perplexity."""
    try:
        if not api_key:
            return "us"  # Default to US if no key

        prompt = f"""Is the organization "{org}" (associated with {name}) primarily a US-based entity or an International/Non-US entity? 
        If it's a multinational with a major US presence, count as US-based.
        Return ONLY valid JSON: {{ "location": "us" }} or {{ "location": "international" }}."""

        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "sonar",
                    "messages": [
                        {"role": "system", "content": "You are a location detector. Return only valid JSON."},
                        {"role": "user", "content": prompt}
                    ]
                }
            )
            
            if response.status_code == 200:
                content = response.json()['choices'][0]['message']['content']
                import json
                import re
                match = re.search(r'\{.*\}', content, re.DOTALL)
                if match:
                    data = json.loads(match.group(0))
                    return data.get("location", "us")
    except Exception as e:
        logger.error(f"Location detection failed: {e}")
    
    return "us"  # Fallback

@app.post("/api/lor/generate")
async def generate_lor(
    persona: str = Form(...),
    beneficiary_name: str = Form(...),
    visa_type: str = Form("EB-2 NIW"),
    field: str = Form(None),
    recommender_name: str = Form(...),
    recommender_title: str = Form(...),
    recommender_org: str = Form(...),
    recommender_email: str = Form(...),
    recommender_phone: str = Form(None),
    relationship: str = Form(...),
    years_known: str = Form("several"),
    custom_paragraphs: str = Form(""),
    international_context: str = Form("auto"),
    include_national_importance: bool = Form(True),
    include_prong3: bool = Form(True),
    output_format: str = Form("docx"),
    salutation: str = Form(None),
    closing: str = Form(None),
):
    """Generate a Letter of Recommendation."""
    logger.info(f"Generating LOR for {beneficiary_name} - persona: {persona}, format: {output_format}")
    try:
        generator = LORGenerator(
            persona=persona,
            beneficiary_name=beneficiary_name,
            visa_type=visa_type,
            field=field if field else None,
            output_dir=str(OUTPUT_DIR),
        )

        # Resolve Auto Context
        if international_context == "auto":
            # Detect
            international_context = await detect_location_context(recommender_name, recommender_org, PERPLEXITY_API_KEY)
            logger.info(f"Auto-detected location context for {recommender_org}: {international_context}")

        # Build paragraphs
        paragraphs = []

        # Opening
        paragraphs.append(generator.get_opening_paragraph(relationship, years_known))

        # Custom paragraphs
        if custom_paragraphs.strip():
            for para in custom_paragraphs.strip().split("\n\n"):
                if para.strip():
                    paragraphs.append(para.strip())

        # National importance
        if include_national_importance and field:
            ni_para = generator.get_national_importance_paragraph()
            if ni_para:
                paragraphs.append(ni_para)

        # Prong 3
        if include_prong3 and field:
            paragraphs.append(generator.get_prong3_paragraph())

        # Add Location Context Paragraph
        if international_context == "international":
            paragraphs.append(f"From my international perspective as an expert at {recommender_org}, I can attest that {beneficiary_name}'s work has achieved global recognition and is frequently cited by researchers worldwide, demonstrating an influence that extends far beyond the United States.")
        elif international_context == "us":
             # Optional: Add domestic focus if needed, or leave implied
             pass

        # Conclusion
        paragraphs.append(f"I strongly support the approval of {beneficiary_name}'s {visa_type} petition.")

        # Generate document (always creates DOCX first)
        docx_filepath = generator.create_document(
            recommender_name=recommender_name,
            recommender_title=recommender_title,
            recommender_org=recommender_org,
            recommender_email=recommender_email,
            recommender_phone=recommender_phone if recommender_phone else None,
            relationship=relationship,
            paragraphs=paragraphs,
            custom_salutation=salutation if salutation else None,
            custom_closing=closing if closing else None,
        )

        # Convert to PDF if requested
        if output_format == "pdf":
            pdf_filepath = docx_filepath.replace(".docx", ".pdf")
            docx_to_pdf(docx_filepath, pdf_filepath)
            # Remove the DOCX file, keep only PDF
            os.remove(docx_filepath)
            filepath = pdf_filepath
        else:
            filepath = docx_filepath

        filename = Path(filepath).name
        logger.info(f"LOR generated: {filename}")
        return JSONResponse({
            "success": True,
            "filename": filename,
            "download_url": f"/download/{filename}",
        })

    except Exception as e:
        logger.error(f"LOR generation failed: {str(e)}", exc_info=True)
        return JSONResponse({
            "success": False,
            "error": str(e),
        }, status_code=400)


# =============================================================================
# PS ENDPOINTS
# =============================================================================

@app.post("/api/ps/generate")
async def generate_ps(
    beneficiary_name: str = Form(...),
    field: str = Form(...),
    overview: str = Form(""),
    national_importance: str = Form(""),
    practical_impact: str = Form(""),
    well_positioned: str = Form(""),
    conclusion: str = Form(""),
    salutation: str = Form(None),
    closing: str = Form(None),
):
    """Generate a Personal Statement."""
    try:
        generator = PSGenerator(
            beneficiary_name=beneficiary_name,
            field=field,
            output_dir=str(OUTPUT_DIR),
        )

        # Use templates if sections are empty
        sections = {
            "overview": overview.strip() or generator.get_section_template("overview"),
            "national_importance": national_importance.strip() or generator.get_section_template("national_importance"),
            "practical_impact": practical_impact.strip() or generator.get_section_template("practical_impact"),
            "well_positioned": well_positioned.strip() or generator.get_section_template("well_positioned"),
            "conclusion": conclusion.strip() or generator.get_section_template("conclusion"),
        }

        filepath = generator.create_document(
            sections=sections,
            salutation=salutation if salutation else None,
            closing=closing if closing else None,
        )
        filename = Path(filepath).name

        return JSONResponse({
            "success": True,
            "filename": filename,
            "download_url": f"/download/{filename}",
        })

    except Exception as e:
        return JSONResponse({
            "success": False,
            "error": str(e),
        }, status_code=400)


# =============================================================================
# FILE CONVERSION UTILITIES
# =============================================================================

@app.post("/api/ps/extract")
async def extract_ps_data(
    file: UploadFile = File(...),
    context: str = Form(""),
    fix_errors: str = Form("false"),  # JS FormData sends boolean as string
    complete_sections: str = Form("false"),
    enhance_language: str = Form("false"),
):
    """Extract and generate Personal Statement content from an uploaded document."""
    logger.info(f"Extracting PS data from {file.filename}")
    
    # Convert string booleans
    do_fix = fix_errors.lower() == 'true'
    do_complete = complete_sections.lower() == 'true'
    do_enhance = enhance_language.lower() == 'true'
    
    temp_path = None
    try:
        # Save file temporarily
        ext = Path(file.filename).suffix.lower()
        if ext not in ['.docx', '.pdf', '.txt', '.doc']:
             return JSONResponse({"success": False, "error": "Unsupported file format"}, status_code=400)
             
        temp_path = UPLOAD_DIR / f"ps_extract_{uuid.uuid4()}{ext}"
        with open(temp_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
            
        # Extract text logic (reusing extract_cv_text logic or similar)
        content = ""
        if ext == '.docx':
            doc = DocxDocument(temp_path)
            content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        elif ext == '.txt':
            with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        elif ext == '.pdf':
            reader = PdfReader(temp_path)
            content = '\n'.join([page.extract_text() or '' for page in reader.pages])
        
        if not content.strip():
             return JSONResponse({"success": False, "error": "Could not extract text from file"}, status_code=400)

        # Truncate content if too long for API
        content = content[:25000] 

        # Build Prompt
        prompt = f"""You are an expert immigration attorney assistant. Analyze the provided document text and extract information to draft or improve a Personal Statement for an EB-2 National Interest Waiver (NIW) or EB-1A petition.
        
        CONTEXT provided by user: "{context}"
        DOCUMENT TEXT:
        {content}
        
        INSTRUCTIONS:
        1. Extract the beneficiary name and field of endeavor.
        2. Draft or refine the following sections based on the content:
           - Overview (Professional Summary)
           - National Importance (Prong 1: specific endeavors, broader impact)
           - Practical Impact (Real world benefits)
           - Well Positioned (Prong 2: education, experience, record)
           - Conclusion
        { "3. Fix grammar and spelling errors." if do_fix else "" }
        { "4. Complete any missing logic or sparse sections using reasonable professional inferences." if do_complete else "" }
        { "5. Enhance the language to be more persuasive and professional." if do_enhance else "" }

        Output strictly valid JSON with no markdown formatting:
        {{
            "beneficiary_name": "Name found",
            "field": "Field found",
            "overview": "text content...",
            "national_importance": "text content...",
            "practical_impact": "text content...",
            "well_positioned": "text content...",
            "conclusion": "text content...",
            "suggestions": ["suggestion 1", "suggestion 2", "suggestion 3"]
        }}
        """
        
        # Call Perplexity
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers={
                    "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "sonar",
                    "messages": [
                        {"role": "system", "content": "You are an expert legal assistant. Return only valid JSON."},
                        {"role": "user", "content": prompt}
                    ]
                }
            )
            
            if response.status_code != 200:
                logger.error(f"AI API Error: {response.text}")
                return JSONResponse({"success": False, "error": "AI service unavailable"}, status_code=500)
                
            data = response.json()
            completion = data['choices'][0]['message']['content']
            
            # Parse JSON
            import json
            import re
            
            # Find JSON block
            match = re.search(r'\{[\s\S]*\}', completion)
            if not match:
                logger.error(f"Could not parse JSON from AI response: {completion[:200]}...")
                return JSONResponse({"success": False, "error": "AI response was not valid JSON"}, status_code=500)
                
            extracted_data = json.loads(match.group(0))
            
            return JSONResponse({
                "success": True,
                "extracted": extracted_data
            })

    except Exception as e:
        logger.error(f"Extract failed: {str(e)}", exc_info=True)
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)
    finally:
        if temp_path and temp_path.exists():
            os.remove(temp_path)
            


def image_to_pdf_bytes(image_path: str) -> str:
    """Convert an image to PDF and return the output path."""
    output_path = str(image_path) + ".pdf"
    img = Image.open(image_path)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(output_path, "PDF", resolution=100.0)
    return output_path


def txt_to_pdf(txt_path: str) -> str:
    """Convert a text file to PDF."""
    output_path = str(txt_path) + ".pdf"

    with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()

    c = rl_canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    # Settings
    margin = 1 * inch
    line_height = 14
    max_width = width - 2 * margin
    y = height - margin

    c.setFont("Helvetica", 11)

    for line in text.split('\n'):
        if y < margin:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - margin

        # Wrap long lines
        while len(line) > 90:
            c.drawString(margin, y, line[:90])
            line = line[90:]
            y -= line_height
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - margin

        c.drawString(margin, y, line)
        y -= line_height

    c.save()
    return output_path




def convert_file_to_pdf(file_path: str) -> str:
    """Convert any supported file to PDF."""
    ext = Path(file_path).suffix.lower()

    if ext == '.pdf':
        return file_path
    elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.tiff', '.bmp']:
        return image_to_pdf_bytes(file_path)
    elif ext == '.docx':
        return docx_to_pdf(file_path)
    elif ext == '.txt':
        return txt_to_pdf(file_path)
    elif ext == '.doc':
        raise ValueError("DOC format not supported. Please convert to DOCX first.")
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# =============================================================================
# PDF TOOLS ENDPOINTS
# =============================================================================

@app.post("/api/pdf/merge")
async def merge_files(files: List[UploadFile] = File(...)):
    """Merge multiple files (PDF, DOCX, TXT, images) into a single PDF."""
    try:
        # Supported extensions
        supported_ext = {'.pdf', '.jpg', '.jpeg', '.png', '.gif', '.webp', '.tiff', '.bmp', '.docx', '.txt'}

        # Save uploaded files
        temp_files = []
        converted_files = []

        for file in files:
            ext = Path(file.filename).suffix.lower()
            if ext not in supported_ext:
                continue

            temp_path = UPLOAD_DIR / f"{uuid.uuid4()}_{file.filename}"
            with open(temp_path, "wb") as f:
                shutil.copyfileobj(file.file, f)
            temp_files.append(str(temp_path))

            # Convert to PDF
            try:
                pdf_path = convert_file_to_pdf(str(temp_path))
                converted_files.append(pdf_path)
            except Exception as e:
                # Cleanup and return error
                for tf in temp_files:
                    if os.path.exists(tf):
                        os.remove(tf)
                return JSONResponse({
                    "success": False,
                    "error": f"Failed to convert {file.filename}: {str(e)}",
                }, status_code=400)

        if len(converted_files) < 1:
            return JSONResponse({
                "success": False,
                "error": "No valid files to merge. Supported: PDF, DOCX, TXT, JPG, PNG, GIF, WebP",
            }, status_code=400)

        # Merge all PDFs
        output_filename = f"merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = OUTPUT_DIR / output_filename

        merger = PdfMerger()
        for pdf_path in converted_files:
            merger.append(pdf_path)
        merger.write(str(output_path))
        merger.close()

        # Cleanup temp and converted files
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.remove(temp_file)
        for conv_file in converted_files:
            if conv_file not in temp_files and os.path.exists(conv_file):
                os.remove(conv_file)

        return JSONResponse({
            "success": True,
            "filename": output_filename,
            "download_url": f"/download/{output_filename}",
        })

    except Exception as e:
        return JSONResponse({
            "success": False,
            "error": str(e),
        }, status_code=400)


@app.post("/api/pdf/images-to-pdf")
async def images_to_pdf(files: List[UploadFile] = File(...)):
    """Convert images to PDF."""
    try:
        # Save uploaded files
        temp_files = []
        valid_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.tiff'}

        for file in files:
            ext = Path(file.filename).suffix.lower()
            if ext not in valid_extensions:
                continue
            temp_path = UPLOAD_DIR / f"{uuid.uuid4()}_{file.filename}"
            with open(temp_path, "wb") as f:
                shutil.copyfileobj(file.file, f)
            temp_files.append(str(temp_path))

        if not temp_files:
            return JSONResponse({
                "success": False,
                "error": "No valid image files uploaded",
            }, status_code=400)

        # Convert
        output_filename = f"images_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = OUTPUT_DIR / output_filename
        convert_images_to_pdf(temp_files, str(output_path))

        # Cleanup
        for temp_file in temp_files:
            os.remove(temp_file)

        return JSONResponse({
            "success": True,
            "filename": output_filename,
            "download_url": f"/download/{output_filename}",
        })

    except Exception as e:
        return JSONResponse({
            "success": False,
            "error": str(e),
        }, status_code=400)


# =============================================================================
# PACKAGE BUILDER ENDPOINT
# =============================================================================

def create_separator_page(title: str, subtitle: str, output_path: str) -> str:
    """Create a separator/title page PDF."""
    c = rl_canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    # Center the content vertically
    center_y = height / 2

    # Title (large, bold)
    c.setFont("Helvetica-Bold", 36)
    c.drawCentredString(width / 2, center_y + 40, title)

    # Subtitle
    c.setFont("Helvetica", 24)
    c.drawCentredString(width / 2, center_y - 20, subtitle)

    # Decorative line
    c.setStrokeColorRGB(0.3, 0.3, 0.3)
    c.setLineWidth(2)
    c.line(width * 0.25, center_y - 60, width * 0.75, center_y - 60)

    c.save()
    return output_path


@app.post("/api/package/build")
async def build_package(
    files: List[UploadFile] = File(default=[]),
    structure: str = Form(...),
    include_cover: bool = Form(False),
    beneficiary_name: str = Form(""),
    visa_type: str = Form("EB-2 NIW"),
    case_number: str = Form(""),
    include_toc: bool = Form(False),
    add_page_nums: bool = Form(False),
    add_watermark_text: str = Form("")
):
    """Build a complete package with separators, cover page, TOC, and page numbers."""
    import json

    logger.info(f"Building package - include_cover={include_cover}, include_toc={include_toc}, page_nums={add_page_nums}")

    try:
        structure_data = json.loads(structure)
        temp_files = []
        pdf_files_to_merge = []
        exhibits_info = []  # For TOC generation

        # Track page numbers for each exhibit
        current_page = 1

        # Reserve pages for cover and TOC if included
        if include_cover:
            current_page += 1  # Cover is 1 page
        if include_toc:
            current_page += 1  # TOC is typically 1 page

        # Process each exhibit
        for exhibit in structure_data:
            letter_code = exhibit.get('letter', '')
            name = exhibit.get('name', '')
            separator = exhibit.get('separator')
            exhibit_files = exhibit.get('files', [])

            # Track this exhibit's start page
            exhibit_start_page = current_page

            # Create separator page if exists
            if separator:
                sep_title = separator.get('title', f'EXHIBIT {letter_code}')
                sep_subtitle = separator.get('subtitle', name)
                sep_path = UPLOAD_DIR / f"sep_{uuid.uuid4()}.pdf"
                create_separator_page(sep_title, sep_subtitle, str(sep_path))
                pdf_files_to_merge.append(str(sep_path))
                temp_files.append(str(sep_path))
                current_page += 1

            # Process files for this exhibit
            exhibit_page_count = 0
            for file_info in exhibit_files:
                file_index = file_info.get('index', 0)
                if file_index < len(files):
                    file = files[file_index]

                    # Save uploaded file
                    temp_path = UPLOAD_DIR / f"{uuid.uuid4()}_{file.filename}"
                    with open(temp_path, "wb") as f:
                        shutil.copyfileobj(file.file, f)
                    temp_files.append(str(temp_path))

                    # Convert to PDF
                    try:
                        pdf_path = convert_file_to_pdf(str(temp_path))
                        pdf_files_to_merge.append(pdf_path)
                        if pdf_path != str(temp_path):
                            temp_files.append(pdf_path)

                        # Count pages in this file
                        reader = PdfReader(pdf_path)
                        file_pages = len(reader.pages)
                        current_page += file_pages
                        exhibit_page_count += file_pages
                        logger.info(f"Added {file.filename} - {file_pages} pages")

                    except Exception as e:
                        logger.error(f"Failed to convert {file.filename}: {str(e)}")
                        for tf in temp_files:
                            if os.path.exists(tf):
                                os.remove(tf)
                        return JSONResponse({
                            "success": False,
                            "error": f"Failed to convert {file.filename}: {str(e)}",
                        }, status_code=400)

            # Add to TOC info if this exhibit has content
            if separator or exhibit_files:
                exhibits_info.append({
                    'letter': letter_code,
                    'name': name,
                    'page_start': exhibit_start_page,
                    'page_count': exhibit_page_count + (1 if separator else 0)
                })

        if len(pdf_files_to_merge) < 1 and not include_cover:
            return JSONResponse({
                "success": False,
                "error": "No files to merge in package",
            }, status_code=400)

        # Create cover page if requested
        cover_path = None
        if include_cover and beneficiary_name:
            cover_path = UPLOAD_DIR / f"cover_{uuid.uuid4()}.pdf"
            create_cover_page(
                beneficiary_name=beneficiary_name,
                visa_type=visa_type,
                case_number=case_number if case_number else None,
                output_path=str(cover_path)
            )
            temp_files.append(str(cover_path))

        # Create TOC if requested
        toc_path = None
        if include_toc and exhibits_info:
            toc_path = UPLOAD_DIR / f"toc_{uuid.uuid4()}.pdf"
            create_toc_page(exhibits_info, str(toc_path))
            temp_files.append(str(toc_path))

        # Build final list: cover + TOC + exhibits
        final_pdf_list = []
        if cover_path:
            final_pdf_list.append(str(cover_path))
        if toc_path:
            final_pdf_list.append(str(toc_path))
        final_pdf_list.extend(pdf_files_to_merge)

        # Merge all PDFs
        output_filename = f"package_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = OUTPUT_DIR / output_filename

        merger = PdfMerger()
        total_pages = 0

        for pdf_path in final_pdf_list:
            try:
                merger.append(pdf_path)
                reader = PdfReader(pdf_path)
                total_pages += len(reader.pages)
            except Exception as e:
                logger.error(f"Failed to merge PDF {pdf_path}: {str(e)}")
                for tf in temp_files:
                    if os.path.exists(tf):
                        os.remove(tf)
                return JSONResponse({
                    "success": False,
                    "error": f"Failed to merge PDF: {str(e)}",
                }, status_code=400)

        merger.write(str(output_path))
        merger.close()
        logger.info(f"Package merged: {output_filename} ({total_pages} pages)")

        # Add page numbers if requested
        if add_page_nums:
            numbered_path = OUTPUT_DIR / f"numbered_{output_filename}"
            add_page_numbers(str(output_path), str(numbered_path))
            os.remove(str(output_path))
            os.rename(str(numbered_path), str(output_path))

        # Add watermark if requested
        if add_watermark_text:
            watermarked_path = OUTPUT_DIR / f"watermarked_{output_filename}"
            add_watermark(str(output_path), str(watermarked_path), add_watermark_text)
            os.remove(str(output_path))
            os.rename(str(watermarked_path), str(output_path))

        # Cleanup temp files
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.remove(temp_file)

        logger.info(f"Package built successfully: {output_filename}")

        return JSONResponse({
            "success": True,
            "filename": output_filename,
            "download_url": f"/download/{output_filename}",
            "page_count": total_pages,
            "exhibits_count": len(exhibits_info),
        })

    except Exception as e:
        logger.error(f"Package build failed: {str(e)}", exc_info=True)
        return JSONResponse({
            "success": False,
            "error": str(e),
        }, status_code=400)


# =============================================================================
# PACKAGE VALIDATION ENDPOINT
# =============================================================================

@app.post("/api/package/validate")
async def validate_package(structure: str = Form(...)):
    """
    Validate package structure before building.

    Checks for:
    - Required exhibits (A, B, C, D)
    - Recommended number of LORs in Exhibit D
    - Common issues and warnings
    """
    import json

    logger.info("Validating package structure")

    try:
        structure_data = json.loads(structure)

        issues = []      # Critical issues that should be fixed
        warnings = []    # Recommendations
        info = []        # Informational messages

        # Collect present exhibits
        present_exhibits = []
        exhibit_details = {}

        for exhibit in structure_data:
            letter = exhibit.get('letter', '')
            files = exhibit.get('files', [])
            separator = exhibit.get('separator')

            if files or separator:
                present_exhibits.append(letter)
                exhibit_details[letter] = {
                    'file_count': len(files),
                    'has_separator': bool(separator),
                    'name': exhibit.get('name', '')
                }

        # Check required exhibits (based on USCIS standard practice)
        required_exhibits = {
            'A': 'Forms (I-140, G-28, etc.)',
            'B': 'Brief/Cover Letter',
            'C': 'Self Petitioner Information (CV, diplomas)',
            'D': 'Letters of Recommendation',
        }

        for letter, description in required_exhibits.items():
            if letter not in present_exhibits:
                issues.append({
                    "type": "missing_required",
                    "exhibit": letter,
                    "message": f"Missing required Exhibit {letter}: {description}"
                })

        # Check Exhibit D (LORs) - recommend 3-6 letters
        if 'D' in exhibit_details:
            lor_count = exhibit_details['D']['file_count']
            if lor_count < 3:
                warnings.append({
                    "type": "insufficient_lors",
                    "exhibit": "D",
                    "count": lor_count,
                    "message": f"Exhibit D has only {lor_count} LOR(s). USCIS recommends 3-6 strong letters."
                })
            elif lor_count > 8:
                info.append({
                    "type": "many_lors",
                    "exhibit": "D",
                    "count": lor_count,
                    "message": f"Exhibit D has {lor_count} LORs. Consider quality over quantity (5-6 is typical)."
                })

        # Check for evidence exhibits (at least some should have content)
        evidence_exhibits = ['E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M']
        evidence_count = sum(1 for e in evidence_exhibits if e in present_exhibits)

        if evidence_count < 2:
            warnings.append({
                "type": "limited_evidence",
                "message": f"Only {evidence_count} evidence exhibit(s) included. Consider adding more supporting evidence (Exhibits E-M)."
            })

        # Check for separator pages
        exhibits_without_separators = [
            letter for letter, details in exhibit_details.items()
            if not details['has_separator'] and details['file_count'] > 0
        ]

        if exhibits_without_separators:
            info.append({
                "type": "no_separators",
                "exhibits": exhibits_without_separators,
                "message": f"Exhibits {', '.join(exhibits_without_separators)} have no separator pages. Consider adding them for clarity."
            })

        # Summary stats
        total_files = sum(d['file_count'] for d in exhibit_details.values())
        exhibits_with_content = len([e for e in exhibit_details.values() if e['file_count'] > 0])

        logger.info(f"Package validation complete: {len(issues)} issues, {len(warnings)} warnings")

        return JSONResponse({
            "valid": len(issues) == 0,
            "issues": issues,
            "warnings": warnings,
            "info": info,
            "summary": {
                "exhibits_present": present_exhibits,
                "exhibits_with_content": exhibits_with_content,
                "total_files": total_files,
                "required_present": [e for e in required_exhibits.keys() if e in present_exhibits],
                "required_missing": [e for e in required_exhibits.keys() if e not in present_exhibits],
            }
        })

    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in package structure: {str(e)}")
        return JSONResponse({
            "valid": False,
            "issues": [{"type": "invalid_json", "message": f"Invalid package structure: {str(e)}"}],
            "warnings": [],
            "info": [],
        }, status_code=400)
    except Exception as e:
        logger.error(f"Package validation failed: {str(e)}", exc_info=True)
        return JSONResponse({
            "valid": False,
            "issues": [{"type": "error", "message": str(e)}],
            "warnings": [],
            "info": [],
        }, status_code=400)


# =============================================================================
# RECOMMENDER PROFILE ANALYZER
# =============================================================================

import httpx

# Load from environment variable (never hardcode secrets!)
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY", "")

def extract_cv_text(file_path: str) -> str:
    """Extract text content from CV file (DOCX, PDF, TXT)."""
    ext = Path(file_path).suffix.lower()

    if ext == '.docx':
        doc = DocxDocument(file_path)
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or '')
        return '\n'.join(text)
    return ""


async def analyze_recommender_profile(name: str, title: str, organization: str, cv_text: str = "", style_description: str = "") -> dict:
    """Analyze recommender and generate a writing profile based on public information."""

    # Build search query
    search_parts = [name]
    if title:
        search_parts.append(title)
    if organization:
        search_parts.append(organization)
    search_query = ' '.join(search_parts)

    # Build prompt for analysis
    cv_section = f"\n\nCV/Background Information:\n{cv_text[:3000]}" if cv_text else ""

    # User description with professionalism guardrails
    style_section = ""
    if style_description:
        style_section = f"""

User's Description of Recommender's Style:
"{style_description[:1000]}"

IMPORTANT: Use this description to inform the persona, but ALWAYS maintain strict professional language standards.
The final writing style must be:
- Suitable for formal immigration petition letters
- Professional and business-appropriate
- Free of slang, colloquialisms, or inappropriate language
- Respectful and dignified in tone
Adapt the user's observations into a professional writing persona."""

    analysis_prompt = f"""Analyze this professional for writing style recommendations:

Name: {name}
Title: {title}
Organization: {organization}
{cv_section}
{style_section}

Based on their professional background and any style descriptions provided, determine:
1. Likely communication style (formal/academic/corporate/technical)
2. Writing tone (authoritative/collaborative/mentoring/peer)
3. Key areas of expertise they would emphasize
4. Relationship context (academic supervisor, industry leader, colleague, etc.)
5. Recommended font style (formal like Times New Roman, modern like Calibri, academic like Garamond)
6. Appropriate salutation and closing for letters they would write

Respond in JSON format only:
{{
    "style": "formal|academic|corporate|technical|casual",
    "tone": "authoritative|collaborative|mentoring|peer|diplomatic",
    "expertise_areas": ["area1", "area2", "area3"],
    "relationship_type": "supervisor|colleague|industry_leader|academic|mentor",
    "font": "Times New Roman|Calibri|Arial|Garamond|Georgia",
    "font_size": 11 or 12,
    "use_tabs": true or false,
    "salutation": "appropriate opening",
    "closing": "appropriate closing",
    "emphasis": ["what they emphasize", "in letters"],
    "profile_summary": "One sentence describing their writing persona"
}}"""

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                "https://api.perplexity.ai/chat/completions",
                headers={
                    "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "sonar",
                    "messages": [
                        {"role": "system", "content": "You are a professional writing style analyst. Respond only in valid JSON format."},
                        {"role": "user", "content": analysis_prompt}
                    ]
                }
            )

            if response.status_code == 200:
                data = response.json()
                content = data.get('choices', [{}])[0].get('message', {}).get('content', '')

                # Parse JSON from response
                import json
                import re

                # Extract JSON from response (may have markdown or extra text)
                json_match = re.search(r'\{[\s\S]*\}', content)
                if json_match:
                    profile = json.loads(json_match.group())
                    return {
                        "success": True,
                        "profile": profile
                    }
                else:
                    # Return default profile
                    return {
                        "success": True,
                        "profile": {
                            "style": "formal",
                            "tone": "authoritative",
                            "expertise_areas": ["professional expertise"],
                            "relationship_type": "colleague",
                            "font": "Arial",
                            "font_size": 11,
                            "use_tabs": False,
                            "salutation": "Dear Reviewing Officer,",
                            "closing": "Sincerely,",
                            "emphasis": ["professional achievements", "industry impact"],
                            "profile_summary": f"Professional profile for {name}"
                        }
                    }
            else:
                logger.error(f"Profile analysis request failed: {response.status_code}")
                return {"success": False, "error": "Analysis service unavailable"}

    except Exception as e:
        logger.error(f"Profile analysis error: {str(e)}")
        return {"success": False, "error": str(e)}


@app.post("/api/persona/analyze")
async def analyze_persona(
    name: str = Form(...),
    title: str = Form(""),
    organization: str = Form(""),
    style_description: str = Form(""),
    cv_file: Optional[UploadFile] = File(None)
):
    """Analyze a recommender and generate a custom writing profile."""
    logger.info(f"Analyzing profile for: {name}, {title}, {organization}, style: {style_description[:100] if style_description else 'None'}")

    cv_text = ""
    temp_path = None

    try:
        # Extract CV text if provided
        if cv_file and cv_file.filename:
            ext = Path(cv_file.filename).suffix.lower()
            if ext in ['.docx', '.pdf', '.txt']:
                temp_path = UPLOAD_DIR / f"cv_{uuid.uuid4()}{ext}"
                with open(temp_path, "wb") as f:
                    shutil.copyfileobj(cv_file.file, f)
                cv_text = extract_cv_text(str(temp_path))
                logger.info(f"Extracted {len(cv_text)} chars from CV")

        # Analyze profile
        result = await analyze_recommender_profile(name, title, organization, cv_text, style_description)

        # Cleanup
        if temp_path and temp_path.exists():
            os.remove(temp_path)

        if result.get("success"):
            return JSONResponse({
                "success": True,
                "persona": {
                    "name": f"Custom: {name.split()[0] if name else 'Profile'}",
                    "recommender_name": name,
                    "recommender_title": title,
                    "recommender_org": organization,
                    **result["profile"]
                }
            })
        else:
            return JSONResponse({
                "success": False,
                "error": result.get("error", "Analysis failed")
            }, status_code=500)

    except Exception as e:
        logger.error(f"Persona analysis failed: {str(e)}", exc_info=True)
        if temp_path and temp_path.exists():
            os.remove(temp_path)
        return JSONResponse({
            "success": False,
            "error": str(e)
        }, status_code=500)



# Link download endpoint



# =============================================================================
# DOWNLOAD ENDPOINT
# =============================================================================

@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download a generated file. SECURED against path traversal."""
    # Security: Prevent path traversal attacks
    if ".." in filename or "/" in filename or "\\" in filename:
        logger.warning(f"Path traversal attempt blocked: {filename}")
        raise HTTPException(status_code=400, detail="Invalid filename")

    # Only allow alphanumeric, dash, underscore, and dot
    import re
    if not re.match(r'^[\w\-\.]+$', filename):
        logger.warning(f"Invalid filename characters blocked: {filename}")
        raise HTTPException(status_code=400, detail="Invalid filename characters")

    filepath = OUTPUT_DIR / filename

    # Double-check: ensure resolved path is within OUTPUT_DIR
    try:
        resolved = filepath.resolve()
        if not resolved.is_relative_to(OUTPUT_DIR.resolve()):
            logger.warning(f"Path escape attempt blocked: {filename} -> {resolved}")
            raise HTTPException(status_code=403, detail="Access denied")
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid path")

    if not filepath.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type="application/octet-stream",
    )


# =============================================================================
# FEEDBACK SYSTEM
# =============================================================================

class FeedbackSubmit(BaseModel):
    message: str
    page: str
    user_email: Optional[str] = None

@app.post("/api/feedback")
async def submit_feedback(data: FeedbackSubmit, request: Request):
    """Receive feedback and send to Google Chat + Email."""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    user_agent = request.headers.get("user-agent", "unknown")
    ip = request.client.host if request.client else "unknown"

    logger.info(f"Feedback received from {ip}: {data.message[:50]}...")

    # 1. Send to Google Chat via Webhook
    webhook_url = os.getenv("GOOGLE_CHAT_WEBHOOK_URL")
    if webhook_url:
        try:
            chat_message = {
                "text": f"🐛 *Novo Feedback - CaseHub Tools*\n\n"
                        f"*Página:* {data.page}\n"
                        f"*Usuário:* {data.user_email or 'Anônimo'}\n"
                        f"*IP:* {ip}\n\n"
                        f"*Mensagem:*\n{data.message}"
            }
            async with httpx.AsyncClient() as client:
                await client.post(webhook_url, json=chat_message, timeout=10.0)
            logger.info("Feedback sent to Google Chat")
        except Exception as e:
            logger.error(f"Failed to send to Google Chat: {e}")

    # 2. Send email notification
    feedback_email = os.getenv("FEEDBACK_EMAIL")
    smtp_host = os.getenv("SMTP_HOST")
    smtp_port = os.getenv("SMTP_PORT", "587")
    smtp_user = os.getenv("SMTP_USER")
    smtp_password = os.getenv("SMTP_PASSWORD")

    if feedback_email and smtp_host and smtp_user and smtp_password:
        try:
            msg = MIMEMultipart()
            msg['From'] = smtp_user
            msg['To'] = feedback_email
            msg['Subject'] = f"[CaseHub Tools Feedback] {data.page}"

            body = f"""Novo Feedback Recebido - CaseHub Tools

Página: {data.page}
Usuário: {data.user_email or 'Anônimo'}
IP: {ip}
Data: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Mensagem:
{data.message}

---
Este email foi enviado automaticamente pelo sistema CaseHub Tools.
"""
            msg.attach(MIMEText(body, 'plain'))

            with smtplib.SMTP(smtp_host, int(smtp_port)) as server:
                server.starttls()
                server.login(smtp_user, smtp_password)
                server.send_message(msg)

            logger.info(f"Feedback email sent to {feedback_email}")
        except Exception as e:
            logger.error(f"Failed to send feedback email: {e}")
    else:
        logger.warning("Email not configured - feedback not sent via email")

    # 3. Log feedback to file for backup
    feedback_log = LOG_DIR / "feedback.log"
    with open(feedback_log, "a") as f:
        f.write(f"\n{'='*60}\n")
        f.write(f"Date: {datetime.now().isoformat()}\n")
        f.write(f"Page: {data.page}\n")
        f.write(f"User: {data.user_email or 'Anonymous'}\n")
        f.write(f"IP: {ip}\n")
        f.write(f"Message:\n{data.message}\n")

    return {"success": True, "message": "Feedback enviado! Obrigado."}


# =============================================================================
# TASK NOTIFICATION EMAILS
# =============================================================================

async def send_task_assignment_email(to_email: str, person_name: str, task_title: str, task_id: str):
    """Send email notification when a person is assigned to a task."""
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    smtp_host = os.getenv("SMTP_HOST")
    smtp_port = os.getenv("SMTP_PORT", "587")
    smtp_user = os.getenv("SMTP_USER")
    smtp_password = os.getenv("SMTP_PASSWORD")
    from_email = os.getenv("EMAIL_FROM", smtp_user)

    if not all([smtp_host, smtp_user, smtp_password]):
        logger.warning("SMTP not configured, skipping email notification")
        return

    msg = MIMEMultipart("alternative")
    msg["Subject"] = f"Nova tarefa atribuída: {task_title}"
    msg["From"] = f"CaseHub <{from_email}>"
    msg["To"] = to_email

    task_url = f"https://casehub.app/tools/dashboard"

    text_content = f"""
Olá {person_name},

Você foi atribuído(a) a uma nova tarefa:

📋 {task_title}

Acesse o CaseHub para ver os detalhes:
{task_url}

---
CaseHub
casehub.app
"""

    html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; color: #333; }}
        .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 10px 10px 0 0; }}
        .content {{ background: #f8f9fa; padding: 30px; border-radius: 0 0 10px 10px; }}
        .task-box {{ background: white; border-left: 4px solid #667eea; padding: 15px; margin: 20px 0; border-radius: 4px; }}
        .btn {{ display: inline-block; background: #667eea; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; margin-top: 20px; }}
        .footer {{ text-align: center; margin-top: 30px; color: #666; font-size: 12px; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1 style="margin: 0;">📋 Nova Tarefa</h1>
        </div>
        <div class="content">
            <p>Olá <strong>{person_name}</strong>,</p>
            <p>Você foi atribuído(a) a uma nova tarefa:</p>
            <div class="task-box">
                <h3 style="margin: 0; color: #667eea;">{task_title}</h3>
            </div>
            <a href="{task_url}" class="btn">Ver no CaseHub</a>
        </div>
        <div class="footer">
            <p>CaseHub • casehub.app</p>
        </div>
    </div>
</body>
</html>
"""

    msg.attach(MIMEText(text_content, "plain"))
    msg.attach(MIMEText(html_content, "html"))

    # Send email (synchronous, run in thread pool)
    import asyncio
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, lambda: _send_email_sync(smtp_host, smtp_port, smtp_user, smtp_password, msg))


def _send_email_sync(smtp_host, smtp_port, smtp_user, smtp_password, msg):
    """Synchronous email sending helper."""
    import smtplib
    with smtplib.SMTP(smtp_host, int(smtp_port)) as server:
        server.starttls()
        server.login(smtp_user, smtp_password)
        server.send_message(msg)


async def send_google_chat_notification(task_title: str, assignee_names: list, task_id: str):
    """Send a card notification to Google Chat space(s) via webhook."""
    import httpx

    webhook_urls_str = os.getenv("GOOGLE_CHAT_WEBHOOK_URL")
    if not webhook_urls_str:
        logger.warning("GOOGLE_CHAT_WEBHOOK_URL not configured, skipping Chat notification")
        return False

    # Support multiple webhooks (comma-separated)
    webhook_urls = [url.strip() for url in webhook_urls_str.split(",") if url.strip()]

    # Format assignees
    assignees_text = ", ".join(assignee_names) if assignee_names else "Não atribuído"

    # Build a beautiful Card v2
    card_message = {
        "cardsV2": [{
            "cardId": f"task-{task_id[:8]}",
            "card": {
                "header": {
                    "title": "📋 Nova Atribuição de Tarefa",
                    "subtitle": "CaseHub",
                    "imageUrl": "https://casehub.app/favicon.ico",
                    "imageType": "CIRCLE"
                },
                "sections": [
                    {
                        "header": "Detalhes da Tarefa",
                        "collapsible": False,
                        "widgets": [
                            {
                                "decoratedText": {
                                    "startIcon": {"knownIcon": "DESCRIPTION"},
                                    "topLabel": "Tarefa",
                                    "text": f"<b>{task_title}</b>"
                                }
                            },
                            {
                                "decoratedText": {
                                    "startIcon": {"knownIcon": "PERSON"},
                                    "topLabel": "Responsável(is)",
                                    "text": assignees_text
                                }
                            }
                        ]
                    },
                    {
                        "widgets": [
                            {
                                "buttonList": {
                                    "buttons": [
                                        {
                                            "text": "Ver no CaseHub",
                                            "onClick": {
                                                "openLink": {
                                                    "url": "https://casehub.app/tools/dashboard"
                                                }
                                            }
                                        }
                                    ]
                                }
                            }
                        ]
                    }
                ]
            }
        }]
    }

    success_count = 0
    try:
        async with httpx.AsyncClient() as client:
            for webhook_url in webhook_urls:
                try:
                    response = await client.post(
                        webhook_url,
                        json=card_message,
                        headers={"Content-Type": "application/json"}
                    )

                    if response.status_code == 200:
                        logger.info(f"Google Chat notification sent to {webhook_url[:50]}... for task: {task_title}")
                        success_count += 1
                    else:
                        logger.error(f"Google Chat webhook error: {response.status_code} - {response.text}")
                except Exception as e:
                    logger.error(f"Failed to send to webhook {webhook_url[:50]}...: {e}")

            if success_count > 0:
                logger.info(f"Sent to {success_count}/{len(webhook_urls)} Google Chat spaces")
                return True
            else:
                logger.error(f"Failed to send to any Google Chat space")
                return False

    except Exception as e:
        logger.error(f"Failed to send Google Chat notification: {e}")
        return False


# =============================================================================
# NOTION TASKS INTEGRATION
# =============================================================================

class TaskCreate(BaseModel):
    title: str
    assignee: str
    due_date: Optional[str] = None
    description: Optional[str] = None
    origem: Optional[str] = "Manual"

@app.get("/api/tasks")
async def list_tasks(
    assignee: Optional[str] = None,
    status: Optional[str] = None,
    origem: Optional[str] = None,
    user: dict = Depends(require_auth)
):
    """List tasks from Notion database."""
    notion_key = os.getenv("NOTION_API_KEY")
    database_id = os.getenv("NOTION_DATABASE_ID")

    if not notion_key or not database_id:
        raise HTTPException(status_code=503, detail="Notion integration not configured")

    try:
        import httpx

        headers = {
            "Authorization": f"Bearer {notion_key}",
            "Content-Type": "application/json",
            "Notion-Version": "2022-06-28"
        }

        # Build request body
        body = {}
        filters = []

        if assignee:
            filters.append({
                "property": "Responsável",
                "select": {"equals": assignee}
            })
        if status:
            filters.append({
                "property": "Status",
                "status": {"equals": status}
            })
        if origem:
            filters.append({
                "property": "Origem",
                "select": {"equals": origem}
            })

        if filters:
            body["filter"] = {"and": filters} if len(filters) > 1 else filters[0]

        # Query Notion API directly
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"https://api.notion.com/v1/databases/{database_id}/query",
                headers=headers,
                json=body
            )

        if response.status_code != 200:
            logger.error(f"Notion API error: {response.status_code} - {response.text}")
            raise HTTPException(status_code=response.status_code, detail="Notion API error")

        data = response.json()

        # Parse results - adapting to actual Notion database structure
        # Schema: Tarefa(title), Pessoa(people), Responsável(relation), Caso(relation),
        #         Cliente(relation), Status(select), Deadline(date), Prioridade(select)
        tasks = []
        for page in data.get("results", []):
            props = page.get("properties", {})

            # Extract title from "Tarefa" column (title type)
            tarefa = ""
            tarefa_prop = props.get("Tarefa", {})
            if tarefa_prop.get("title"):
                title_list = tarefa_prop["title"]
                tarefa = title_list[0].get("text", {}).get("content", "") if title_list else ""

            # Extract "Pessoa" (people type - assigned team member)
            pessoa = {"name": "", "ids": []}
            pessoa_prop = props.get("Pessoa", {})
            if pessoa_prop.get("people"):
                people_list = pessoa_prop["people"]
                if people_list:
                    pessoa["name"] = people_list[0].get("name", "")
                    pessoa["ids"] = [p.get("id") for p in people_list]

            # Extract "Responsável" (relation type - links to team database)
            responsavel = {"name": "", "ids": []}
            resp_prop = props.get("Responsável", {})
            if resp_prop.get("relation"):
                rel_list = resp_prop["relation"]
                responsavel["ids"] = [r.get("id") for r in rel_list]

            # Extract "Caso" (relation type)
            caso = {"name": "", "ids": []}
            caso_prop = props.get("Caso", {})
            if caso_prop.get("relation"):
                rel_list = caso_prop["relation"]
                caso["ids"] = [r.get("id") for r in rel_list]

            # Extract "Cliente" (relation type)
            cliente = {"name": "", "ids": []}
            cliente_prop = props.get("Cliente", {})
            if cliente_prop.get("relation"):
                rel_list = cliente_prop["relation"]
                cliente["ids"] = [r.get("id") for r in rel_list]

            # Extract status (select type)
            status_val = ""
            status_prop = props.get("Status", {})
            if status_prop.get("select"):
                status_val = status_prop["select"].get("name", "")

            # Extract deadline (date type)
            deadline = ""
            deadline_prop = props.get("Deadline", {})
            if deadline_prop.get("date"):
                deadline = deadline_prop["date"].get("start", "")

            # Extract priority (select type)
            prioridade = ""
            prio_prop = props.get("Prioridade", {})
            if prio_prop.get("select"):
                prioridade = prio_prop["select"].get("name", "")

            # Extract origem (select type)
            origem_val = ""
            origem_prop = props.get("Origem", {})
            if origem_prop.get("select"):
                origem_val = origem_prop["select"].get("name", "")

            tasks.append({
                "id": page["id"],
                "tarefa": tarefa,
                "pessoa": pessoa,
                "responsavel": responsavel,
                "caso": caso,
                "cliente": cliente,
                "status": status_val,
                "deadline": deadline,
                "prioridade": prioridade,
                "origem": origem_val,
                "url": page.get("url", "")
            })

        return {"tasks": tasks, "count": len(tasks)}

    except httpx.HTTPError as e:
        logger.error(f"HTTP error querying Notion: {e}")
        raise HTTPException(status_code=503, detail="Failed to connect to Notion")
    except Exception as e:
        logger.error(f"Notion API error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/tasks")
async def create_task(data: TaskCreate, user: dict = Depends(require_auth)):
    """Create a new task in Notion database."""
    notion_key = os.getenv("NOTION_API_KEY")
    database_id = os.getenv("NOTION_DATABASE_ID")

    if not notion_key or not database_id:
        raise HTTPException(status_code=503, detail="Notion integration not configured")

    try:
        from notion_client import Client
        notion = Client(auth=notion_key)

        # Schema: Tarefa(title), Status(select), Prioridade(select), Deadline(date), Origem(select)
        # Relations: Pessoa(people), Responsável(relation), Caso(relation), Cliente(relation)
        properties = {
            "Tarefa": {"title": [{"text": {"content": data.title}}]},
            "Status": {"select": {"name": "To Do"}},
            "Origem": {"select": {"name": data.origem or "Manual"}}
        }

        if data.due_date:
            properties["Deadline"] = {"date": {"start": data.due_date}}

        new_page = notion.pages.create(
            parent={"database_id": database_id},
            properties=properties
        )

        logger.info(f"Task created: {data.title}")

        return {
            "success": True,
            "task_id": new_page["id"],
            "url": new_page.get("url", "")
        }

    except ImportError:
        raise HTTPException(status_code=503, detail="notion-client not installed")
    except Exception as e:
        logger.error(f"Notion API error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/tasks/assignees")
async def list_assignees(user: dict = Depends(require_auth)):
    """Get list of possible assignees from users."""
    users = auth.load_users()
    assignees = [{"email": u["email"], "name": u["full_name"]} for u in users.values() if u["status"] == "active"]
    return {"assignees": assignees}


# Schema cache for Notion database
_notion_schema_cache = {"data": None, "timestamp": 0}
SCHEMA_CACHE_TTL = 300  # 5 minutes

@app.get("/api/notion/schema")
async def get_notion_schema(user: dict = Depends(require_auth)):
    """Get Notion database schema with select options and relation items."""
    import time
    global _notion_schema_cache

    # Check cache
    if _notion_schema_cache["data"] and (time.time() - _notion_schema_cache["timestamp"]) < SCHEMA_CACHE_TTL:
        return _notion_schema_cache["data"]

    notion_key = os.getenv("NOTION_API_KEY")
    database_id = os.getenv("NOTION_DATABASE_ID")

    if not notion_key or not database_id:
        raise HTTPException(status_code=503, detail="Notion integration not configured")

    try:
        import httpx

        headers = {
            "Authorization": f"Bearer {notion_key}",
            "Content-Type": "application/json",
            "Notion-Version": "2022-06-28"
        }

        async with httpx.AsyncClient() as client:
            # Get database schema
            db_response = await client.get(
                f"https://api.notion.com/v1/databases/{database_id}",
                headers=headers
            )

            if db_response.status_code != 200:
                raise HTTPException(status_code=db_response.status_code, detail="Failed to get database schema")

            db_data = db_response.json()
            properties = db_data.get("properties", {})

            schema = {
                "selects": {},
                "relations": {},
                "properties": {}
            }

            # Extract select options and relation database IDs
            for prop_name, prop_data in properties.items():
                prop_type = prop_data.get("type")
                schema["properties"][prop_name] = prop_type

                if prop_type == "select":
                    options = prop_data.get("select", {}).get("options", [])
                    schema["selects"][prop_name] = [
                        {"id": opt.get("id"), "name": opt.get("name"), "color": opt.get("color")}
                        for opt in options
                    ]

                elif prop_type == "status":
                    options = prop_data.get("status", {}).get("options", [])
                    schema["selects"][prop_name] = [
                        {"id": opt.get("id"), "name": opt.get("name"), "color": opt.get("color")}
                        for opt in options
                    ]

                elif prop_type == "relation":
                    relation_db_id = prop_data.get("relation", {}).get("database_id")
                    if relation_db_id:
                        # Query related database to get items
                        rel_response = await client.post(
                            f"https://api.notion.com/v1/databases/{relation_db_id}/query",
                            headers=headers,
                            json={"page_size": 100}
                        )

                        if rel_response.status_code == 200:
                            rel_data = rel_response.json()
                            items = []
                            for page in rel_data.get("results", []):
                                # Try to get title from properties
                                page_props = page.get("properties", {})
                                title = ""
                                for p_name, p_val in page_props.items():
                                    if p_val.get("type") == "title":
                                        title_arr = p_val.get("title", [])
                                        if title_arr:
                                            title = title_arr[0].get("text", {}).get("content", "")
                                        break

                                items.append({
                                    "id": page["id"],
                                    "name": title or page["id"][:8]
                                })

                            schema["relations"][prop_name] = {
                                "database_id": relation_db_id,
                                "items": items
                            }

            # Get workspace users (people)
            try:
                users_response = await client.get(
                    "https://api.notion.com/v1/users",
                    headers=headers
                )
                if users_response.status_code == 200:
                    users_data = users_response.json()
                    schema["people"] = [
                        {
                            "id": u["id"],
                            "name": u.get("name", ""),
                            "email": u.get("person", {}).get("email", ""),
                            "avatar": u.get("avatar_url", "")
                        }
                        for u in users_data.get("results", [])
                        if u.get("type") == "person"
                    ]
                else:
                    schema["people"] = []
            except Exception as e:
                logger.warning(f"Could not fetch Notion users: {e}")
                schema["people"] = []

            # Cache the result
            _notion_schema_cache = {"data": schema, "timestamp": time.time()}

            return schema

    except httpx.HTTPError as e:
        logger.error(f"HTTP error getting Notion schema: {e}")
        raise HTTPException(status_code=503, detail="Failed to connect to Notion")
    except Exception as e:
        logger.error(f"Error getting Notion schema: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class TaskUpdate(BaseModel):
    property: str
    value: Optional[str] = None
    relation_ids: Optional[List[str]] = None
    send_email: Optional[bool] = False  # Send email notifications to new assignees


@app.patch("/api/tasks/{task_id}")
async def update_task(task_id: str, data: TaskUpdate, user: dict = Depends(require_auth)):
    """Update a specific property of a task."""
    notion_key = os.getenv("NOTION_API_KEY")

    if not notion_key:
        raise HTTPException(status_code=503, detail="Notion integration not configured")

    try:
        import httpx

        headers = {
            "Authorization": f"Bearer {notion_key}",
            "Content-Type": "application/json",
            "Notion-Version": "2022-06-28"
        }

        # Build properties update based on property type
        properties = {}
        prop_name = data.property

        # Map property names to Notion format
        if prop_name == "tarefa" or prop_name == "Tarefa":
            properties["Tarefa"] = {
                "title": [{"text": {"content": data.value or ""}}]
            }

        elif prop_name == "status" or prop_name == "Status":
            properties["Status"] = {
                "select": {"name": data.value} if data.value else None
            }

        elif prop_name == "prioridade" or prop_name == "Prioridade":
            properties["Prioridade"] = {
                "select": {"name": data.value} if data.value else None
            }

        elif prop_name == "deadline" or prop_name == "Deadline":
            properties["Deadline"] = {
                "date": {"start": data.value} if data.value else None
            }

        elif prop_name == "pessoa" or prop_name == "Pessoa":
            # People type - needs user IDs
            if data.relation_ids:
                properties["Pessoa"] = {
                    "people": [{"id": uid} for uid in data.relation_ids]
                }
            else:
                properties["Pessoa"] = {"people": []}

        elif prop_name in ["responsavel", "Responsável"]:
            # Responsável is a relation type in Notion (links to team database)
            if data.relation_ids:
                properties["Responsável"] = {
                    "relation": [{"id": rid} for rid in data.relation_ids]
                }
            else:
                properties["Responsável"] = {"relation": []}

        elif prop_name in ["caso", "Caso", "cliente", "Cliente"]:
            # Relation types
            notion_prop_name = {
                "caso": "Caso",
                "Caso": "Caso",
                "cliente": "Cliente",
                "Cliente": "Cliente"
            }.get(prop_name, prop_name)

            if data.relation_ids:
                properties[notion_prop_name] = {
                    "relation": [{"id": rid} for rid in data.relation_ids]
                }
            else:
                properties[notion_prop_name] = {"relation": []}

        else:
            raise HTTPException(status_code=400, detail=f"Unknown property: {prop_name}")

        async with httpx.AsyncClient() as client:
            response = await client.patch(
                f"https://api.notion.com/v1/pages/{task_id}",
                headers=headers,
                json={"properties": properties}
            )

            if response.status_code != 200:
                logger.error(f"Notion API error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=response.status_code, detail="Failed to update task")

            logger.info(f"Task {task_id} updated: {prop_name} = {data.value or data.relation_ids}")

            # Send notifications to new assignees
            emails_sent = []
            chat_notified = False
            assignee_names = []

            if data.send_email and prop_name in ["responsavel", "Responsável"] and data.relation_ids:
                # Get task title for notifications
                task_response = await client.get(
                    f"https://api.notion.com/v1/pages/{task_id}",
                    headers=headers
                )
                task_title = "Nova tarefa"
                if task_response.status_code == 200:
                    task_data = task_response.json()
                    title_prop = task_data.get("properties", {}).get("Tarefa", {})
                    if title_prop.get("title"):
                        task_title = title_prop["title"][0].get("text", {}).get("content", "Nova tarefa")

                # Get user info from cache
                if _notion_schema_cache.get("data") and _notion_schema_cache["data"].get("people"):
                    people = _notion_schema_cache["data"]["people"]
                    for user_id in data.relation_ids:
                        person = next((p for p in people if p["id"] == user_id), None)
                        if person:
                            assignee_names.append(person["name"])
                            # Send email notification
                            if person.get("email"):
                                try:
                                    await send_task_assignment_email(
                                        to_email=person["email"],
                                        person_name=person["name"],
                                        task_title=task_title,
                                        task_id=task_id
                                    )
                                    emails_sent.append(person["email"])
                                    logger.info(f"Assignment email sent to {person['email']}")
                                except Exception as email_error:
                                    logger.error(f"Failed to send email to {person['email']}: {email_error}")

                # Send Google Chat notification (one message for all assignees)
                if assignee_names:
                    try:
                        chat_notified = await send_google_chat_notification(
                            task_title=task_title,
                            assignee_names=assignee_names,
                            task_id=task_id
                        )
                    except Exception as chat_error:
                        logger.error(f"Failed to send Google Chat notification: {chat_error}")

            return {"success": True, "task_id": task_id, "emails_sent": emails_sent, "chat_notified": chat_notified}

    except httpx.HTTPError as e:
        logger.error(f"HTTP error updating task: {e}")
        raise HTTPException(status_code=503, detail="Failed to connect to Notion")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating task: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/tasks/{task_id}")
async def delete_task(task_id: str, user: dict = Depends(require_auth)):
    """Delete (archive) a task."""
    notion_key = os.getenv("NOTION_API_KEY")

    if not notion_key:
        raise HTTPException(status_code=503, detail="Notion integration not configured")

    try:
        import httpx

        headers = {
            "Authorization": f"Bearer {notion_key}",
            "Content-Type": "application/json",
            "Notion-Version": "2022-06-28"
        }

        async with httpx.AsyncClient() as client:
            # Archive the page (Notion doesn't truly delete)
            response = await client.patch(
                f"https://api.notion.com/v1/pages/{task_id}",
                headers=headers,
                json={"archived": True}
            )

            if response.status_code != 200:
                logger.error(f"Notion API error: {response.status_code} - {response.text}")
                raise HTTPException(status_code=response.status_code, detail="Failed to delete task")

            logger.info(f"Task {task_id} archived/deleted")

            return {"success": True, "task_id": task_id}

    except httpx.HTTPError as e:
        logger.error(f"HTTP error deleting task: {e}")
        raise HTTPException(status_code=503, detail="Failed to connect to Notion")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting task: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# LEADS CRM API
# =============================================================================

LEADS_WEBHOOK_API_KEY = os.getenv("LEADS_WEBHOOK_API_KEY", "ilc-leads-webhook-2026")


# --- Sync (must be before /{lead_id} routes to avoid path conflicts) ---

@app.post("/api/leads/sync/moskit")
async def sync_leads_from_moskit(admin: dict = Depends(require_admin)):
    """Import/sync leads from Moskit CRM."""
    try:
        data = leads_manager.load_leads()
        result = await leads_manager.sync_from_moskit(data)
        leads_manager.save_leads(data)
        logger.info(f"Moskit sync triggered by {admin['email']}: {result}")
        return {"success": True, **result}
    except Exception as e:
        logger.error(f"Error syncing from Moskit: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/leads/sync/notion")
async def sync_leads_to_notion(admin: dict = Depends(require_admin)):
    """Sync all leads to Notion database."""
    try:
        data = leads_manager.load_leads()
        result = await leads_manager.sync_all_to_notion(data)
        leads_manager.save_leads(data)
        logger.info(f"Notion sync triggered by {admin['email']}: {result}")
        return {"success": True, **result}
    except Exception as e:
        logger.error(f"Error syncing to Notion: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/leads/sync/status")
async def get_sync_status(admin: dict = Depends(require_admin)):
    """Get sync status information."""
    try:
        data = leads_manager.load_leads()
        recent_logs = data.get("sync_log", [])[-20:]
        return {
            "last_moskit_sync": data.get("last_moskit_sync"),
            "last_notion_sync": data.get("last_notion_sync"),
            "last_updated": data.get("last_updated"),
            "total_leads": len(data.get("leads", {})),
            "recent_sync_log": recent_logs
        }
    except Exception as e:
        logger.error(f"Error getting sync status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Metrics (must be before /{lead_id} routes to avoid path conflicts) ---

@app.get("/api/leads/metrics")
async def get_leads_metrics(admin: dict = Depends(require_admin)):
    """Get dashboard metrics for leads."""
    try:
        data = leads_manager.load_leads()
        metrics = leads_manager.get_metrics(data)
        return metrics
    except Exception as e:
        logger.error(f"Error getting leads metrics: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/leads/metrics/pipeline")
async def get_pipeline_metrics(admin: dict = Depends(require_admin)):
    """Get pipeline/funnel metrics."""
    try:
        data = leads_manager.load_leads()
        metrics = leads_manager.get_pipeline_metrics(data)
        return metrics
    except Exception as e:
        logger.error(f"Error getting pipeline metrics: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/leads/metrics/trends")
async def get_trend_metrics(
    days: int = 30,
    admin: dict = Depends(require_admin)
):
    """Get trend metrics (daily lead counts)."""
    try:
        data = leads_manager.load_leads()
        metrics = leads_manager.get_trend_metrics(data, days=days)
        return metrics
    except Exception as e:
        logger.error(f"Error getting trend metrics: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- Webhook (for WhatsApp bot dual-write, no JWT auth) ---

@app.post("/api/leads/webhook")
async def leads_webhook(request: Request):
    """
    Receive leads from WhatsApp bot (dual-write).
    Authenticated via X-API-Key header (not JWT).
    """
    api_key = request.headers.get("X-API-Key", "")
    if api_key != LEADS_WEBHOOK_API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")

    try:
        body = await request.json()
        data = leads_manager.load_leads()
        lead = leads_manager.upsert_from_webhook(data, body)
        leads_manager.save_leads(data)
        logger.info(f"Webhook: lead upserted - {lead['id']} ({lead.get('phone', 'no-phone')})")
        return {"success": True, "lead_id": lead["id"], "action": "upserted"}
    except Exception as e:
        logger.error(f"Error processing webhook: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# --- CRUD ---

@app.get("/api/leads")
async def list_leads(
    request: Request,
    stage: Optional[str] = None,
    source: Optional[str] = None,
    status: Optional[str] = None,
    lead_status: Optional[str] = None,
    score_min: Optional[int] = None,
    score_max: Optional[int] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    search: Optional[str] = None,
    sort_by: str = "created_at",
    sort_order: str = "desc",
    page: int = 1,
    per_page: int = 50,
    admin: dict = Depends(require_admin)
):
    """List leads with filters, search, sorting, and pagination."""
    try:
        data = leads_manager.load_leads()

        leads, total = leads_manager.get_all_leads(
            data,
            search=search,
            stage=stage,
            source=source,
            status=status,
            lead_status=lead_status,
            score_min=score_min,
            score_max=score_max,
            date_from=date_from,
            date_to=date_to,
            sort_by=sort_by,
            sort_order=sort_order,
            page=page,
            per_page=per_page
        )

        return {
            "leads": leads,
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": (total + per_page - 1) // per_page if per_page > 0 else 1
        }
    except Exception as e:
        logger.error(f"Error listing leads: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/leads")
async def create_lead(lead_data: dict, admin: dict = Depends(require_admin)):
    """Create a new lead."""
    try:
        data = leads_manager.load_leads()
        lead = leads_manager.create_lead(data, lead_data)
        leads_manager.save_leads(data)
        logger.info(f"Lead created by {admin['email']}: {lead['id']}")
        return {"success": True, "lead": lead}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error creating lead: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/leads/{lead_id}")
async def get_lead(lead_id: str, admin: dict = Depends(require_admin)):
    """Get a single lead by ID."""
    try:
        data = leads_manager.load_leads()
        lead = data["leads"].get(lead_id)
        if not lead:
            raise HTTPException(status_code=404, detail="Lead not found")
        return {"lead": lead}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting lead {lead_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/leads/{lead_id}")
async def update_lead(lead_id: str, updates: dict, admin: dict = Depends(require_admin)):
    """Update a lead."""
    try:
        data = leads_manager.load_leads()
        lead = leads_manager.update_lead(data, lead_id, updates)
        leads_manager.save_leads(data)
        logger.info(f"Lead {lead_id} updated by {admin['email']}")
        return {"success": True, "lead": lead}
    except KeyError:
        raise HTTPException(status_code=404, detail="Lead not found")
    except Exception as e:
        logger.error(f"Error updating lead {lead_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/leads/{lead_id}")
async def delete_lead(lead_id: str, admin: dict = Depends(require_admin)):
    """Soft-delete a lead."""
    try:
        data = leads_manager.load_leads()
        leads_manager.delete_lead(data, lead_id)
        leads_manager.save_leads(data)
        logger.info(f"Lead {lead_id} deleted by {admin['email']}")
        return {"success": True}
    except KeyError:
        raise HTTPException(status_code=404, detail="Lead not found")
    except Exception as e:
        logger.error(f"Error deleting lead {lead_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/leads/{lead_id}/stage")
async def update_lead_stage(lead_id: str, body: dict, admin: dict = Depends(require_admin)):
    """Move a lead to a different pipeline stage (Kanban drag-and-drop)."""
    try:
        stage = body.get("stage")
        if not stage:
            raise HTTPException(status_code=400, detail="Missing 'stage' field")

        valid_stages = list(leads_manager.MOSKIT_STAGES.keys())
        if stage not in valid_stages:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid stage. Must be one of: {valid_stages}"
            )

        data = leads_manager.load_leads()
        lead = leads_manager.update_lead(data, lead_id, {"pipeline_stage": stage})
        leads_manager.save_leads(data)
        logger.info(f"Lead {lead_id} moved to stage {stage} by {admin['email']}")
        return {"success": True, "lead": lead}
    except KeyError:
        raise HTTPException(status_code=404, detail="Lead not found")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating lead stage {lead_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/leads/{lead_id}/notes")
async def add_lead_note(lead_id: str, body: dict, admin: dict = Depends(require_admin)):
    """Add a note to a lead."""
    try:
        content = body.get("content")
        if not content:
            raise HTTPException(status_code=400, detail="Missing 'content' field")

        note_type = body.get("type", "note")
        actor = admin.get("name", admin.get("email", "staff"))

        data = leads_manager.load_leads()
        note = leads_manager.add_note(data, lead_id, content, note_type=note_type, actor=actor)
        leads_manager.save_leads(data)
        logger.info(f"Note added to lead {lead_id} by {admin['email']}")
        return {"success": True, "note": note}
    except KeyError:
        raise HTTPException(status_code=404, detail="Lead not found")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error adding note to lead {lead_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# TRANSCRIPTION SERVICE
# =============================================================================

# Lazy import to avoid loading Whisper model on startup
transcription_service = None

def get_transcription_svc():
    """Get transcription service with lazy loading."""
    global transcription_service
    if transcription_service is None:
        try:
            from transcription_service import get_transcription_service
            transcription_service = get_transcription_service()
        except Exception as e:
            logger.error(f"Failed to load transcription service: {e}")
            raise HTTPException(status_code=503, detail="Transcription service not available")
    return transcription_service


@app.post("/api/transcribe")
async def transcribe_audio(
    audio: UploadFile = File(...),
    language: str = Form(default="pt"),
    generate_srt: bool = Form(default=True),
    generate_minutes: bool = Form(default=True)
):
    """
    Transcribe audio file and optionally generate meeting minutes.

    - **audio**: Audio file (mp3, wav, ogg, m4a, etc.)
    - **language**: Language code ('pt', 'en', 'es', 'auto')
    - **generate_srt**: Generate SRT subtitles (default: true)
    - **generate_minutes**: Generate meeting minutes (default: true)

    Returns transcription, SRT, and meeting minutes.
    """
    # Validate file type
    allowed_types = [
        "audio/mpeg", "audio/mp3", "audio/wav", "audio/wave", "audio/x-wav",
        "audio/ogg", "audio/x-ogg", "audio/mp4", "audio/m4a", "audio/webm",
        "video/webm", "application/ogg"
    ]
    content_type = audio.content_type or ""
    filename = audio.filename or "audio"

    # Also check by extension
    ext = Path(filename).suffix.lower()
    allowed_extensions = [".mp3", ".wav", ".ogg", ".m4a", ".webm", ".mp4", ".opus"]

    if content_type not in allowed_types and ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported audio format. Allowed: {', '.join(allowed_extensions)}"
        )

    # Save uploaded file temporarily
    temp_path = Path(tempfile.gettempdir()) / f"upload_{uuid.uuid4()}{ext or '.audio'}"

    try:
        # Write uploaded file
        content = await audio.read()
        temp_path.write_bytes(content)

        logger.info(f"Processing transcription: {filename} ({len(content)} bytes)")

        # Get transcription service and process
        svc = get_transcription_svc()
        result = svc.process(
            str(temp_path),
            language=language,
            generate_srt=generate_srt,
            generate_minutes=generate_minutes
        )

        if not result["success"]:
            raise HTTPException(status_code=500, detail=result.get("error", "Transcription failed"))

        return {
            "success": True,
            "transcription": result["transcription"],
            "srt": result["srt"],
            "minutes": result["minutes"],
            "language": result["language"],
            "duration_seconds": result["duration_seconds"]
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Transcription error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        # Cleanup
        if temp_path.exists():
            temp_path.unlink()


@app.get("/api/transcribe/status")
async def transcription_status():
    """Check if transcription service is available."""
    try:
        svc = get_transcription_svc()
        return {
            "available": True,
            "model_size": svc.model_size,
            "whisper_loaded": svc._whisper_loaded,
            "llm_loaded": svc._llm_loaded
        }
    except Exception as e:
        return {
            "available": False,
            "error": str(e)
        }


# =============================================================================
# RUN
# =============================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
