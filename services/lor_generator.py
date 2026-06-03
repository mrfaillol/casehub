#!/usr/bin/env python3
"""
CaseHub Letter of Recommendation Generator
========================================

Generates professionally formatted LORs for EB-2 NIW, EB-1A, and O-1 petitions.
Based on proven templates used for clients like Prince, Loyo, Edwin, Sagar, etc.

PERSONAS (5 types with specific formatting):
- Executive: Arial 11pt, no tabs, direct/results-oriented
- Technical: Calibri 11pt, no tabs, understated/factual
- Academic: Garamond 12pt, WITH tabs, scholarly/methodological
- Mentor: Georgia 11pt, WITH tabs, personal/narrative
- Corporate: Times New Roman 12pt, no tabs, formal/structured

CRITICAL RULES:
- NO em-dash (-) ever, use commas or colons
- NO "I am writing to recommend..."
- NO empty superlatives (brilliant, extraordinary)
- Mandatory variation between letters in same case
- Include footnotes with official sources
- Include Prong 3 (Matter of Dhanasar) justification
"""

import os
import logging
import random
from datetime import datetime
from typing import Literal, Optional, List, Dict

logger = logging.getLogger(__name__)

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# SECTION TEMPLATES FOR CONTENT EXPANSION (Added 2026-01-04)
# =============================================================================
# These templates are used to generate rich content for each LOR section
# Target: 600-900 words per LOR (currently generating ~360, need 2x expansion)

SECTION_TEMPLATES = {
    "recommender_credentials": {
        "academic": [
            "My name is {recommender_name}, and I am {recommender_title} at {recommender_org}. I hold a {degree} from {university}, specializing in {specialty}. With over {years_experience} years of experience in this field, I have supervised numerous graduate students, published extensively in peer-reviewed journals, and served as a reviewer for leading publications in this domain. My expertise in {specific_area} qualifies me to assess professionals working at the intersection of {field} and practical application. I have been recognized with {recognition} for my contributions to this field.",
            "I am {recommender_name}, {recommender_title} at {recommender_org}. My academic credentials include a {degree} from {university} and {additional_credentials}. Throughout my {years_experience}-year career, I have contributed to the advancement of {field} through both research and practical application. I have authored {publications} peer-reviewed articles and served on editorial boards for leading journals. This background positions me to evaluate professionals like {beneficiary_name} with authority and objectivity.",
        ],
        "executive": [
            "I am {recommender_name}, {recommender_title} at {recommender_org}. In my {years_experience} years leading teams of {team_size} professionals across {scope}, I have developed a keen ability to identify talent that drives organizational success. My responsibilities include {responsibilities}, giving me direct insight into what distinguishes exceptional performers from their peers. I have successfully led initiatives generating {revenue_impact} in value and have been recognized for {achievements}.",
            "My name is {recommender_name}. As {recommender_title} at {recommender_org}, I oversee {department_description}. With {years_experience} years in senior leadership positions at companies including {previous_companies}, I have developed expertise in identifying and developing talent. I have directly managed professionals who have gone on to leadership roles at major organizations, which qualifies me to assess {beneficiary_name}'s capabilities and potential.",
        ],
        "technical": [
            "I am {recommender_name}, {recommender_title} at {recommender_org}, where I lead {technical_responsibilities}. I hold {technical_credentials} and have {years_experience} years of hands-on experience in {technical_field}. My work has resulted in {patents_or_innovations} and has been implemented in systems serving {scale_of_impact}. This technical background allows me to objectively evaluate the contributions of professionals like {beneficiary_name}.",
            "My name is {recommender_name}. I serve as {recommender_title} at {recommender_org}, with responsibility for {scope_of_work}. My technical expertise spans {technical_areas}, developed through {years_experience} years of practical experience. I have contributed to {technical_achievements} and hold certifications in {certifications}. This positions me to assess technical contributions with precision and objectivity.",
        ],
        "mentor": [
            "I am {recommender_name}, {recommender_title} at {recommender_org}. Throughout my {years_experience}-year career in {field}, I have had the privilege of mentoring many emerging professionals, some of whom have gone on to leadership positions at organizations including {notable_mentees_orgs}. My approach to mentorship emphasizes both technical excellence and professional development. It is from this perspective that I write in support of {beneficiary_name}.",
        ],
        "corporate": [
            "My name is {recommender_name}, and I serve as {recommender_title} at {recommender_org}. Our organization {org_description}. In my role, I am responsible for {responsibilities}, which has given me extensive experience evaluating professional contributions and their impact on organizational objectives. With {years_experience} years in this position, I am well-qualified to assess the work of professionals in {field}.",
        ],
    },
    "relationship_context": {
        "supervisor": [
            "I have had the privilege of serving as {beneficiary_name}'s direct supervisor for {duration}. During this time, I have directly observed {his_her} work on {projects_description}, where {he_she} demonstrated exceptional capabilities in {specific_skills}. Our working relationship has given me comprehensive insight into {beneficiary_name}'s technical abilities, professional conduct, and potential for continued impact.",
            "As {beneficiary_name}'s supervisor since {start_date}, I have had extensive opportunity to observe {his_her} contributions to our organization. I have reviewed {his_her} work product, participated in {his_her} performance evaluations, and witnessed firsthand how {he_she} approaches complex challenges. This close professional relationship qualifies me to speak authoritatively about {beneficiary_name}'s capabilities.",
        ],
        "mentor": [
            "I first met {beneficiary_name} in {year} when {meeting_context}. Since then, I have served as {his_her} mentor, guiding {his_her} professional development through {mentorship_activities}. Over our {duration} of working together, I have witnessed remarkable growth and consistently impressive performance. Our mentoring relationship has given me unique insight into {beneficiary_name}'s character, work ethic, and professional capabilities.",
        ],
        "colleague": [
            "I have worked alongside {beneficiary_name} at {organization} for {duration}, collaborating on {projects_description}. Our peer relationship has allowed me to observe {his_her} work closely while also benefiting from {his_her} expertise. Through this collaboration, I have developed a clear understanding of {beneficiary_name}'s technical capabilities and professional qualities.",
        ],
        "independent": [
            "Although I have not worked directly with {beneficiary_name}, I have become familiar with {his_her} work through {how_familiar}. This independent perspective allows me to offer an objective assessment based purely on the quality and impact of {his_her} contributions to our field. I believe this objectivity adds credibility to my evaluation.",
            "I came to know of {beneficiary_name}'s work through {context}. As an independent evaluator with no personal or professional relationship, I can assess {his_her} contributions objectively. My familiarity with {his_her} work stems from {specific_exposure}, which has given me sufficient insight to comment on both the quality and significance of {his_her} contributions.",
        ],
    },
    "work_description": {
        "technical": [
            "{beneficiary_name} has made significant technical contributions in {technical_area}. Specifically, {he_she} developed {technical_achievement_1}, which {impact_1}. Additionally, {his_her} work on {technical_achievement_2} resulted in {impact_2}. These contributions demonstrate not only technical proficiency but also the ability to translate complex concepts into practical solutions. The systems {he_she} designed now serve {scale} and have {measurable_outcome}.",
            "During {his_her} time at {organization}, {beneficiary_name} has been responsible for {responsibilities}. {He_She} successfully completed {project_1}, which {outcome_1}. {His_Her} approach to {challenge} demonstrated exceptional problem-solving ability. Furthermore, {his_her} work on {project_2} resulted in {outcome_2}, directly benefiting {stakeholders}.",
        ],
        "research": [
            "{beneficiary_name}'s research contributions have advanced the field of {field} significantly. {His_Her} work on {research_topic} has been published in {publications} and cited {citation_count} times. The methodology {he_she} developed for {methodology_description} has been adopted by researchers at {institutions}. This represents a meaningful contribution to scientific knowledge with practical applications in {applications}.",
        ],
        "leadership": [
            "Beyond technical expertise, {beneficiary_name} has demonstrated exceptional leadership capabilities. {He_She} led a team of {team_size} professionals on {project}, delivering results that {outcomes}. Under {his_her} leadership, the team achieved {achievements}, including {specific_achievement}. {His_Her} ability to mentor junior team members while driving strategic initiatives makes {him_her} an asset to any organization.",
        ],
    },
    "national_importance": {
        "statistics": [
            "The national importance of {beneficiary_name}'s endeavor is substantial. {national_statistic}. This challenge affects {affected_population} Americans and costs the nation approximately {cost_estimate} annually. {beneficiary_name}'s work directly addresses this challenge by {how_addresses}. Without professionals like {beneficiary_name}, the United States will struggle to maintain its competitive position in {competitive_area}.",
            "The field in which {beneficiary_name} works addresses critical national priorities. {statistic_1}. Furthermore, {statistic_2}. The shortage of qualified professionals in this field threatens {what_threatens}. {beneficiary_name}'s expertise directly contributes to addressing these challenges.",
        ],
        "government_alignment": [
            "{beneficiary_name}'s work aligns directly with priorities established by the federal government. {executive_order_citation} specifically calls for {what_eo_calls_for}. The {strategic_plan} identifies {field} as a critical area requiring immediate attention. {beneficiary_name}'s contributions to {specific_contribution} directly support these national objectives.",
            "The United States government has recognized {field} as a national priority. {government_document} establishes {what_establishes}. {beneficiary_name}'s work on {work_area} directly supports these objectives by {how_supports}. This alignment with established national priorities demonstrates the broader significance of {his_her} contributions.",
        ],
    },
    "unique_skills": [
        "{beneficiary_name} possesses a unique combination of skills that positions {him_her} exceptionally well to advance this endeavor. {He_She} combines {skill_1} with {skill_2}, a rare combination that allows {him_her} to {unique_capability}. Unlike many professionals who specialize in only one aspect, {beneficiary_name} can {bridge_capability}. This dual expertise is precisely what is needed to {what_is_needed}.",
        "What distinguishes {beneficiary_name} from other professionals in this field is {distinguishing_factor}. While many practitioners possess {common_skill}, few combine this with {rare_skill}. {beneficiary_name}'s background in both {background_1} and {background_2} enables {him_her} to approach problems from multiple perspectives, resulting in more effective and innovative solutions.",
    ],
    "waiver_justification": [
        "I believe it would serve the national interest to waive the labor certification requirement for {beneficiary_name}. The traditional labor certification process, designed to protect American workers, would delay {his_her} contributions by 12 to 18 months. Given the urgent national priorities in {field} and the documented shortage of qualified professionals, such delays would be counterproductive. {beneficiary_name}'s work benefits not just a single employer but the broader national effort in {national_effort}. As established in Matter of Dhanasar, when the benefits of a petitioner's work extend beyond the immediate scope of their employment, waiving labor certification serves the national interest.",
        "The national interest would be well served by waiving the labor certification requirement. {field_urgency_statement}. Every month of delay in deploying qualified professionals like {beneficiary_name} represents a real cost to national priorities. The skills {he_she} possesses are in critically short supply, and the traditional labor certification process would unnecessarily delay contributions that the nation urgently needs. Under the Matter of Dhanasar framework, this broader national benefit justifies a waiver.",
    ],
    "conclusion": [
        "In summary, {beneficiary_name} is exactly the type of professional the United States needs to advance its strategic interests in {field}. {His_Her} exceptional abilities, demonstrated through {key_achievements}, combined with {his_her} alignment with national priorities, make {him_her} an ideal candidate for the EB-2 NIW classification. I offer my strongest recommendation for this petition and believe that approving it would serve the national interest.",
        "For the reasons stated above, I strongly support {beneficiary_name}'s petition for an EB-2 National Interest Waiver. {His_Her} contributions to {field} have already benefited {beneficiaries}, and {his_her} continued work in the United States will produce even greater benefits. I am confident that {beneficiary_name} will continue to make significant contributions that serve the national interest. Please do not hesitate to contact me if you require any additional information.",
    ],
}

# Minimum words per section to reach 600+ word target
SECTION_MIN_WORDS = {
    "recommender_credentials": 80,
    "relationship_context": 60,
    "work_description": 120,
    "national_importance": 100,
    "unique_skills": 80,
    "waiver_justification": 100,
    "conclusion": 60,
}


# =============================================================================
# PERSONA CONFIGURATIONS
# =============================================================================

PERSONAS = {
    "executive": {
        "font": "Arial",
        "size": 11,
        "use_tabs": False,
        "salutation": "To the Reviewing Officer:",
        "alt_salutation": "Dear Officer:",
        "closing": "Sincerely,",
        "alt_closing": "Respectfully submitted,",
        "tone": "direct, results-oriented, business language",
        "emphasis": ["ROI", "business impact", "leadership", "deliverables"],
        "typical_phrases": [
            "I'll be direct:",
            "The bottom line is...",
            "From a business perspective...",
            "He/She delivered.",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "It is my pleasure",
            "I wholeheartedly recommend",
        ],
    },
    "technical": {
        "font": "Calibri",
        "size": 11,
        "use_tabs": False,
        "salutation": "To the Officer reviewing this petition:",
        "alt_salutation": "Dear Immigration Officer:",
        "closing": "Sincerely,",
        "alt_closing": "Best regards,",
        "tone": "understated, factual, data-driven",
        "emphasis": ["technical competence", "problem-solving", "innovation", "metrics"],
        "typical_phrases": [
            "Technically speaking...",
            "The data shows...",
            "The system achieved...",
            "From an engineering standpoint...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to recommend",
            "outstanding individual",
        ],
    },
    "academic": {
        "font": "Garamond",
        "size": 12,
        "use_tabs": True,
        "salutation": "Dear Officer,",
        "alt_salutation": "Dear Members of the USCIS Review Panel:",
        "closing": "Respectfully,",
        "alt_closing": "Sincerely,",
        "tone": "formal, methodological, detailed, scholarly",
        "emphasis": ["publications", "methodology", "scientific contribution", "research"],
        "typical_phrases": [
            "From a methodological standpoint...",
            "The significance of this work...",
            "In academic terms...",
            "Research demonstrates...",
        ],
        "forbidden_phrases": [
            "I am writing",
            "highly recommend",
            "exceptional talent",
        ],
    },
    "mentor": {
        "font": "Georgia",
        "size": 11,
        "use_tabs": True,
        "salutation": "Dear Sir or Madam:",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Sincerely,",
        "alt_closing": "With appreciation,",
        "tone": "genuine, warm but professional, narrative",
        "emphasis": ["personal qualities", "development", "growth", "impact on others"],
        "typical_phrases": [
            "I first met...",
            "I still remember when...",
            "Over the years, I watched...",
            "What struck me most...",
        ],
        "forbidden_phrases": [
            "I am writing to",
            "pleasure to write",
            "best student",
        ],
    },
    "corporate": {
        "font": "Times New Roman",
        "size": 12,
        "use_tabs": False,
        "salutation": "To Whom It May Concern:",
        "alt_salutation": "Dear Immigration Officer:",
        "closing": "Respectfully submitted,",
        "alt_closing": "Sincerely,",
        "tone": "professional, structured, objective, methodological",
        "emphasis": ["project management", "coordination", "measurable results", "process"],
        "typical_phrases": [
            "From a project management perspective...",
            "The process involved...",
            "In terms of deliverables...",
            "The results were measurable...",
        ],
        "forbidden_phrases": [
            "I am writing",
            "happy to recommend",
            "strongly recommend",
        ],
    },
    # NEW PERSONAS (added 2026-01-02)
    "collaborator": {
        "font": "Courier New",
        "size": 10,
        "use_tabs": False,
        "salutation": "Dear Reviewing Officer,",
        "alt_salutation": "To the Immigration Officer:",
        "closing": "Best regards,",
        "alt_closing": "Sincerely,",
        "tone": "peer-focused, collaborative, emphasizes joint achievements",
        "emphasis": ["teamwork", "joint publications", "collaborative projects", "peer relationship"],
        "typical_phrases": [
            "As a colleague and collaborator...",
            "We worked together on...",
            "Our joint research demonstrated...",
            "In our collaborative efforts...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "subordinate",
            "my student",
        ],
    },
    "industry_expert": {
        "font": "Verdana",
        "size": 11,
        "use_tabs": False,
        "salutation": "To the Immigration Officer:",
        "alt_salutation": "Dear USCIS Review Panel:",
        "closing": "Respectfully,",
        "alt_closing": "Sincerely,",
        "tone": "authoritative, standards-focused, regulatory insight",
        "emphasis": ["industry standards", "certifications", "regulatory compliance", "market impact"],
        "typical_phrases": [
            "From an industry perspective...",
            "The standards in this field require...",
            "Based on my experience in the industry...",
            "The regulatory implications...",
        ],
        "forbidden_phrases": [
            "I am writing",
            "personally know",
            "pleasure to recommend",
        ],
    },
    "client_partner": {
        "font": "Arial",
        "size": 11,
        "use_tabs": False,
        "salutation": "Dear Sir or Madam:",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Sincerely,",
        "alt_closing": "Best regards,",
        "tone": "business-oriented, results-focused, commercial relationship",
        "emphasis": ["business value", "client results", "commercial impact", "ROI"],
        "typical_phrases": [
            "As a client of...",
            "The business impact was...",
            "Our company benefited from...",
            "The return on investment...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to write",
            "outstanding individual",
        ],
    },
    # NEW PERSONAS (added 2026-01-02 - from real LOR analysis)
    "institute_president": {
        "font": "Times New Roman",
        "size": 12,
        "use_tabs": False,
        "salutation": "Dear Adjudication Officer,",
        "alt_salutation": "To Whom It May Concern:",
        "closing": "Sincerely,",
        "alt_closing": "It is, therefore, without reservation, that I recommend",
        "tone": "formal, institutional, global perspective, alumni-focused",
        "emphasis": ["global impact", "institutional recognition", "alumni achievement", "transformative leadership"],
        "typical_phrases": [
            "I do so in my official capacity as...",
            "distinguished alumnus of our institute...",
            "epitomizes the caliber of global changemakers...",
            "transcends national borders...",
            "world-class expert...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to write",
        ],
    },
    "former_student": {
        "font": "Georgia",
        "size": 11,
        "use_tabs": True,
        "salutation": "To Whom It May Concern,",
        "alt_salutation": "Dear USCIS Officer,",
        "closing": "Sincerely,",
        "alt_closing": "I strongly and wholeheartedly recommend",
        "tone": "personal, emotional, gratitude-filled, mentee perspective",
        "emphasis": ["mentorship impact", "career transformation", "lasting influence", "personal growth"],
        "typical_phrases": [
            "I am writing with great pride and deep respect...",
            "I studied under...",
            "whose academic and professional journey was shaped profoundly by...",
            "I am proud to count myself among them...",
            "It is no exaggeration to say...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "objectively speaking",
        ],
    },
    "international_colleague": {
        "font": "Arial",
        "size": 11,
        "use_tabs": False,
        "salutation": "Dear USCIS Officer,",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Sincerely,",
        "alt_closing": "For these reasons, I strongly recommend",
        "tone": "professional, cross-border perspective, regional expertise",
        "emphasis": ["international experience", "cross-border collaboration", "regional market knowledge", "global standards"],
        "typical_phrases": [
            "I first crossed paths with...",
            "worked in a parallel capacity...",
            "From my international perspective...",
            "Having worked across multiple regions...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to recommend",
        ],
    },
    "conference_acquaintance": {
        "font": "Garamond",
        "size": 12,
        "use_tabs": True,
        "salutation": "Dear Adjudication Officer,",
        "alt_salutation": "Dear Officer,",
        "closing": "Sincerely,",
        "alt_closing": "Respectfully,",
        "tone": "professional, independent observer, symposium-based knowledge",
        "emphasis": ["professional reputation", "presentation quality", "international recognition", "research impact"],
        "typical_phrases": [
            "I first met... at the...",
            "The symposium had X participants from Y countries...",
            "presented a paper titled...",
            "I have followed with great interest the research of...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "close colleague",
            "worked directly together",
        ],
    },
    "hiring_manager": {
        "font": "Arial",
        "size": 11,
        "use_tabs": False,
        "salutation": "To Whom It May Concern,",
        "alt_salutation": "Dear Immigration Officer:",
        "closing": "Sincerely,",
        "alt_closing": "Best regards,",
        "tone": "direct, hiring perspective, career progression focus",
        "emphasis": ["hiring decision", "rapid advancement", "career transformation", "organizational impact"],
        "typical_phrases": [
            "I had the privilege of hiring...",
            "As the person who hired...",
            "later witnessing rapid advancement to...",
            "The impact was immediate and lasting...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to write",
        ],
    },
    "strategic_advisor": {
        "font": "Calibri",
        "size": 11,
        "use_tabs": False,
        "salutation": "Dear Sir or Madam,",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Respectfully,",
        "alt_closing": "Sincerely,",
        "tone": "analytical, strategic, advisory relationship focus",
        "emphasis": ["strategic guidance", "mentorship", "methodology review", "market positioning"],
        "typical_phrases": [
            "I have worked closely with... as their strategic advisor...",
            "entered into a formal mentorship and strategic advisory relationship...",
            "My knowledge of... work is direct and ongoing...",
            "In that capacity, I have reviewed methodology...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to recommend",
        ],
    },
    # NEW PERSONAS (added 2026-01-02 - from real LOR analysis session)
    "nonprofit_leader": {
        "font": "Georgia",
        "size": 11,
        "use_tabs": True,
        "salutation": "Dear USCIS Officer,",
        "alt_salutation": "To Whom It May Concern:",
        "closing": "Sincerely,",
        "alt_closing": "With respect,",
        "tone": "mission-focused, practical, community impact",
        "emphasis": ["nonprofit mission", "community service", "practical outcomes", "sector expertise"],
        "typical_phrases": [
            "In fifteen years of nonprofit leadership, I have learned that...",
            "The sector needs professionals who...",
            "From a mission-driven perspective...",
            "The community impact has been...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to recommend",
        ],
    },
    "long_term_colleague": {
        "font": "Arial",
        "size": 11,
        "use_tabs": False,
        "salutation": "Dear Immigration Officer,",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Sincerely,",
        "alt_closing": "Best regards,",
        "tone": "storytelling, experiential, 10+ years perspective",
        "emphasis": ["long-term observation", "professional growth", "consistent excellence", "trusted relationship"],
        "typical_phrases": [
            "I first met [name] in [year]...",
            "Over the past [X] years, I have watched...",
            "Our professional relationship spans...",
            "Through the years, I have observed...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to write",
        ],
    },
    "academic_advisor": {
        "font": "Times New Roman",
        "size": 12,
        "use_tabs": True,
        "salutation": "Dear Sir or Madam:",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Respectfully,",
        "alt_closing": "Sincerely,",
        "tone": "formal academic, professor-student dynamic, research-focused",
        "emphasis": ["academic rigor", "research methodology", "scholarly contribution", "intellectual growth"],
        "typical_phrases": [
            "As [name]'s academic advisor...",
            "During graduate studies under my supervision...",
            "From an academic standpoint...",
            "The scholarly contribution demonstrates...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "brilliant student",
        ],
    },
    "alumni_colleague": {
        "font": "Calibri",
        "size": 11,
        "use_tabs": False,
        "salutation": "To Whom It May Concern:",
        "alt_salutation": "Dear Immigration Officer:",
        "closing": "Sincerely,",
        "alt_closing": "Best regards,",
        "tone": "collegial, shared history, long-term relationship",
        "emphasis": ["shared educational background", "professional network", "career trajectory", "mutual respect"],
        "typical_phrases": [
            "I have known [name] for over [X] years...",
            "We first met as graduate students at...",
            "Our shared background allows me to...",
            "Having followed [name]'s career since...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to recommend",
        ],
    },
    "fellow_award_recipient": {
        "font": "Garamond",
        "size": 12,
        "use_tabs": True,
        "salutation": "Dear Adjudication Officer,",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Respectfully,",
        "alt_closing": "Sincerely,",
        "tone": "peer recognition, shared credential, mutual achievement",
        "emphasis": ["shared recognition", "competitive selection", "peer validation", "industry standing"],
        "typical_phrases": [
            "As a fellow recipient of the [award]...",
            "Having been recognized alongside [name] for...",
            "The selection process for this award...",
            "Among the recipients, [name] stood out...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to recommend",
        ],
    },
    "global_engineering_colleague": {
        "font": "Arial",
        "size": 11,
        "use_tabs": False,
        "salutation": "Dear USCIS Officer,",
        "alt_salutation": "To the Reviewing Officer:",
        "closing": "Sincerely,",
        "alt_closing": "Best regards,",
        "tone": "cross-site collaboration, multinational perspective, engineering standards",
        "emphasis": ["global standards", "cross-site coordination", "multinational projects", "engineering excellence"],
        "typical_phrases": [
            "I am writing to provide my assessment...",
            "In my role coordinating across multiple sites...",
            "From a global engineering perspective...",
            "Working across international boundaries...",
        ],
        "forbidden_phrases": [
            "I am writing to recommend",
            "pleasure to recommend",
        ],
    },
}

PersonaType = Literal["executive", "technical", "academic", "mentor", "corporate", "collaborator", "industry_expert", "client_partner", "institute_president", "former_student", "international_colleague", "conference_acquaintance", "hiring_manager", "strategic_advisor", "nonprofit_leader", "long_term_colleague", "academic_advisor", "alumni_colleague", "fellow_award_recipient", "global_engineering_colleague"]


# =============================================================================
# OPENING STYLE CONFIGURATIONS (based on real LOR analysis)
# =============================================================================

OPENING_STYLES = {
    "direct_personal": {
        "description": "Start with direct personal introduction",
        "template": "I'm {recommender_name}, {recommender_title} at {recommender_org}, and I hereby express my support for {beneficiary_name}'s petition for the {visa_type}.",
        "example": "I'm John Duren, National Account Manager at MPSW..."
    },
    "formal_declaration": {
        "description": "Start with formal declaration of support",
        "template": "My name is {recommender_name}. I am writing this letter in strong support of {beneficiary_name} and the petition for an {visa_type}.",
        "example": "My name is Federica Bartoli, and I am writing this letter..."
    },
    "credentials_first": {
        "description": "Lead with recommender's credentials before mentioning beneficiary",
        "template": "My name is {recommender_name}, and I am {recommender_title} at {recommender_org}. With {years_experience} years of experience in this field, I am qualified to assess professionals in this domain. I write to support {beneficiary_name}.",
        "example": "My name is Glendon J. Parker, Ph.D. I am an Adjunct Associate Professor..."
    },
    "anecdote_opener": {
        "description": "Start with a memorable story or moment",
        "template": "When {beneficiary_name} joined my team in {year}, we had a gap. We needed {role_description}. Finding someone who could handle that scope independently is harder than it sounds. {beneficiary_name} filled that gap from day one.",
        "example": "When Abhishek Sahu joined my team at Arctic Wolf..."
    },
    "context_first": {
        "description": "Establish context before introducing people",
        "template": "I am pleased to write this letter of recommendation in support of {beneficiary_name}'s petition for a {visa_type}. As the {recommender_title} of {recommender_org}, a {org_description}, I have extensive expertise in {field_description}.",
        "example": "I am pleased to write this letter of recommendation..."
    },
    "symposium_meeting": {
        "description": "Explain meeting at professional event",
        "template": "I first met {beneficiary_name} in {year}, at the {event_name} held in {location}. The {event_type} had {participant_count} participants from {country_count} countries. {beneficiary_name} presented a paper titled '{paper_title}.'",
        "example": "I first met Anselm in 2012, at the 4th International Symposium..."
    },
    "project_collaboration": {
        "description": "Start with specific project context",
        "template": "I came to know {beneficiary_name} during the work on {project_name}. At that time, we collaborated closely on {project_description} across multiple sites.",
        "example": "I came to know Omar during his tenure as Technical Service Manager..."
    },
    "independent_evaluator": {
        "description": "Emphasize objective, independent perspective",
        "template": "I have followed with great interest the research of {beneficiary_name}, particularly the innovative work on {research_area}. Although I did not work directly with {beneficiary_name}, I became familiar with this work while evaluating {evaluation_context}. This makes my perspective objective and credible.",
        "example": "I have followed with great interest the research of Professor..."
    },
    # NEW OPENING STYLES (added 2026-01-02 - from real LOR analysis)
    "strong_support": {
        "description": "Lead with strong support declaration",
        "template": "I am writing to offer my strong support for {beneficiary_name} in the application for permanent residency in the United States, under the {visa_type} visa category. My name is {recommender_name}, and I am {recommender_title} at {recommender_org}.",
        "example": "I am writing to offer my strong support for Mr. Omar Abdelmotelb..."
    },
    "strongest_support": {
        "description": "Emphatic support with professional context",
        "template": "I would like to offer my strongest support to {beneficiary_name}'s {visa_type} application. My name is {recommender_name}, and I am currently {recommender_title} at {recommender_org}.",
        "example": "I would like to offer my strongest support to Mr. Omar Abdelmotelb's EB-2 NIW application..."
    },
    "recognition": {
        "description": "Frame as recognition letter",
        "template": "I am pleased to write this letter in recognition of {beneficiary_name} and the invaluable contributions to {field_description}. My name is {recommender_name}, and I serve as {recommender_title} at {recommender_org}.",
        "example": "I am pleased to write this letter in recognition of Ms. Kosalee Galkaduwa..."
    },
    "pride_and_respect": {
        "description": "Personal, emotional opening for mentee relationships",
        "template": "I am writing with great pride and deep respect to recommend {beneficiary_name} for classification under the {visa_type} immigrant visa category. I do so not only as a {relationship}, but also as one of the many individuals whose {impact_description} were shaped profoundly by {beneficiary_name}'s mentorship, vision, and exceptional contributions.",
        "example": "I am writing with great pride and deep respect to recommend Professor Anselm..."
    },
    "official_capacity": {
        "description": "Formal institutional declaration",
        "template": "I, {recommender_name}, am pleased to write this letter in support of the petition of {beneficiary_name} for classification as an individual of {classification_type} under the United States immigrant visa program. I do so in my official capacity as the {recommender_title} of {recommender_org}, with over {years_experience} years of experience in this position.",
        "example": "I, Dr. Joseph Shevel, am pleased to write this letter in support..."
    },
    "present_letter": {
        "description": "Direct acknowledgment opening",
        "template": "In the present letter, I want to acknowledge {beneficiary_name}'s outstanding {achievement_type} by contextualizing some key achievements. My name is {recommender_name}, and I serve as {recommender_title} at {recommender_org}.",
        "example": "In the present letter, I want to acknowledge Ms. Kosalee's outstanding leadership..."
    },
    "hiring_relationship": {
        "description": "Focus on hiring and career progression",
        "template": "I am pleased to write in support of {beneficiary_name}'s petition for permanent residence in the United States under the {visa_type} category. I had the privilege of hiring {beneficiary_name} as {initial_role} and later witnessing rapid advancement to {current_role}, a leadership position in which {beneficiary_name} reshaped and elevated our {department} practices.",
        "example": "I had the privilege of hiring Ms. Galkaduwa as a Production Planning Coordinator..."
    },
    "appreciation_letter": {
        "description": "Direct appreciation addressed to candidate (internal commendation style)",
        "template": "Dear {beneficiary_name},\n\nI hope this message finds you well. I am writing to formally express my appreciation for the outstanding work and dedication you have demonstrated since joining {org_name} as our {role}.",
        "example": "Dear Kosalee, I hope this message finds you well. I am writing to formally express..."
    },
    "sectioned_intro": {
        "description": "Lead into structured sections",
        "template": "I am pleased to write this letter in support of {beneficiary_name}'s petition. As the {recommender_title} at {recommender_org}, I was pleased to witness firsthand exceptional skills, professionalism, and dedication to our organization.",
        "example": "As the Executive Office Manager at Gelberg Signs, I was pleased to witness..."
    },
    # NEW OPENING STYLES (added 2026-01-02 - from additional LOR analysis)
    "wisdom_opener": {
        "description": "Start with wisdom/experience statement",
        "template": "In {years_experience} years of {field} leadership, I have learned that the most impactful professionals are those who combine technical excellence with a genuine commitment to their field. {beneficiary_name} exemplifies this principle.",
        "example": "In fifteen years of nonprofit leadership, I have learned that the most impactful professionals..."
    },
    "storytelling_opener": {
        "description": "Narrative opening with first meeting",
        "template": "I first met {beneficiary_name} in {year}, when {context}. Since then, I have watched {beneficiary_name}'s career evolve from {initial_role} to {current_role}, consistently demonstrating {qualities}.",
        "example": "I first met Abhishek in 2014, when he joined as a junior engineer. Since then, I have watched..."
    },
    "long_term_relationship": {
        "description": "Emphasize duration of relationship",
        "template": "I have known {beneficiary_name} for over {years} years, first as {initial_relationship} and now as {current_relationship}. This extended professional relationship gives me a unique perspective on {beneficiary_name}'s capabilities and contributions.",
        "example": "I have known Edwin for over twenty years, first as a graduate student and now as a respected colleague..."
    },
    "assessment_framing": {
        "description": "Frame as professional assessment",
        "template": "I am writing to provide my assessment of {beneficiary_name}'s qualifications for the {visa_type} petition. In my capacity as {recommender_title} at {recommender_org}, I have directly observed {beneficiary_name}'s work and can speak to both technical competence and professional impact.",
        "example": "I am writing to provide my assessment of Mr. Marcondes's qualifications for the EB-2 NIW petition..."
    },
}


# =============================================================================
# LETTER STRUCTURE OPTIONS
# =============================================================================

LETTER_STRUCTURES = {
    "sectioned": {
        "description": "Explicit sections with headers",
        "sections": [
            "How I am qualified to evaluate {beneficiary_name}'s work",
            "My relationship with {beneficiary_name} and what I personally observed",
            "{beneficiary_name}'s contributions and why they matter in the national interest",
            "Evidence of leadership and critical role",
            "Why the U.S. benefits from {beneficiary_name} continuing this work here",
            "Conclusion",
        ],
        "style_notes": "Use bold headers, clear separation between sections"
    },
    "flowing_narrative": {
        "description": "Continuous prose without explicit headers",
        "sections": None,
        "style_notes": "Natural flow, transitions between topics, personal voice"
    },
    "problem_solution": {
        "description": "Structure around problems and solutions",
        "sections": [
            "The challenge we faced",
            "How {beneficiary_name} addressed it",
            "The measurable results",
            "National implications",
            "Why this matters",
        ],
        "style_notes": "Focus on concrete problems and quantifiable outcomes"
    },
    "chronological": {
        "description": "Temporal progression of relationship and achievements",
        "sections": None,
        "style_notes": "Timeline-based narrative, growth over time"
    },
    # NEW STRUCTURES (added 2026-01-02 - from real LOR analysis)
    "question_headers": {
        "description": "Structure with question-style headers",
        "sections": [
            "Who She/He Is",
            "What She/He Did",
            "Projects Led or Played a Crucial Role In",
            "Contribution to Success",
            "Impact of Efforts",
        ],
        "style_notes": "Clear section headers, concise summaries under each"
    },
    "conference_narrative": {
        "description": "Based on meeting at professional event",
        "sections": [
            "How We Met",
            "The Work Presented",
            "Significance of Research",
            "Continued Impact",
            "National Interest",
        ],
        "style_notes": "Opens with conference/symposium meeting, builds credibility through professional context"
    },
    "career_progression": {
        "description": "Tracks career growth under recommender's observation",
        "sections": [
            "Initial Hiring/Meeting",
            "Early Contributions",
            "Growth and Advancement",
            "Current Impact",
            "Future Potential",
        ],
        "style_notes": "Hiring manager perspective, shows development arc"
    },
    "appreciation_format": {
        "description": "Internal commendation/appreciation letter style",
        "sections": None,
        "style_notes": "Addressed directly to candidate, more personal, internal recognition focus"
    },
    "international_perspective": {
        "description": "Cross-border collaboration focus",
        "sections": [
            "International Context",
            "Cross-Border Collaboration",
            "Technical Achievements",
            "Global Standards Alignment",
            "U.S. National Interest",
        ],
        "style_notes": "Emphasizes global perspective, regional expertise, transferable knowledge"
    },
}


# =============================================================================
# RELATIONSHIP TYPES (expanded based on real LOR analysis)
# =============================================================================

RELATIONSHIP_TYPES = {
    "direct_supervisor": {
        "label": "Direct Supervisor/Manager",
        "strength": "strong",
        "perspective": "direct oversight of work quality and impact",
        "typical_phrases": [
            "In my role managing {name}...",
            "As {name}'s direct supervisor...",
            "During my oversight of {name}'s projects...",
        ]
    },
    "mentor": {
        "label": "Mentor/Advisor",
        "strength": "strong",
        "perspective": "professional development and growth",
        "typical_phrases": [
            "As {name}'s mentor...",
            "Through my role advising {name}...",
            "Having guided {name}'s professional development...",
        ]
    },
    "colleague_peer": {
        "label": "Colleague/Peer",
        "strength": "moderate",
        "perspective": "day-to-day collaboration and technical work",
        "typical_phrases": [
            "As a colleague of {name}...",
            "Working alongside {name}...",
            "In our collaborative work...",
        ]
    },
    "independent_evaluator": {
        "label": "Independent Evaluator (knows work by reputation)",
        "strength": "very_strong",
        "perspective": "objective, third-party assessment",
        "typical_phrases": [
            "Although I did not work directly with {name}...",
            "I became familiar with {name}'s work through...",
            "My independent assessment is based on...",
        ]
    },
    "client_partner": {
        "label": "Client/Business Partner",
        "strength": "moderate",
        "perspective": "business value and results delivered",
        "typical_phrases": [
            "As a client of {name}'s services...",
            "In our business relationship...",
            "Our company benefited from {name}'s...",
        ]
    },
    "conference_acquaintance": {
        "label": "Professional Conference/Symposium Connection",
        "strength": "moderate",
        "perspective": "professional reputation and presentation",
        "typical_phrases": [
            "I first met {name} at...",
            "We connected through the {conference}...",
            "After {name}'s presentation at...",
        ]
    },
    "thesis_advisor": {
        "label": "Thesis/Dissertation Advisor",
        "strength": "strong",
        "perspective": "academic rigor and research capability",
        "typical_phrases": [
            "As {name}'s thesis advisor...",
            "During {name}'s graduate studies under my supervision...",
            "I oversaw {name}'s doctoral research on...",
        ]
    },
    "project_collaborator": {
        "label": "Project Collaborator (specific project)",
        "strength": "moderate",
        "perspective": "specific project contributions",
        "typical_phrases": [
            "During our collaboration on {project}...",
            "We worked together on the {project}...",
            "In the context of {project}...",
        ]
    },
    "potential_employer": {
        "label": "Potential Employer/Recruiter",
        "strength": "moderate",
        "perspective": "market value and industry demand",
        "typical_phrases": [
            "I am currently considering {name} for...",
            "In my assessment of {name} for a potential role...",
            "Our organization is interested in {name}'s expertise...",
        ]
    },
    "industry_peer": {
        "label": "Industry Expert/Peer (same field)",
        "strength": "strong",
        "perspective": "technical competence and industry standing",
        "typical_phrases": [
            "As a fellow professional in {field}...",
            "From my industry perspective...",
            "Based on my experience in {field}...",
        ]
    },
    # NEW RELATIONSHIP TYPES (added 2026-01-02 - from real LOR analysis)
    "distinguished_alumnus": {
        "label": "Institute President/Alumnus Relationship",
        "strength": "strong",
        "perspective": "institutional achievement and global recognition",
        "typical_phrases": [
            "As a distinguished alumnus of our institute...",
            "During training with us, exhibited exceptional...",
            "We have proudly followed this career as...",
            "epitomizes the caliber of global changemakers we strive to develop...",
        ]
    },
    "former_student_mentee": {
        "label": "Former Student/Mentee (reverse perspective)",
        "strength": "moderate",
        "perspective": "personal impact of mentorship, career transformation",
        "typical_phrases": [
            "I studied under...",
            "As a former student of...",
            "Whose academic and professional journey was shaped profoundly by...",
            "I am proud to count myself among the many individuals...",
        ]
    },
    "hiring_manager": {
        "label": "Person Who Hired the Candidate",
        "strength": "strong",
        "perspective": "career progression, organizational impact",
        "typical_phrases": [
            "I had the privilege of hiring...",
            "As the person who hired {name}...",
            "I witnessed firsthand {name}'s rapid advancement...",
            "The impact was immediate and lasting...",
        ]
    },
    "strategic_advisor": {
        "label": "Strategic Advisor/Formal Mentor",
        "strength": "strong",
        "perspective": "strategic guidance, methodology review, market positioning",
        "typical_phrases": [
            "I have worked closely with {name} as strategic advisor...",
            "Entered into a formal mentorship and strategic advisory relationship...",
            "My knowledge of {name}'s work is direct and ongoing...",
            "In that capacity, I have reviewed methodology and advised on...",
        ]
    },
    "parallel_capacity": {
        "label": "Parallel/Cross-Functional Colleague",
        "strength": "moderate",
        "perspective": "observation from adjacent role, cross-team coordination",
        "typical_phrases": [
            "I first crossed paths with {name}...",
            "Worked in a parallel capacity as part of the same...",
            "This structure allowed me to directly observe work in real time...",
            "Through our collaboration, I saw how {name} navigated...",
        ]
    },
    "company_founder": {
        "label": "Company Founder/CEO",
        "strength": "very_strong",
        "perspective": "organizational vision, strategic impact",
        "typical_phrases": [
            "As founder and CEO, I had the privilege of working with...",
            "Our company's growth is directly tied to...",
            "{name}'s mandate required taking ownership of...",
            "transformed them into structured, data-driven systems...",
        ]
    },
    "research_observer": {
        "label": "Research Observer (knows work by publications/reputation)",
        "strength": "very_strong",
        "perspective": "academic reputation, citation impact, field influence",
        "typical_phrases": [
            "I have followed with great interest the research of...",
            "Data from the Google Scholar Citation Index confirms...",
            "publications have been referenced by researchers based in...",
            "His/her work advances critical themes in...",
        ]
    },
}


# =============================================================================
# HEADER FORMAT OPTIONS
# =============================================================================

HEADER_FORMATS = {
    "full_letterhead": {
        "description": "Complete letterhead with all contact details",
        "elements": ["name", "title", "organization", "address", "email", "phone", "date"],
        "alignment": "left"
    },
    "simple_header": {
        "description": "Just name and date",
        "elements": ["name", "date"],
        "alignment": "left"
    },
    "institutional": {
        "description": "Organization-first format",
        "elements": ["organization", "name", "title", "email", "date"],
        "alignment": "left"
    },
    "date_address_first": {
        "description": "Date and USCIS address before recommender info",
        "elements": ["date", "uscis_address", "re_line", "name", "title", "organization"],
        "alignment": "left"
    },
}


# =============================================================================
# ADDITIONAL SALUTATIONS (from real LOR analysis)
# =============================================================================

ADDITIONAL_SALUTATIONS = [
    "To Whom It May Concern:",
    "Dear Sir or Madam:",
    "To the Reviewing Officer:",
    "Dear Adjudication Officer:",
    "USCIS Officer,",
    "To the United States Citizenship and Immigration Services:",
    "Dear Immigration Officer:",
    "Dear Members of the USCIS Review Panel:",
    "To the Officer reviewing this petition:",
    "Dear Officer,",
    # NEW SALUTATIONS (added 2026-01-02 - from real LOR analysis)
    "Dear USCIS Officer,",
    "Adjudication Officer,",  # Without "Dear"
    "Dear Sir or Madam,",  # Comma variant
    "RE: Letter of Support for {beneficiary_name}'s {visa_type} Petition",  # RE line format
]


# =============================================================================
# CLOSING VARIATIONS (from real LOR analysis)
# =============================================================================

CLOSING_VARIATIONS = [
    "Sincerely,",
    "Respectfully submitted,",
    "Best regards,",
    "Regards,",
    "Respectfully,",
    "With appreciation,",
    "Very truly yours,",
    # NEW CLOSINGS (added 2026-01-02 - from real LOR analysis)
    "Warm regards,",  # Personal/appreciation style
    "Sincerely yours,",
    "I offer my unequivocal support for this petition.",
    "I recommend {beneficiary_name} without reservation.",
    "I strongly and wholeheartedly recommend {beneficiary_name}.",
    "For these reasons, I strongly recommend {beneficiary_name}.",
    "I support this petition and believe that allowing {beneficiary_name} to continue this work will produce outsized benefits.",
    "It is, therefore, without reservation, that I recommend {beneficiary_name}.",
]


# =============================================================================
# EVIDENCE PRESENTATION STYLES
# =============================================================================

EVIDENCE_STYLES = {
    "footnoted": {
        "description": "Academic style with numbered footnotes",
        "use_footnotes": True,
        "citation_style": "numbered",
        "example": "...as documented in research studies.[1]"
    },
    "inline_citations": {
        "description": "Citations within the text",
        "use_footnotes": False,
        "citation_style": "inline",
        "example": "According to the CyberSeek 2024 report..."
    },
    "statistics_heavy": {
        "description": "Lead with numbers and data",
        "use_footnotes": True,
        "citation_style": "numbered",
        "example": "The U.S. faces a 500,000+ professional shortage..."
    },
    "narrative_with_data": {
        "description": "Story-driven with supporting statistics",
        "use_footnotes": False,
        "citation_style": "embedded",
        "example": "This represents a real problem: industry reports show..."
    },
}


# =============================================================================
# PARAGRAPH LENGTH PREFERENCES
# =============================================================================

PARAGRAPH_STYLES = {
    "short_punchy": {
        "description": "Short, impactful paragraphs (2-4 sentences)",
        "max_sentences": 4,
        "style": "direct, executive tone"
    },
    "medium_balanced": {
        "description": "Standard paragraphs (4-6 sentences)",
        "max_sentences": 6,
        "style": "balanced professional tone"
    },
    "long_detailed": {
        "description": "Comprehensive paragraphs (6+ sentences)",
        "max_sentences": 10,
        "style": "academic, detailed exposition"
    },
}


# =============================================================================
# INTERNATIONAL CONTEXT OPTIONS
# =============================================================================

INTERNATIONAL_CONTEXTS = {
    "us_based": {
        "description": "Recommender based in the United States",
        "perspective_notes": "Domestic expertise, local market knowledge"
    },
    "international_perspective": {
        "description": "Recommender based outside the US",
        "perspective_notes": "Global view, comparative expertise, international recognition",
        "typical_phrases": [
            "From my international perspective...",
            "Having worked in {country} where {context}...",
            "Compared to other markets globally...",
        ]
    },
    "multinational_experience": {
        "description": "Recommender with experience in multiple countries",
        "perspective_notes": "Cross-border expertise, global standards",
        "typical_phrases": [
            "In my experience across {countries}...",
            "Having led projects in multiple countries...",
            "From a global industry standpoint...",
        ]
    },
}


# =============================================================================
# NATIONAL INTEREST TEXTS BY FIELD
# =============================================================================

NATIONAL_INTEREST_TEXTS = {
    "cybersecurity": {
        "intro": "The United States has identified cybersecurity as a critical priority for national security and economic stability",
        "context": "Cyber threats to American infrastructure, businesses, and citizens continue to escalate in sophistication and frequency. The average cost of a data breach in the U.S. reached $9.48 million in 2023.",
        "scarcity": "The cybersecurity workforce gap in the United States exceeds 500,000 unfilled positions according to CyberSeek.",
        "executive_orders": [
            "Executive Order 14028, 'Improving the Nation's Cybersecurity,' 86 Fed. Reg. 26633 (May 12, 2021).",
            "Executive Order 14144, 'Strengthening and Promoting Innovation in the Nation's Cybersecurity,' 90 Fed. Reg. 4551 (January 16, 2025).",
        ],
        "sources": [
            "IBM Security, 'Cost of a Data Breach Report 2023,' IBM Corporation, 2023.",
            "CyberSeek, 'Cybersecurity Supply/Demand Heat Map,' NICE, 2024.",
            "The White House, 'National Cybersecurity Strategy,' March 2023.",
        ],
    },
    "ai_ml": {
        "intro": "The United States has identified artificial intelligence and machine learning as critical and emerging technologies",
        "context": "AI and ML technologies are transforming industries from healthcare to national security, and maintaining American leadership in this field is a national priority.",
        "scarcity": "Professionals with advanced expertise in AI/ML systems are in critically short supply, with demand growing faster than educational institutions can produce qualified candidates.",
        "executive_orders": [
            "Executive Order 14110, 'Safe, Secure, and Trustworthy Development and Use of Artificial Intelligence,' 88 Fed. Reg. 75191 (October 30, 2023).",
        ],
        "sources": [
            "The White House, 'National Artificial Intelligence R&D Strategic Plan, 2023 Update,' May 2023.",
            "Critical and Emerging Technologies List Update, February 2024, Executive Office of the President.",
        ],
    },
    "clean_energy": {
        "intro": "The United States has identified clean energy and advanced manufacturing as strategic priorities",
        "context": "The transition to renewable energy sources is essential for addressing climate change and achieving energy independence.",
        "scarcity": "Engineers and scientists with expertise in solar, wind, battery technology, and energy storage are in high demand.",
        "executive_orders": [
            "Executive Order on Tackling the Climate Crisis at Home and Abroad, January 2021.",
        ],
        "sources": [
            "Critical and Emerging Technologies List Update, February 2024, Executive Office of the President.",
            "Inflation Reduction Act of 2022, Pub. L. No. 117-169.",
        ],
    },
    "stem_education": {
        "intro": "The United States faces a well-documented shortage of qualified STEM educators",
        "context": "Access to qualified STEM educators is one of the strongest predictors of student success in science, technology, engineering, and mathematics fields.",
        "scarcity": "The national shortage of mathematics and science teachers, especially in urban and rural schools serving disadvantaged populations, represents a significant threat to American competitiveness.",
        "executive_orders": [
            "Executive Order 14081, 'Advancing Biotechnology and Biomanufacturing Innovation' (Sept. 12, 2022), emphasizing the importance of STEM education for national competitiveness.",
        ],
        "sources": [
            "U.S. Department of Education, Teacher Shortage Areas Report (2024).",
            "National Center for Education Statistics, Characteristics of Public School Teachers (2023).",
            "Learning Policy Institute, 'A Coming Crisis in Teaching? Teacher Supply, Demand, and Shortages in the U.S.' (2023).",
        ],
    },
    "biotech": {
        "intro": "The United States has recognized biotechnology and biomanufacturing as areas of critical national importance",
        "context": "Advances in biotechnology have implications for healthcare, agriculture, environmental sustainability, and national security.",
        "scarcity": "Researchers with the specialized expertise required for breakthrough discoveries in this field are exceptionally rare.",
        "executive_orders": [
            "Executive Order on Advancing Biotechnology and Biomanufacturing Innovation for a Sustainable, Safe, and Secure American Bioeconomy, 87 Fed. Reg. 56849 (September 12, 2022).",
        ],
        "sources": [
            "National Institutes of Health, 'NIH-Wide Strategic Plan for Fiscal Years 2021-2025.'",
        ],
    },
    "semiconductor": {
        "intro": "The United States has identified advanced semiconductor technology as essential to national security and economic competitiveness",
        "context": "Semiconductor chips are foundational to virtually every sector of the modern economy, from consumer electronics to defense systems.",
        "scarcity": "Engineers with expertise in advanced semiconductor design and manufacturing are in critically short supply domestically.",
        "executive_orders": [],
        "sources": [
            "CHIPS and Science Act of 2022, Creating Helpful Incentives to Produce Semiconductors for America.",
        ],
    },
    "healthcare": {
        "intro": "The United States faces significant challenges in healthcare delivery and medical innovation",
        "context": "Improving healthcare outcomes while controlling costs requires innovative solutions and skilled professionals.",
        "scarcity": "Healthcare professionals with specialized expertise are in critically short supply across the country.",
        "executive_orders": [],
        "sources": [
            "National Institutes of Health, 'NIH-Wide Strategic Plan for Fiscal Years 2021-2025.'",
            "U.S. Department of Health and Human Services, Health Resources and Services Administration.",
        ],
    },
    # NEW FIELDS (added 2026-01-02)
    "quantum_computing": {
        "intro": "The United States has identified quantum computing as a critical emerging technology essential for national security and technological leadership",
        "context": "Quantum computing represents a paradigm shift in computational capability, with applications ranging from cryptography to drug discovery and materials science.",
        "scarcity": "The U.S. faces a critical shortage of quantum computing experts, with only approximately 3,000 professionals nationwide possessing the advanced skills required.",
        "executive_orders": [
            "National Quantum Initiative Act of 2018, establishing a national quantum computing research program.",
            "Executive Order 13885, 'Establishing the National Quantum Initiative Advisory Committee' (August 30, 2019).",
        ],
        "sources": [
            "National Science and Technology Council, 'National Strategic Overview for Quantum Information Science' (2018).",
            "Quantum Economic Development Consortium (QED-C), 'Quantum Workforce Development' Report (2024).",
        ],
    },
    "fintech": {
        "intro": "The United States recognizes financial technology innovation as essential for maintaining global financial leadership and economic competitiveness",
        "context": "Fintech is transforming banking, payments, lending, and investment services, creating new opportunities while requiring specialized expertise.",
        "scarcity": "The financial technology sector requires over 50,000 new specialists annually, with demand far outpacing the supply of qualified professionals.",
        "executive_orders": [
            "Treasury Department Strategic Plan 2022-2026, emphasizing financial innovation and technology.",
            "SEC Innovation Initiative, promoting responsible fintech development.",
        ],
        "sources": [
            "U.S. Department of the Treasury, 'A Financial System That Creates Economic Opportunities: Nonbank Financials, Fintech, and Innovation' (2018).",
            "Financial Stability Oversight Council, Annual Report (2024).",
        ],
    },
    "advanced_manufacturing": {
        "intro": "The United States has identified advanced manufacturing as critical for economic security, supply chain resilience, and national competitiveness",
        "context": "Reshoring critical manufacturing and developing next-generation production capabilities are national priorities that require specialized talent.",
        "scarcity": "The manufacturing skills gap is projected to result in 2.1 million unfilled jobs by 2030, threatening American industrial competitiveness.",
        "executive_orders": [
            "CHIPS and Science Act of 2022, investing in domestic semiconductor and advanced manufacturing.",
            "Executive Order 14017, 'America's Supply Chains' (February 24, 2021), requiring resilient domestic manufacturing.",
        ],
        "sources": [
            "Deloitte and The Manufacturing Institute, 'Creating Pathways for Tomorrow's Workforce Today' (2023).",
            "National Institute of Standards and Technology, 'Manufacturing USA Strategic Plan' (2024).",
        ],
    },
    "space_technology": {
        "intro": "The United States has identified space technology and commercial space development as critical national priorities for security and economic leadership",
        "context": "Space dominance is essential for national security, communications, Earth observation, and emerging commercial opportunities.",
        "scarcity": "The commercial space industry is growing at 9% annually, with talent demand significantly exceeding the supply of qualified aerospace professionals.",
        "executive_orders": [
            "Space Policy Directive-1, 'Reinvigorating America's Human Space Exploration Program' (December 11, 2017).",
            "National Space Strategy, establishing American leadership in space (2020).",
        ],
        "sources": [
            "Space Foundation, 'The Space Report 2024: Q2.'",
            "NASA Workforce Strategy (2023).",
            "FAA Office of Commercial Space Transportation, Annual Report (2024).",
        ],
    },
    # NEW FIELDS (added 2026-01-02 - from real LOR analysis)
    "hvac": {
        "intro": "The United States recognizes heating, ventilation, and air conditioning (HVAC) systems as critical infrastructure for public health, energy efficiency, and climate resilience",
        "context": "HVAC systems are essential for healthcare facilities, data centers, and commercial buildings. Higher ventilation rates reduce transmission of infectious agents and improve patient outcomes. The quality of indoor air impacts asthma symptoms and respiratory health.",
        "scarcity": "The HVACR industry holds over 137,000 jobs with significant projected growth. The market is projected to reach $31.4 billion by 2034, driven by smart and sustainable systems.",
        "executive_orders": [
            "Building Performance Standards, promoting energy-efficient building systems.",
        ],
        "sources": [
            "U.S. Bureau of Labor Statistics, HVACR Industry Employment Data (2024).",
            "IBISWorld, HVAC Industry Market Report (2024).",
            "CDC Guidelines for Environmental Infection Control in Healthcare Facilities.",
        ],
    },
    "procurement": {
        "intro": "The United States recognizes resilient supply chain management and procurement as critical to economic security and business continuity",
        "context": "Resilient, transparent, and fraud-aware procurement in digital commerce protects consumers, legitimate vendors, and market integrity. Data-driven procurement with embedded risk controls touches multiple sectors.",
        "scarcity": "Professionals with expertise in strategic sourcing, vendor risk management, and procurement analytics are in high demand as organizations strengthen supply chains.",
        "executive_orders": [
            "Executive Order 14017, 'America's Supply Chains' (February 24, 2021), requiring resilient supply chain management.",
        ],
        "sources": [
            "U.S. Department of Commerce, Supply Chain Resilience Report (2024).",
            "Institute for Supply Management, Supply Management Salary Survey (2024).",
        ],
    },
    "agriculture": {
        "intro": "The United States has identified sustainable agriculture and food security as critical national priorities",
        "context": "Improving agricultural productivity while reducing environmental impact requires innovative solutions. Climate-smart agriculture addresses challenges of food security and environmental sustainability.",
        "scarcity": "Agricultural scientists with expertise in sustainable farming practices, alternative feed systems, and climate adaptation are in critically short supply.",
        "executive_orders": [
            "USDA Climate Smart Agriculture Initiative.",
        ],
        "sources": [
            "USDA, 'Climate Smart Agriculture and Forestry Strategy' (2024).",
            "National Academies of Sciences, 'Science Breakthroughs to Advance Food and Agricultural Research by 2030.'",
        ],
    },
    "animal_science": {
        "intro": "The United States recognizes animal science and livestock nutrition as essential for food security and sustainable agriculture",
        "context": "Research in animal nutrition, alternative feed systems, and sustainable livestock production addresses challenges of food security and environmental sustainability. This work advances Sustainable Development Goals including Zero Hunger and Responsible Consumption.",
        "scarcity": "Researchers with specialized expertise in animal nutrition, feed innovation, and sustainable livestock systems are exceptionally rare.",
        "executive_orders": [],
        "sources": [
            "USDA National Institute of Food and Agriculture, Animal Science Research Programs.",
            "FAO, 'The State of Food and Agriculture' (2024).",
            "UN Sustainable Development Goals Progress Report.",
        ],
    },
    "project_management": {
        "intro": "The United States recognizes effective project management as essential for delivering complex technology and infrastructure initiatives",
        "context": "Project managers with expertise in regulatory compliance, cross-functional coordination, and financial management are essential for banking, healthcare, and technology sectors.",
        "scarcity": "Professionals with combined expertise in project management, regulatory compliance, and technical implementation are in high demand.",
        "executive_orders": [],
        "sources": [
            "Project Management Institute, 'Talent Gap Report' (2024).",
            "U.S. Bureau of Labor Statistics, Project Management Occupational Outlook.",
        ],
    },
    # NEW FIELDS (added 2026-01-02 - from LOR analysis session)
    "nonprofit_finance": {
        "intro": "The United States nonprofit sector represents a $1 trillion economic force employing 12.5 million workers",
        "context": "Nonprofits deliver critical services in healthcare, education, social services, and community development. Effective financial management ensures mission sustainability and maximizes social impact.",
        "scarcity": "Finance professionals with nonprofit expertise who understand compliance, grant management, and mission-driven budgeting are in critically short supply.",
        "executive_orders": [],
        "sources": [
            "Independent Sector, 'Health of the Nonprofit Sector' (2024).",
            "Bureau of Labor Statistics, Nonprofit Employment Data.",
            "Nonprofit Finance Fund, 'State of the Nonprofit Sector' Report (2024).",
        ],
    },
    "transportation": {
        "intro": "The United States has identified transportation infrastructure as a critical national priority for economic competitiveness and public safety",
        "context": "The Infrastructure Investment and Jobs Act allocates $550 billion for transportation improvements. Efficient transportation systems are essential for economic competitiveness, supply chain resilience, and quality of life.",
        "scarcity": "Transportation engineers and logistics professionals with expertise in multimodal systems, fleet electrification, and supply chain optimization are in high demand.",
        "executive_orders": [
            "Infrastructure Investment and Jobs Act (2021).",
        ],
        "sources": [
            "U.S. Department of Transportation, Strategic Plan 2022-2026.",
            "American Society of Civil Engineers, Infrastructure Report Card (2024).",
            "Bureau of Transportation Statistics, National Transportation Statistics.",
        ],
    },
    "electric_vehicles": {
        "intro": "The United States has committed to electric vehicle adoption as part of the clean energy transition",
        "context": "The transition to electric vehicles addresses climate change, reduces dependence on foreign oil, and creates new manufacturing opportunities. Commercial vehicle electrification is particularly important for reducing fleet emissions.",
        "scarcity": "Engineers with expertise in EV systems, battery technology, charging infrastructure, and commercial vehicle electrification are critically needed.",
        "executive_orders": [
            "Executive Order 14008, 'Tackling the Climate Crisis at Home and Abroad' (January 27, 2021).",
        ],
        "sources": [
            "U.S. Department of Energy, National Blueprint for Transportation Decarbonization.",
            "Inflation Reduction Act of 2022, EV tax credits and incentives.",
            "Department of Energy, Alternative Fuels Data Center.",
        ],
    },
    "renewable_energy": {
        "intro": "The United States recognizes renewable energy development as essential for energy security and climate goals",
        "context": "Solar, wind, and other renewable sources are critical for achieving net-zero emissions and reducing dependence on fossil fuels. The clean energy sector is experiencing rapid growth and transformation.",
        "scarcity": "Engineers and scientists with expertise in solar technology, grid integration, and renewable energy systems are in extremely high demand.",
        "executive_orders": [
            "Executive Order 14008, 'Tackling the Climate Crisis at Home and Abroad' (January 27, 2021).",
        ],
        "sources": [
            "U.S. Department of Energy, Solar Futures Study (2021).",
            "National Renewable Energy Laboratory, Renewable Energy Research.",
            "Inflation Reduction Act of 2022, clean energy provisions.",
        ],
    },
}


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def set_font(run, font_name: str, font_size: int):
    """Apply font styling to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.insert(0, rFonts)


def validate_content(text: str, persona_config: dict) -> List[str]:
    """Validate content against forbidden patterns. Returns list of issues."""
    issues = []

    # Check for em-dash (Unicode character)
    if "\u2014" in text:
        issues.append("Content contains em-dash. Use regular dash (-) or rewrite.")

    # Check for forbidden phrases
    text_lower = text.lower()
    for phrase in persona_config.get("forbidden_phrases", []):
        if phrase.lower() in text_lower:
            issues.append(f"Content contains forbidden phrase: '{phrase}'")

    return issues


# =============================================================================
# PRONG 3 DETRIMENT ANALYSIS TEMPLATES
# =============================================================================

PRONG3_DETRIMENT_TEMPLATES = {
    "cybersecurity": "Requiring {name} to undergo the lengthy labor certification process would delay critical cybersecurity contributions for 12 to 18 months. Given that cyber threats evolve daily and attack sophistication increases constantly, this delay would directly undermine the national security priorities established in Executive Order 14028. The United States cannot afford to wait while adversaries continue to develop new attack vectors.",
    "ai_ml": "A traditional labor certification for {name} would take 12 to 24 months, during which the United States would lose ground in the global AI race. China and other competitors continue to advance rapidly in artificial intelligence capabilities. Any delay in deploying top AI talent compromises American technological leadership and economic competitiveness.",
    "clean_energy": "The labor certification process would delay {name}'s contributions to clean energy innovation for over a year. With climate change accelerating and international competition for green technology leadership intensifying, the United States cannot afford such delays. The national interest in energy independence and environmental sustainability requires immediate access to specialized talent.",
    "stem_education": "Requiring {name} to wait 12 to 18 months for labor certification would deprive American students of quality STEM instruction during a critical shortage. Every semester without qualified educators compounds the national skills gap. The immediate need for STEM teachers in underserved communities makes traditional labor certification contrary to the national interest.",
    "biotech": "The labor certification timeline would delay {name}'s biotechnology research for over a year, potentially allowing foreign competitors to achieve breakthroughs first. In a field where timing determines patent priority and market leadership, such delays compromise American competitiveness in the bioeconomy.",
    "semiconductor": "Requiring {name} to complete labor certification would delay critical semiconductor expertise for 12 to 18 months. Given the national security implications of chip technology and the current supply chain vulnerabilities, the United States cannot afford to wait. The CHIPS Act specifically recognizes the urgency of building domestic semiconductor talent.",
    "healthcare": "The labor certification process would delay {name}'s healthcare contributions for over a year, leaving patients without access to specialized care. In underserved communities facing severe provider shortages, such delays directly harm public health outcomes.",
    "quantum_computing": "Requiring {name} to undergo labor certification would delay quantum computing expertise for 12 to 18 months. In this nascent field where technological leadership is still being established, such delays could allow China and other competitors to achieve quantum advantage first, with serious implications for national security and cryptography.",
    "fintech": "A traditional labor certification for {name} would take over a year, during which the U.S. financial technology sector would lose competitive momentum. As global financial markets increasingly rely on innovative technology, delays in attracting fintech talent compromise American financial sector leadership.",
    "advanced_manufacturing": "The labor certification timeline would delay {name}'s manufacturing expertise for 12 to 18 months. With supply chain vulnerabilities exposed and reshoring urgently needed, such delays undermine the national interest in manufacturing resilience established in Executive Order 14017.",
    "space_technology": "Requiring {name} to complete labor certification would delay space technology contributions for over a year. As commercial space competition intensifies and space becomes increasingly strategic, the United States cannot afford to wait for specialized aerospace talent.",
    "default": "Requiring {name} to undergo the traditional labor certification process would delay contributions to the national interest for 12 to 18 months or longer. Given the documented shortage of professionals in this field and the urgent national priorities at stake, such delays would be contrary to the national interest that the NIW waiver is designed to serve.",
    # NEW PRONG 3 TEMPLATES (added 2026-01-02 - from real LOR analysis)
    "hvac": "Requiring {name} to complete labor certification would delay critical HVAC expertise for 12 to 18 months. Given that HVAC systems are essential infrastructure for healthcare facilities, data centers, and commercial buildings, such delays would undermine public health and energy efficiency goals. The demand for professionals who can design and implement HVAC systems capable of sustaining performance in extreme operational scenarios is urgent.",
    "procurement": "The labor certification timeline would delay {name}'s procurement expertise for over a year. As U.S. organizations strengthen their supply chains and develop fraud-aware procurement practices, such delays compromise resilience and competitiveness. {name}'s ability to deploy data-driven procurement analytics and vendor-risk rigor serves the national interest more effectively than requiring a traditional labor certification tied to one role.",
    "agriculture": "Requiring {name} to undergo labor certification would delay agricultural research contributions for 12 to 18 months. With climate change intensifying and food security challenges growing, the United States cannot afford such delays. {name}'s expertise in sustainable agriculture and climate-smart practices addresses challenges of global significance aligned with U.S. national and international priorities.",
    "animal_science": "The labor certification process would delay {name}'s animal science expertise for over a year, potentially allowing foreign competitors to achieve breakthroughs first. The demand for researchers who can improve livestock production while reducing environmental impact is urgent. {name}'s work advances critical themes in climate-smart agriculture and food security.",
    "project_management": "Requiring {name} to complete labor certification would delay critical project management expertise for 12 to 18 months. Given the complexity of regulatory compliance projects in banking, healthcare, and technology sectors, such delays would undermine organizational resilience and competitiveness. {name}'s ability to deliver time-sensitive regulatory projects serves the national interest.",
}


def get_prong3_detriment_paragraph(field: str, beneficiary_name: str) -> str:
    """
    Generate a Prong 3 paragraph explaining what the US LOSES if the petition is denied.

    This addresses the "detriment" aspect of the Matter of Dhanasar analysis,
    explaining why it would be against the national interest to require labor certification.

    Args:
        field: The beneficiary's field of expertise
        beneficiary_name: The name of the beneficiary

    Returns:
        A paragraph suitable for inclusion in an LOR
    """
    template = PRONG3_DETRIMENT_TEMPLATES.get(field, PRONG3_DETRIMENT_TEMPLATES["default"])
    return template.format(name=beneficiary_name)


# =============================================================================
# LOR GENERATOR CLASS
# =============================================================================

class LORGenerator:
    """
    Generate Letters of Recommendation with persona-specific formatting.

    Example usage:
        generator = LORGenerator(
            persona="executive",
            beneficiary_name="John Doe",
            visa_type="EB-2 NIW",
            field="cybersecurity",
        )

        filepath = generator.create_document(
            recommender_name="Jane Smith",
            recommender_title="VP of Engineering",
            recommender_org="TechCorp Inc.",
            recommender_email="jane@techcorp.com",
            relationship="supervisor",
            paragraphs=[...],
        )
    """

    def __init__(
        self,
        persona: PersonaType,
        beneficiary_name: str,
        visa_type: str = "EB-2 NIW",
        field: Optional[str] = None,
        output_dir: str = "output",
    ):
        self.persona = persona
        self.config = PERSONAS[persona]
        self.beneficiary_name = beneficiary_name
        self.visa_type = visa_type
        self.field = field
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_document(
        self,
        recommender_name: str,
        recommender_title: str,
        recommender_org: str,
        recommender_email: str,
        relationship: str,
        paragraphs: List[str],
        recommender_phone: Optional[str] = None,
        recommender_address: Optional[str] = None,
        include_footnotes: bool = True,
        custom_footnotes: Optional[List[str]] = None,
        date: Optional[str] = None,
        use_alt_salutation: bool = False,
        use_alt_closing: bool = False,
        custom_salutation: Optional[str] = None,
        custom_closing: Optional[str] = None,
    ) -> str:
        """
        Create a complete LOR document.

        Args:
            recommender_name: Full name of recommender
            recommender_title: Title/position
            recommender_org: Organization name
            recommender_email: Email address
            relationship: How recommender knows beneficiary
            paragraphs: List of paragraph texts (main content)
            recommender_phone: Optional phone number
            recommender_address: Optional address
            include_footnotes: Whether to add footnotes section
            custom_footnotes: Custom footnotes to add
            date: Optional date string, defaults to today
            use_alt_salutation: Use alternate salutation
            use_alt_closing: Use alternate closing

        Returns:
            Path to generated document
        """
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        font_name = self.config["font"]
        font_size = self.config["size"]
        use_tabs = self.config["use_tabs"]

        # =====================================================================
        # HEADER - Recommender info
        # =====================================================================
        header = doc.add_paragraph()
        run = header.add_run(f"{recommender_name}\n")
        run.bold = True
        set_font(run, font_name, font_size)

        run = header.add_run(f"{recommender_title}\n")
        set_font(run, font_name, font_size)

        run = header.add_run(f"{recommender_org}\n")
        set_font(run, font_name, font_size)

        if recommender_address:
            run = header.add_run(f"{recommender_address}\n")
            set_font(run, font_name, font_size)

        run = header.add_run(f"{recommender_email}\n")
        set_font(run, font_name, font_size)

        if recommender_phone:
            run = header.add_run(f"{recommender_phone}\n")
            set_font(run, font_name, font_size)

        run = header.add_run(f"\n{date or datetime.now().strftime('%B %d, %Y')}")
        set_font(run, font_name, font_size)

        # =====================================================================
        # SALUTATION
        # =====================================================================
        doc.add_paragraph()
        salutation = doc.add_paragraph()
        if custom_salutation:
            salutation_text = custom_salutation
        elif use_alt_salutation:
            salutation_text = self.config.get("alt_salutation")
        else:
            salutation_text = self.config["salutation"]
        run = salutation.add_run(salutation_text)
        run.bold = True
        set_font(run, font_name, font_size)

        # =====================================================================
        # BODY PARAGRAPHS
        # =====================================================================
        for para_text in paragraphs:
            # Validate content
            issues = validate_content(para_text, self.config)
            if issues:
                logger.warning(f"{issues}")

            doc.add_paragraph()
            p = doc.add_paragraph()

            if use_tabs:
                p.paragraph_format.first_line_indent = Inches(0.5)

            run = p.add_run(para_text)
            set_font(run, font_name, font_size)

            # Justify long paragraphs
            if len(para_text) > 100:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            p.paragraph_format.line_spacing = 1.15

        # =====================================================================
        # CLOSING
        # =====================================================================
        doc.add_paragraph()
        closing = doc.add_paragraph()
        if custom_closing:
            closing_text = custom_closing
        elif use_alt_closing:
            closing_text = self.config.get("alt_closing")
        else:
            closing_text = self.config["closing"]
        run = closing.add_run(f"{closing_text}\n\n\n")
        set_font(run, font_name, font_size)

        # Signature
        run = closing.add_run(f"{recommender_name}\n")
        run.bold = True
        set_font(run, font_name, font_size)

        run = closing.add_run(f"{recommender_title}\n")
        set_font(run, font_name, font_size)

        run = closing.add_run(recommender_org)
        set_font(run, font_name, font_size)

        # =====================================================================
        # FOOTNOTES
        # =====================================================================
        if include_footnotes:
            self._add_footnotes(doc, custom_footnotes)

        # =====================================================================
        # SAVE
        # =====================================================================
        safe_name = "".join(c for c in self.beneficiary_name if c.isalnum() or c in " _-").replace(" ", "_")
        safe_recommender = "".join(c for c in recommender_name if c.isalnum() or c in " _-").replace(" ", "_")
        filename = f"LOR_{safe_name}_{safe_recommender}_{self.persona}.docx"
        filepath = os.path.join(self.output_dir, filename)

        doc.save(filepath)
        logger.info(f"Generated: {filepath}")
        return filepath

    def _add_footnotes(self, doc: Document, custom_footnotes: Optional[List[str]] = None):
        """Add footnotes section."""
        font_name = self.config["font"]

        footnotes = []

        # Add Matter of Dhanasar
        footnotes.append(
            "Matter of Dhanasar, 26 I&N Dec. 884 (AAO 2016), establishing the three-prong framework for National Interest Waiver petitions."
        )

        # Add field-specific sources
        if self.field and self.field in NATIONAL_INTEREST_TEXTS:
            field_data = NATIONAL_INTEREST_TEXTS[self.field]
            for eo in field_data.get("executive_orders", []):
                footnotes.append(eo)
            for source in field_data.get("sources", []):
                footnotes.append(source)

        # Add custom footnotes
        if custom_footnotes:
            footnotes.extend(custom_footnotes)

        if not footnotes:
            return

        # Add separator
        doc.add_paragraph()
        p_sep = doc.add_paragraph()
        run = p_sep.add_run("_" * 50)
        set_font(run, font_name, 10)

        # Add footnotes
        for i, footnote in enumerate(footnotes, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {footnote}")
            run.font.size = Pt(9)
            run.font.italic = True
            set_font(run, font_name, 9)
            p.paragraph_format.left_indent = Inches(0.25)

    def get_prong3_paragraph(self) -> str:
        """Generate a Prong 3 (labor certification waiver) paragraph."""
        field_data = NATIONAL_INTEREST_TEXTS.get(self.field, {})
        scarcity = field_data.get("scarcity", "qualified professionals in this field are in critically short supply")

        templates = [
            f"Critically, I believe it would benefit the United States to waive the labor certification requirement for {self.beneficiary_name}. The traditional labor certification process is designed to protect American workers, but in this specialized field, {scarcity.lower()}. Requiring labor certification would be counterproductive when the nation needs experts like {self.beneficiary_name} contributing immediately. As the Administrative Appeals Office recognized in Matter of Dhanasar, when a petitioner's work benefits the nation beyond the immediate scope of their employment, waiving labor certification serves the national interest.",
            f"From my professional perspective, waiving the labor certification requirement for {self.beneficiary_name} would serve the national interest. {scarcity} The standard labor certification process would delay contributions that are urgently needed. {self.beneficiary_name}'s work has implications that extend far beyond any single employer, benefiting the broader national effort in this critical area. As established in Matter of Dhanasar, such broader national benefit justifies waiving the labor certification requirement.",
            f"I believe the national interest would be well served by waiving the labor certification requirement for {self.beneficiary_name}. {scarcity}, and the bureaucratic delays of traditional labor certification would be counterproductive. {self.beneficiary_name}'s contributions address urgent national priorities that transcend the typical employer-employee relationship. Under the framework established in Matter of Dhanasar, this type of broader national benefit supports a waiver of the labor certification requirement.",
        ]

        return random.choice(templates)

    def get_national_importance_paragraph(self) -> str:
        """Generate a national importance paragraph based on field."""
        if not self.field or self.field not in NATIONAL_INTEREST_TEXTS:
            return ""

        field_data = NATIONAL_INTEREST_TEXTS[self.field]

        return f"{field_data['intro']}. {field_data['context']} {field_data['scarcity']}"

    def get_opening_paragraph(self, relationship: str, years_known: str = "several") -> str:
        """Generate persona-appropriate opening paragraph."""
        name = self.beneficiary_name

        openings = {
            "executive": [
                f"In over twenty years of leadership, I have worked with professionals across the industry. {name} represents the caliber of talent that American companies need. As {relationship}, I offer my strong support for this petition.",
                f"As {relationship} to {name}, I have directly observed exceptional contributions to our organization. The results speak for themselves.",
            ],
            "technical": [
                f"Having worked with {name} as their {relationship}, I can attest to significant technical expertise and innovative problem-solving capabilities that I have rarely seen in my career.",
                f"From an engineering standpoint, {name}'s technical contributions during our {years_known} years of collaboration have been substantial and measurable.",
            ],
            "academic": [
                f"As a faculty member specializing in this field, I have supervised many students and researchers. {name} stands out as an exceptional practitioner-scholar. I submit this letter in support of the petition.",
                f"During my time as {name}'s {relationship}, I witnessed firsthand remarkable intellectual capacity and dedication to advancing knowledge in the field.",
            ],
            "mentor": [
                f"I first met {name} {years_known} years ago. Since then, I have had the privilege of observing professional growth and development that has been truly remarkable.",
                f"Through my role as {name}'s {relationship}, I have watched consistent professional development and meaningful contributions to our field.",
            ],
            "corporate": [
                f"In my capacity as {relationship} to {name}, I have consistently observed outstanding professional capabilities and contributions over {years_known} years.",
                f"From a project management perspective, {name}'s work during our collaboration demonstrated exceptional coordination and measurable results.",
            ],
        }

        return random.choice(openings.get(self.persona, openings["corporate"]))

    def generate_full_lor(
        self,
        recommender_name: str,
        recommender_title: str,
        recommender_org: str,
        recommender_email: str,
        relationship: str,
        years_known: str = "several years",
        custom_content: Optional[Dict[str, str]] = None,
        recommender_credentials: Optional[Dict[str, str]] = None,
        work_achievements: Optional[Dict[str, str]] = None,
    ) -> List[str]:
        """
        Generate a complete LOR with expanded content.
        Target: 600-900 words.

        Args:
            recommender_name: Full name of recommender
            recommender_title: Title/position
            recommender_org: Organization
            recommender_email: Email
            relationship: Type of relationship (supervisor, mentor, colleague, independent)
            years_known: Duration of relationship
            custom_content: Optional dict with custom content for sections
            recommender_credentials: Optional dict with recommender background info
            work_achievements: Optional dict with beneficiary's specific achievements

        Returns:
            List of paragraphs ready for create_document()
        """
        custom_content = custom_content or {}
        recommender_credentials = recommender_credentials or {}
        work_achievements = work_achievements or {}

        paragraphs = []

        # 1. OPENING (Support Declaration)
        opening = self.get_opening_paragraph(relationship, years_known)
        paragraphs.append(opening)

        # 2. RECOMMENDER CREDENTIALS
        creds_para = self._expand_recommender_credentials(
            recommender_name, recommender_title, recommender_org, recommender_credentials
        )
        if creds_para:
            paragraphs.append(creds_para)

        # 3. RELATIONSHIP CONTEXT
        relationship_para = self._expand_relationship_context(relationship, years_known, custom_content)
        if relationship_para:
            paragraphs.append(relationship_para)

        # 4. WORK DESCRIPTION / ACHIEVEMENTS
        work_para = self._expand_work_description(work_achievements, custom_content)
        if work_para:
            paragraphs.append(work_para)

        # 5. NATIONAL IMPORTANCE (Prong 1)
        national_para = self.get_national_importance_paragraph()
        if national_para:
            paragraphs.append(national_para)

        # 6. GOVERNMENT ALIGNMENT (Prong 1)
        gov_para = self._expand_government_alignment()
        if gov_para:
            paragraphs.append(gov_para)

        # 7. UNIQUE SKILLS (Prong 2)
        skills_para = self._expand_unique_skills(custom_content)
        if skills_para:
            paragraphs.append(skills_para)

        # 8. WAIVER JUSTIFICATION (Prong 3)
        waiver_para = self.get_prong3_paragraph()
        paragraphs.append(waiver_para)

        # 9. CONCLUSION
        conclusion_para = self._expand_conclusion(custom_content)
        paragraphs.append(conclusion_para)

        # Validate word count
        total_words = sum(len(p.split()) for p in paragraphs)
        if total_words < 600:
            # Add elaboration to the work description
            paragraphs = self._add_elaboration(paragraphs, 650)

        return paragraphs

    def _expand_recommender_credentials(
        self,
        recommender_name: str,
        recommender_title: str,
        recommender_org: str,
        credentials: Dict[str, str],
    ) -> str:
        """Expand recommender credentials section."""
        persona_templates = SECTION_TEMPLATES.get("recommender_credentials", {}).get(self.persona, [])
        if not persona_templates:
            persona_templates = SECTION_TEMPLATES.get("recommender_credentials", {}).get("corporate", [])

        if not persona_templates:
            return f"My name is {recommender_name}, and I am {recommender_title} at {recommender_org}."

        template = random.choice(persona_templates)

        # Fill with provided values or defaults
        defaults = {
            "recommender_name": recommender_name,
            "recommender_title": recommender_title,
            "recommender_org": recommender_org,
            "beneficiary_name": self.beneficiary_name,
            "degree": credentials.get("degree", "advanced degree"),
            "university": credentials.get("university", "a leading institution"),
            "specialty": credentials.get("specialty", self.field or "this field"),
            "years_experience": credentials.get("years_experience", "15"),
            "specific_area": credentials.get("specific_area", self.field or "this domain"),
            "field": self.field or "this field",
            "recognition": credentials.get("recognition", "recognition from industry peers"),
            "publications": credentials.get("publications", "numerous"),
            "additional_credentials": credentials.get("additional_credentials", "extensive professional experience"),
            "team_size": credentials.get("team_size", "50+"),
            "scope": credentials.get("scope", "multiple departments"),
            "responsibilities": credentials.get("responsibilities", "strategic planning and talent development"),
            "revenue_impact": credentials.get("revenue_impact", "significant"),
            "achievements": credentials.get("achievements", "leadership excellence"),
            "department_description": credentials.get("department_description", "key organizational functions"),
            "previous_companies": credentials.get("previous_companies", "leading organizations in this industry"),
            "technical_responsibilities": credentials.get("technical_responsibilities", "technical strategy and innovation"),
            "technical_credentials": credentials.get("technical_credentials", "relevant technical certifications"),
            "technical_field": credentials.get("technical_field", self.field or "this technical domain"),
            "patents_or_innovations": credentials.get("patents_or_innovations", "multiple innovations"),
            "scale_of_impact": credentials.get("scale_of_impact", "thousands of users"),
            "scope_of_work": credentials.get("scope_of_work", "critical technical initiatives"),
            "technical_areas": credentials.get("technical_areas", "core technical competencies"),
            "technical_achievements": credentials.get("technical_achievements", "significant technical projects"),
            "certifications": credentials.get("certifications", "industry-recognized certifications"),
            "notable_mentees_orgs": credentials.get("notable_mentees_orgs", "leading organizations in this field"),
            "org_description": credentials.get("org_description", "is a leader in this industry"),
        }

        try:
            return template.format(**defaults)
        except KeyError:
            return f"My name is {recommender_name}, and I am {recommender_title} at {recommender_org}. With extensive experience in {self.field or 'this field'}, I am qualified to assess professionals in this domain."

    def _expand_relationship_context(
        self,
        relationship: str,
        years_known: str,
        custom_content: Dict[str, str],
    ) -> str:
        """Expand relationship context section."""
        relationship_type = "supervisor" if "supervis" in relationship.lower() else \
                           "mentor" if "mentor" in relationship.lower() else \
                           "colleague" if "colleague" in relationship.lower() or "peer" in relationship.lower() else \
                           "independent"

        templates = SECTION_TEMPLATES.get("relationship_context", {}).get(relationship_type, [])
        if not templates:
            templates = SECTION_TEMPLATES.get("relationship_context", {}).get("colleague", [])

        if not templates:
            return f"I have known {self.beneficiary_name} for {years_known} in a professional capacity."

        template = random.choice(templates)

        # Determine pronoun
        he_she = custom_content.get("pronoun", "they")
        his_her = "his" if he_she == "he" else "her" if he_she == "she" else "their"
        him_her = "him" if he_she == "he" else "her" if he_she == "she" else "them"

        defaults = {
            "beneficiary_name": self.beneficiary_name,
            "duration": years_known,
            "he_she": he_she,
            "his_her": his_her,
            "him_her": him_her,
            "projects_description": custom_content.get("projects_description", "key strategic initiatives"),
            "specific_skills": custom_content.get("specific_skills", "technical excellence and leadership"),
            "start_date": custom_content.get("start_date", "the beginning of our collaboration"),
            "year": custom_content.get("year", "several years ago"),
            "meeting_context": custom_content.get("meeting_context", "our professional paths first crossed"),
            "mentorship_activities": custom_content.get("mentorship_activities", "regular guidance and feedback sessions"),
            "organization": custom_content.get("organization", "our organization"),
            "how_familiar": custom_content.get("how_familiar", "reviewing publications and conference presentations"),
            "context": custom_content.get("context", "professional channels"),
            "specific_exposure": custom_content.get("specific_exposure", "professional publications and industry recognition"),
        }

        try:
            return template.format(**defaults)
        except KeyError:
            return f"I have known {self.beneficiary_name} for {years_known} in a professional capacity that has given me direct insight into their capabilities."

    def _expand_work_description(
        self,
        work_achievements: Dict[str, str],
        custom_content: Dict[str, str],
    ) -> str:
        """Expand work description section."""
        work_type = work_achievements.get("type", "technical")
        templates = SECTION_TEMPLATES.get("work_description", {}).get(work_type, [])
        if not templates:
            templates = SECTION_TEMPLATES.get("work_description", {}).get("technical", [])

        if not templates:
            return custom_content.get("custom_paragraphs", "")

        template = random.choice(templates)

        he_she = custom_content.get("pronoun", "they")
        his_her = "his" if he_she == "he" else "her" if he_she == "she" else "their"
        He_She = he_she.capitalize()
        His_Her = his_her.capitalize()

        defaults = {
            "beneficiary_name": self.beneficiary_name,
            "he_she": he_she,
            "his_her": his_her,
            "him_her": "him" if he_she == "he" else "her" if he_she == "she" else "them",
            "He_She": He_She,
            "His_Her": His_Her,
            "technical_area": work_achievements.get("technical_area", self.field or "this technical domain"),
            "technical_achievement_1": work_achievements.get("achievement_1", "innovative solutions"),
            "impact_1": work_achievements.get("impact_1", "significantly improved efficiency"),
            "technical_achievement_2": work_achievements.get("achievement_2", "advanced methodologies"),
            "impact_2": work_achievements.get("impact_2", "enhanced organizational capabilities"),
            "scale": work_achievements.get("scale", "the organization"),
            "measurable_outcome": work_achievements.get("measurable_outcome", "produced measurable improvements"),
            "organization": work_achievements.get("organization", "the organization"),
            "responsibilities": work_achievements.get("responsibilities", "key technical initiatives"),
            "project_1": work_achievements.get("project_1", "a critical project"),
            "outcome_1": work_achievements.get("outcome_1", "delivered exceptional results"),
            "challenge": work_achievements.get("challenge", "complex technical challenges"),
            "project_2": work_achievements.get("project_2", "an important initiative"),
            "outcome_2": work_achievements.get("outcome_2", "achieved significant improvements"),
            "stakeholders": work_achievements.get("stakeholders", "key stakeholders"),
            "field": self.field or "this field",
            "research_topic": work_achievements.get("research_topic", "critical research areas"),
            "publications": work_achievements.get("publications", "peer-reviewed journals"),
            "citation_count": work_achievements.get("citation_count", "numerous"),
            "methodology_description": work_achievements.get("methodology_description", "novel approaches"),
            "institutions": work_achievements.get("institutions", "leading research institutions"),
            "applications": work_achievements.get("applications", "practical applications"),
            "team_size": work_achievements.get("team_size", "a dedicated team"),
            "project": work_achievements.get("project", "strategic initiatives"),
            "outcomes": work_achievements.get("outcomes", "exceeded expectations"),
            "achievements": work_achievements.get("achievements", "notable accomplishments"),
            "specific_achievement": work_achievements.get("specific_achievement", "significant milestones"),
        }

        try:
            result = template.format(**defaults)
            # Append custom paragraphs if provided
            if custom_content.get("custom_paragraphs"):
                result += " " + custom_content["custom_paragraphs"]
            return result
        except KeyError:
            return custom_content.get("custom_paragraphs", f"{self.beneficiary_name} has made significant contributions in {self.field or 'this field'}.")

    def _expand_government_alignment(self) -> str:
        """Generate government alignment paragraph."""
        if not self.field or self.field not in NATIONAL_INTEREST_TEXTS:
            return ""

        field_data = NATIONAL_INTEREST_TEXTS[self.field]
        eos = field_data.get("executive_orders", [])
        sources = field_data.get("sources", [])

        if not eos and not sources:
            return ""

        paragraphs = []

        if eos:
            paragraphs.append(f"{self.beneficiary_name}'s work aligns with federal priorities established in {eos[0]}")

        if sources:
            paragraphs.append(f"As documented in {sources[0]}, the national need in this area is significant.")

        return " ".join(paragraphs)

    def _expand_unique_skills(self, custom_content: Dict[str, str]) -> str:
        """Expand unique skills section."""
        templates = SECTION_TEMPLATES.get("unique_skills", [])
        if not templates:
            return ""

        template = random.choice(templates)

        he_she = custom_content.get("pronoun", "they")
        his_her = "his" if he_she == "he" else "her" if he_she == "she" else "their"
        him_her = "him" if he_she == "he" else "her" if he_she == "she" else "them"
        He_She = he_she.capitalize()

        defaults = {
            "beneficiary_name": self.beneficiary_name,
            "he_she": he_she,
            "his_her": his_her,
            "him_her": him_her,
            "He_She": He_She,
            "skill_1": custom_content.get("skill_1", "deep technical expertise"),
            "skill_2": custom_content.get("skill_2", "practical implementation experience"),
            "unique_capability": custom_content.get("unique_capability", "bridge theory and practice effectively"),
            "bridge_capability": custom_content.get("bridge_capability", "translate complex concepts into actionable solutions"),
            "what_is_needed": custom_content.get("what_is_needed", "advance national priorities in this field"),
            "distinguishing_factor": custom_content.get("distinguishing_factor", "a unique combination of skills"),
            "common_skill": custom_content.get("common_skill", "technical proficiency"),
            "rare_skill": custom_content.get("rare_skill", "strategic vision"),
            "background_1": custom_content.get("background_1", "academic research"),
            "background_2": custom_content.get("background_2", "industry practice"),
        }

        try:
            return template.format(**defaults)
        except KeyError:
            return f"{self.beneficiary_name} possesses a unique combination of skills that distinguishes them in this field."

    def _expand_conclusion(self, custom_content: Dict[str, str]) -> str:
        """Expand conclusion section."""
        templates = SECTION_TEMPLATES.get("conclusion", [])
        if not templates:
            return f"I strongly support {self.beneficiary_name}'s petition and believe approval would serve the national interest."

        template = random.choice(templates)

        his_her = custom_content.get("his_her", "their")
        His_Her = his_her.capitalize()

        defaults = {
            "beneficiary_name": self.beneficiary_name,
            "his_her": his_her,
            "His_Her": His_Her,
            "field": self.field or "this field",
            "key_achievements": custom_content.get("key_achievements", "exceptional contributions to this field"),
            "beneficiaries": custom_content.get("beneficiaries", "many stakeholders"),
        }

        try:
            return template.format(**defaults)
        except KeyError:
            return f"I strongly support {self.beneficiary_name}'s petition and believe approval would serve the national interest."

    def _add_elaboration(self, paragraphs: List[str], target_words: int) -> List[str]:
        """Add elaboration to reach target word count."""
        current_words = sum(len(p.split()) for p in paragraphs)
        if current_words >= target_words:
            return paragraphs

        # Add additional context paragraphs
        additional = []

        if self.field and self.field in NATIONAL_INTEREST_TEXTS:
            field_data = NATIONAL_INTEREST_TEXTS[self.field]
            if field_data.get("sources"):
                additional.append(
                    f"The importance of {self.beneficiary_name}'s work is underscored by authoritative sources. "
                    f"{field_data['sources'][0]} This documentation confirms the national priority of the work "
                    f"that {self.beneficiary_name} is advancing."
                )

        if additional:
            # Insert before conclusion
            paragraphs = paragraphs[:-1] + additional + [paragraphs[-1]]

        return paragraphs


# =============================================================================
# CONVENIENCE FUNCTION
# =============================================================================

def generate_lor(
    persona: PersonaType,
    beneficiary_name: str,
    recommender_name: str,
    recommender_title: str,
    recommender_org: str,
    recommender_email: str,
    relationship: str,
    paragraphs: List[str],
    visa_type: str = "EB-2 NIW",
    field: Optional[str] = None,
    output_dir: str = "output",
    **kwargs,
) -> str:
    """
    High-level function to generate a LOR.

    Args:
        persona: One of executive, technical, academic, mentor, corporate
        beneficiary_name: Name of person being recommended
        recommender_name: Name of recommender
        recommender_title: Title of recommender
        recommender_org: Organization of recommender
        recommender_email: Email of recommender
        relationship: How recommender knows beneficiary
        paragraphs: Content paragraphs for the letter
        visa_type: Type of visa (default EB-2 NIW)
        field: Field for national interest (cybersecurity, ai_ml, etc.)
        output_dir: Directory for output files
        **kwargs: Additional arguments passed to create_document

    Returns:
        Path to generated document
    """
    generator = LORGenerator(
        persona=persona,
        beneficiary_name=beneficiary_name,
        visa_type=visa_type,
        field=field,
        output_dir=output_dir,
    )

    return generator.create_document(
        recommender_name=recommender_name,
        recommender_title=recommender_title,
        recommender_org=recommender_org,
        recommender_email=recommender_email,
        relationship=relationship,
        paragraphs=paragraphs,
        **kwargs,
    )


# =============================================================================
# CLI INTERFACE
# =============================================================================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="CaseHub Letter of Recommendation Generator")
    parser.add_argument("--persona", choices=PERSONAS.keys(), required=True, help="Persona type")
    parser.add_argument("--beneficiary", required=True, help="Beneficiary name")
    parser.add_argument("--recommender", required=True, help="Recommender name")
    parser.add_argument("--title", required=True, help="Recommender title")
    parser.add_argument("--org", required=True, help="Recommender organization")
    parser.add_argument("--email", required=True, help="Recommender email")
    parser.add_argument("--relationship", required=True, help="Relationship to beneficiary")
    parser.add_argument("--visa-type", default="EB-2 NIW", help="Visa type")
    parser.add_argument("--field", choices=NATIONAL_INTEREST_TEXTS.keys(), help="Field for national interest")
    parser.add_argument("--output", default="output", help="Output directory")

    args = parser.parse_args()

    generator = LORGenerator(
        persona=args.persona,
        beneficiary_name=args.beneficiary,
        visa_type=args.visa_type,
        field=args.field,
        output_dir=args.output,
    )

    # Create sample paragraphs for demo
    sample_paragraphs = [
        generator.get_opening_paragraph(args.relationship),
        f"[Add credentials paragraph here describing recommender's qualifications]",
        f"[Add contributions paragraph here describing {args.beneficiary}'s specific achievements]",
    ]

    if args.field:
        sample_paragraphs.append(generator.get_national_importance_paragraph())
        sample_paragraphs.append(generator.get_prong3_paragraph())

    sample_paragraphs.append(f"I strongly support the approval of {args.beneficiary}'s {args.visa_type} petition.")

    filepath = generator.create_document(
        recommender_name=args.recommender,
        recommender_title=args.title,
        recommender_org=args.org,
        recommender_email=args.email,
        relationship=args.relationship,
        paragraphs=sample_paragraphs,
    )

    logger.info(f"Generated LOR: {filepath}")
    logger.info("Note: Replace bracketed sections with actual content.")
